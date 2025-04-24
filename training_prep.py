import pdfplumber
from docx import Document
import uuid
from io import BytesIO
from utilities import logger
from database import save_training_with_chunk
import asyncio
import re
from werkzeug.datastructures import FileStorage
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import tiktoken
import traceback

# Initialize the LLM client with timeout
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)  # 60 second timeout
FAST_LLM = ChatOpenAI(model="gpt-4o-mini", streaming=False, request_timeout=45)  # 45 second timeout

# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt,temperature=0.7, max_tokens=1000, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop, temperature=temperature, max_tokens=max_tokens)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

async def summarize_with_llm(text: str, max_length: int = 150) -> str:
    """Summarize text using the LLM, preserving key factual and actionable elements in the original language."""
    try:
        # Log the length of the text being summarized
        text_length = len(text)
        first_100_chars = text[:100].replace('\n', ' ') if text else "EMPTY TEXT"
        last_100_chars = text[-100:].replace('\n', ' ') if text_length > 100 else ""
        logger.debug(f"Summarizing text of length {text_length}. First 100 chars: '{first_100_chars}...'")
        if last_100_chars:
            logger.debug(f"Last 100 chars: '...{last_100_chars}'")
        
        if not text or text_length < 50:
            logger.warning(f"Text is too short or empty for summarization: '{text}'")
            return "Text provided is too short or empty for meaningful summarization."
        
        # Check if text appears to be a placeholder or metadata only
        if text.count('\n') < 3 and "Document:" in text and "Type:" in text:
            logger.warning(f"Text appears to be metadata only, no content: '{text}'")
            return "The provided text appears to contain only metadata without substantial content to summarize."
        
        # Detect language to explicitly tell the model to maintain it
        # This is a simple detection - the LLM will handle the actual language preservation    
        prompt = (
            f"You are a precise document summarizer specialized in knowledge extraction and preserving cultural and business nuance.\n\n"
            f"TASK: Create a concise summary of the following text in 250 words or less. Your summary MUST CAPTURE ALL critical information, instructions, and key points.\n\n"
            f"GUIDELINES:\n"
            f"1. Extract the most important FACTUAL information (e.g., names, numbers, dates).\n"
            f"2. Prioritize actionable knowledge (e.g., procedures, requirements, steps).\n"
            f"3. Preserve EXACT PHRASING for scripted instructions, questions, and sales prompts.\n"
            f"4. Maintain the EXACT ORDER of steps and processes as in the original text.\n"
            f"5. IMPORTANT: You MUST maintain the EXACT original language of the document. If the document is in Vietnamese or any non-English language, the summary MUST be in that same language. DO NOT translate to English under any circumstances.\n"
            f"6. Preserve all technical terms, proper names, and domain-specific vocabulary exactly as written.\n"
            f"7. Use clear, professional language without filler words.\n"
            f"8. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
            f"9. INCLUDE real-life lessons, practical advice, and applied knowledge from the text.\n"
            f"10. Maintain the emotional tone and perspective of the original document.\n"
            f"11. DO NOT MISS any critical instructions, warnings, exceptions, or conditional cases.\n"
            f"12. Include ALL numerical thresholds, specific quantities, and measurement criteria.\n"
            f"13. Preserve temporal information - when things should happen, timing requirements.\n"
            f"14. Capture conditional logic (if/then/else scenarios) and boundary conditions.\n"
            f"15. NEVER omit critical business process details that could impact operations.\n\n"
            f"OUTPUT FORMAT:\n"
            f"- Provide a clean, flowing summary in plain text without section markers or '# Section' or '## Subsection' headings\n"
            f"- Use natural paragraphs instead of formal section headers\n"
            f"- Organize content logically but do not add structural labels like 'Section:' or 'Subsection:'\n\n"
            f"IMPORTANT WARNING: Missing critical information in your summary risks business process failure. BE THOROUGH while staying concise.\n\n"
            f"TEXT TO SUMMARIZE:\n{text}"
        )
        
        # Log prompt length for debugging
        logger.debug(f"Full prompt length: {len(prompt)} chars")
        logger.debug(f"Summary prompt text length: {len(text)} characters")
        
        response = await invoke_llm_with_retry(LLM, prompt)
        summary = response.content.strip()
        logger.info(f"Generated summary (approx {len(summary.split())} words): {summary}")
        return summary
    except Exception as e:
        logger.error(f"LLM summarization failed: {e}")
        return "Summary unavailable due to processing error."

async def extract_knowledge_with_llm(text: str, domain: str = "", max_items: int = 10) -> str:
    """Extract structured knowledge elements from text using the LLM while preserving the original language."""
    try:
        # Log the length of the text being analyzed
        text_length = len(text)
        first_100_chars = text[:100].replace('\n', ' ') if text else "EMPTY TEXT"
        logger.debug(f"Extracting knowledge from text of length {text_length}. First 100 chars: '{first_100_chars}...'")
        
        if not text or text_length < 50:
            logger.warning(f"Text is too short or empty for knowledge extraction: '{text}'")
            return "Text provided is too short or empty for meaningful knowledge extraction."
        
        # Format domain context if provided
        domain_context = f" in {domain}" if domain else ""
             
        prompt = (
                f"You are a precise knowledge extractor specialized in identifying actionable information {domain_context} while preserving cultural context and nuance.\n\n"
                f"TASK: Extract key knowledge elements from the text below. Your extraction MUST be COMPREHENSIVE and THOROUGH, capturing ALL instructions, procedures, guidelines, and important points.\n\n"
                f"GUIDELINES:\n"
                f"1. Focus on FACTUAL and ACTIONABLE knowledge (e.g., procedures, requirements).\n"
                f"2. Include all specific questions, scripts, and sequential steps EXACTLY as written.\n"
                f"3. Include specific details (e.g., quantities, timelines, criteria).\n"
                f"4. IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all knowledge elements MUST be in that same language. DO NOT translate to English under any circumstances.\n"
                f"5. Maintain all technical terms, proper names, and specific vocabulary exactly as written.\n"
                f"6. Format each element as: 'KEY POINT: [concise statement]'.\n"
                f"7. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
                f"8. PRIORITIZE real-life lessons, practical advice, and applied knowledge.\n"
                f"9. Capture the underlying reasoning and wisdom, not just surface instructions.\n"
                f"10. Include both explicit statements and implied knowledge from the context.\n"
                f"11. Maintain the EXACT ORDER of steps and processes where applicable.\n"
                f"12. DO NOT miss ANY instructions, guidelines, rules, or critical information, no matter how minor it may seem.\n"
                f"13. Pay special attention to conditional statements (if/then/else), exceptions, and special cases.\n"
                f"14. Capture ALL numerical values, thresholds, measurements, and quantitative information.\n\n"
                f"IMPORTANT WARNING: Missing critical information risks business process failure. Be EXHAUSTIVE and THOROUGH in your extraction.\n\n"
                f"TEXT TO ANALYZE:\n{text}"
            )

        logger.debug(f"Knowledge extraction prompt length: {len(prompt)} chars")
        
        logger.debug(f"Prompt for knowledge extraction (first 200 chars): {prompt[:200]}...")
        logger.debug(f"Full text length for knowledge extraction: {len(text)}")
        
        response = await invoke_llm_with_retry(LLM, prompt)
        extracted_knowledge = response.content.strip()
        logger.info(f"Extracted {extracted_knowledge.count('KEY POINT')} knowledge elements")
        return extracted_knowledge
    except Exception as e:
        logger.error(f"LLM knowledge extraction failed: {e}")
        return "Knowledge extraction unavailable due to processing error."



async def refine_document(file: FileStorage = None, text: str = None, user_id: str = "", mode: str = "default", reformat_text: bool = False) -> tuple[bool, dict]:
    """
    Refine a document or text: extract content, summarize, extract knowledge, and optionally reformat as structured text.
    Returns (success: bool, result_dict: dict). Does not save to database.
    
    Args:
        file: Input file (DOCX or PDF), optional if text is provided.
        text: Raw text input, optional if file is provided.
        user_id: User ID for metadata.
        mode: Processing mode (e.g., "default", "pretrain").
        reformat_text: If True, return a structured text representation.
    """
    full_text = ""
    document_metadata = {}
    sections = []

    try:
        # Step 1: Extract text
        if text and text.strip():  # Ensure text is not just whitespace
            logger.debug(f"Processing text input of length {len(text)}")
            full_text = text
            document_metadata = {
                'filename': 'text_input',
                'file_type': 'TXT',
                'paragraph_count': len([p for p in text.split('\n') if p.strip()])
            }
            # Parse text sections (assuming possible Markdown or plain text)
            current_heading = "Document Start"
            current_section = []
            for line in text.split('\n'):
                if line.strip():
                    if line.startswith('# ') or line.startswith('## '):
                        if current_section:
                            sections.append((current_heading, "\n".join(current_section)))
                            current_section = []
                        current_heading = line.lstrip('# ').strip()
                    else:
                        current_section.append(line)
            if current_section:
                sections.append((current_heading, "\n".join(current_section)))
                
            # Verify that we have actual content not just metadata
            if len(sections) == 0 or (len(sections) == 1 and len(sections[0][1]) < 50):
                logger.warning(f"Insufficient text content: {text}")
                return False, {"error": "Insufficient text content."}
                
        elif file:
            file_extension = file.filename.split('.')[-1].lower()
            logger.debug(f"Processing file: {file.filename} ({file_extension})")
            
            # Read file content
            file_content = file.read()
            if not file_content:
                logger.warning(f"Empty file content for {file.filename}")
                return False, {"error": "Empty file content."}

            # Create a fresh BytesIO object for processing
            file_content_bytes = BytesIO(file_content)
            
            if file_extension == 'pdf':
                logger.debug("Opening PDF file")
                try:
                    with pdfplumber.open(file_content_bytes) as pdf:
                        document_metadata = {
                            'page_count': len(pdf.pages),
                            'filename': file.filename,
                            'file_type': 'PDF'
                        }
                        
                        # Extract text with page context
                        pages_text = []
                        for i, page in enumerate(pdf.pages):
                            page_text = page.extract_text() or ""
                            if page_text.strip():
                                logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                                pages_text.append(f"[Page {i+1}] {page_text}")
                            else:
                                logger.debug(f"PDF page {i+1}: empty or non-text content")
                        
                        # Join pages with clear separation
                        full_text = "\n\n".join(pages_text)
                        
                        if not full_text.strip():
                            logger.warning(f"No text extracted from {file.filename}")
                            return False, {"error": "No text extracted from PDF."}
                            
                        # Create sections from pages for better processing
                        sections = [(f"Page {i+1}", page) for i, page in enumerate(pages_text) if page.strip()]
                        logger.debug(f"Created {len(sections)} sections from PDF pages")
                except Exception as pdf_error:
                    logger.error(f"PDF processing error: {type(pdf_error).__name__}: {str(pdf_error)}")
                    return False, {"error": f"PDF processing error: {str(pdf_error)}"}
                    
            elif file_extension == 'docx':
                logger.debug("Opening DOCX file")
                try:
                    # Reset BytesIO position
                    file_content_bytes.seek(0)
                    doc = Document(file_content_bytes)
                    
                    # Extract document properties
                    document_metadata = {
                        'filename': file.filename,
                        'file_type': 'DOCX',
                        'paragraph_count': len(doc.paragraphs)
                    }
                    
                    # Extract text with structural context - similar to docuhandler approach
                    full_text_parts = []
                    current_heading = "Document Start"
                    current_section = []
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            if para.style.name.startswith('Heading'):
                                # Save current section if it exists
                                if current_section:
                                    sections.append((current_heading, "\n".join(current_section)))
                                    current_section = []
                                
                                current_heading = para.text.strip()
                                full_text_parts.append(f"\n[{current_heading}]\n{para.text}")
                            else:
                                current_section.append(para.text)
                                full_text_parts.append(para.text)
                    
                    # Don't forget the last section
                    if current_section:
                        sections.append((current_heading, "\n".join(current_section)))
                    
                    # Combine all the text parts
                    full_text = "\n".join(full_text_parts)
                    
                    if not full_text.strip():
                        logger.warning(f"No text extracted from {file.filename}")
                        return False, {"error": "No text extracted from DOCX."}
                except Exception as docx_error:
                    logger.error(f"DOCX processing error: {type(docx_error).__name__}: {str(docx_error)}")
                    return False, {"error": f"DOCX processing error: {str(docx_error)}"}
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return False, {"error": f"Unsupported file type: {file_extension}"}
        else:
            logger.error("No file or valid text provided")
            return False, {"error": "No file or valid text provided."}

        # Additional validation to ensure we have content to process
        if not full_text.strip():
            logger.warning("No text content extracted")
            return False, {"error": "No text content could be extracted."}
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            return False, {"error": "Text content too short for meaningful processing."}

        # Detect language for logging and validation
        sample_text = full_text[:200]
        is_english = all(ord(c) < 128 for c in sample_text.replace('\n', ' ').replace(' ', ''))
        detected_language = "English" if is_english else "Non-English (possibly multilingual)"
        logger.info(f"Detected language: {detected_language}")

        logger.info(f"Extracted text content of {len(full_text.split())} words from {document_metadata.get('filename', 'unknown')}")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")

        # Step 2: Prepare contextual information - DO NOT include user_id or mode in the context
        context_prefix = f"Document: {document_metadata.get('filename', 'unknown')}\nType: {document_metadata.get('file_type', 'Unknown')}\n"
        if document_metadata.get('file_type') == 'PDF' and document_metadata.get('page_count'):
            context_prefix += f"Pages: {document_metadata.get('page_count')}\n\n"
        else:
            context_prefix += "\n"  # Add a blank line after metadata
        
        # Combine content for summarization but don't include mode or user_id
        combined_text = context_prefix + full_text
        logger.debug(f"Prepared text for summarization: prefix length {len(context_prefix)} chars, content length {len(full_text)} chars")
        
        result = {
            "metadata": document_metadata,
            "text_length": len(full_text.split())
        }

        # Step 3: Reformat as structured text if requested
        if reformat_text and sections:
            async def reformat_document(sections: list[tuple[str, str]]) -> str:
                """Reformat document sections into structured text with meticulous preservation of all knowledge items."""
               
                # Format sections into a string
                formatted_sections = "\n".join([f"Section: {h}\nContent: {c}" for h, c in sections])
                prompt_OK = (
                        f"You are a Professional Copywriter and Real-Life Training Strategist. Your task is to refine the document so it becomes a high-quality, instructive, and emotionally intelligent training manual. This is designed for both AI learning and real human sales training, so accuracy, tone, and practical realism are critical.\n\n"

                        f"INPUT SECTIONS:\n"
                        f"{formatted_sections}\n\n"

                        f"YOUR TASK:\n"
                        f"1. PRESERVE STRATEGIC CONTENT — Retain all factual, instructional, and emotionally meaningful content from the original. Do not delete or simplify concrete data, intent, or emotional appeal.\n"
                        f"2. CLARIFY WITHOUT OVERSIMPLIFYING — Improve readability, structure, and sentence clarity, but retain every nuance and subtle logic from the original.\n"
                        f"3. PRESERVE AND SIGNAL ALL REAL-LIFE EXAMPLES — Retain every real-life example, quote, or scenario from the author. These MUST be clearly preserved and rendered naturally.\n"
                        f"4. FLAG AND ADD CONTEXT-MATCHED EXAMPLES — If a section lacks real examples, insert them using 'Suggested Example:' but only if the example **exactly matches the phase** of the customer journey (e.g., don't persuade during diagnosis).\n"
                        f"5. ADD INTENT-AWARE NOTES — Use 'Note:' to clarify why something matters, how it connects to human behavior, or how it relates to cultural context or emotional safety.\n"
                        f"6. PHASE-CORRECT ILLUSTRATION — All additions or clarifications must respect whether the section is in Discovery, Diagnosis, Motivation, Objection Handling, or Closing. Mismatched tones will break trust.\n"
                        f"7. MAINTAIN FLOW AND HIERARCHY — Ensure smooth transitions and clean structure: context → purpose → instruction → illustration → emotional reinforcement.\n"
                        f"8. MAINTAIN TEMPORAL CONTEXT — If the content includes results, testimonials, or feedback that may evolve, contextualize them as time-sensitive or dynamic.\n"
                        f"9. CRITICAL - PRESERVE ORIGINAL LANGUAGE — Your output MUST BE IN THE EXACT SAME LANGUAGE as the input. If the document is in Vietnamese, your output MUST be in Vietnamese. If it's in English, output must be in English. DO NOT TRANSLATE TO ENGLISH under any circumstances if the original is not in English. Use culturally native expressions of the original language only.\n"
                        f"10. STAY AUTHOR-CENTRIC — Preserve the author's original personality, beliefs, voice, and rhythm. Do not sterilize tone or local phrases. Style = identity.\n"
                        f"11. FORMAT USING MARKERS — Use this format for structured output:\n"
                        f"# Section Heading\n## Subsection Heading\nInstruction + Examples + Notes (if needed)\n\n"
                        f"FINAL CHECK: Before submitting your answer, verify again that you have maintained the original language. If the input is in Vietnamese, Thai, Indonesian, or any other non-English language, your output MUST be in that same language, not in English.\n\n"
                    )
                prompt = (
                    f"You are a Professional Copywriter and Real-Life Training Strategist tasked with refining a raw document into a high-quality, instructive, and emotionally intelligent training manual. This manual is for AI learning and human sales training, requiring accuracy, practical realism, and a tone that resonates with a professional audience. The output must be polished, readable, and culturally authentic.\n\n"
                    f"INPUT SECTIONS:\n"
                    f"{formatted_sections}\n\n"
                    f"YOUR TASK:\n"
                    f"1. PRESERVE ALL CONTENT — Retain every factual, instructional, and emotionally meaningful detail from the original. Do not delete or simplify data, intent, or emotional appeal.\n"
                    f"2. ENHANCE CLARITY AND CONCISENESS — Improve sentence clarity and readability using short sentences (15-20 words), active voice, and precise language. Avoid oversimplifying nuances or subtle logic.\n"
                    f"3. PRESERVE REAL-LIFE EXAMPLES — Retain all original examples, quotes, or scenarios, rendering them naturally and clearly labeled as 'Example:'.\n"
                    f"4. ADD CONTEXT-MATCHED EXAMPLES — For sections lacking examples, insert realistic ones labeled 'Suggested Example:'. Ensure examples match the customer journey phase (e.g., Discovery, Diagnosis, Motivation, Objection Handling, Closing) to maintain tone and trust.\n"
                    f"5. INCLUDE INTENT-AWARE NOTES — Add 'Note:' to explain why content matters, how it connects to human behavior, or its cultural/emotional significance. Notes should enhance instructional value.\n"
                    f"6. RESPECT PHASE-CORRECT TONE — Ensure all additions align with the section's customer journey phase. Mismatched tones (e.g., persuasive in Discovery) will break trust.\n"
                    f"7. ENSURE SMOOTH FLOW AND HIERARCHY — Structure content with context → purpose → instruction → illustration → emotional reinforcement. Use transitions for coherence.\n"
                    f"8. MAINTAIN TEMPORAL CONTEXT — Contextualize results, testimonials, or feedback as time-sensitive (e.g., 'As of 2025, this approach yielded...').\n"
                    f"9. PRESERVE ORIGINAL LANGUAGE — Detect the input language and produce output in the EXACT SAME LANGUAGE. If the input is in Vietnamese, Thai, or any non-English language, the output MUST remain in that language. DO NOT TRANSLATE to English. Use culturally native expressions.\n"
                    f"10. PRESERVE AUTHOR'S VOICE — Maintain the author's personality, beliefs, tone, and rhythm. Do not sterilize local phrases or style, as style = identity.\n"
                    f"11. ENHANCE READABILITY WITH VISUAL AIDS — Use bullet points, numbered lists, or tables where appropriate to break up text and highlight key points.\n"
                    f"12. ENSURE CONSISTENT FORMATTING — Apply uniform headings, fonts, and spacing. Suggest a professional layout (e.g., 1.5 line spacing, clear section breaks).\n"
                    f"13. SUPPORT ITERATIVE REFINEMENT — Flag areas needing human review (e.g., 'Review Suggested Example for accuracy') to facilitate feedback-driven revisions.\n"
                    f"14. FORMAT OUTPUT — Use this structure:\n"
                    f"# Section Heading\n## Subsection Heading\nInstruction\nExample or Suggested Example\nNote (if applicable)\n\n"
                    f"FINAL CHECK:\n"
                    f"- Verify the output is in the same language as the input.\n"
                    f"- Ensure no critical content is lost.\n"
                    f"- Confirm examples and notes align with the section's phase and cultural context.\n"
                    f"- Double-check readability (short sentences, active voice, visual aids).\n\n"
                )
                logger.debug(f"Reformat document prompt length: {len(prompt)} chars")
                logger.debug(f"Number of sections to reformat: {len(sections)}")
                response = await invoke_llm_with_retry(LLM, prompt, temperature=0.2, max_tokens=10000)
                return response.content.strip()
            
            reformatted_text = await reformat_document(sections)
            result["reformatted_text"] = reformatted_text
            
            # Use the reformatted text for summary and knowledge extraction if available
            if reformatted_text and reformatted_text.strip():
                logger.info("Using reformatted text for summary and knowledge extraction")
                text_for_summary = context_prefix + reformatted_text
                text_for_knowledge = reformatted_text
            else:
                logger.info("Using original text for summary and knowledge extraction")
                text_for_summary = combined_text
                text_for_knowledge = full_text
        else:
            # Use original text if no reformatting was done
            text_for_summary = combined_text
            text_for_knowledge = full_text
            
        # Step 4: Run summarization and knowledge extraction on the appropriate text
        logger.info("Starting summarization and knowledge extraction tasks")
        summary_task = summarize_with_llm(text_for_summary)
        knowledge_task = extract_knowledge_with_llm(text_for_knowledge)
        summary, knowledge = await asyncio.gather(summary_task, knowledge_task)
        
        # Add summary and knowledge to result
        result["summary"] = summary
        result["knowledge_elements"] = knowledge
        
        logger.info(f"Processing complete for {document_metadata.get('filename', 'unknown')}")
        return True, result

    except Exception as e:
        logger.error(f"Failed to refine document: {type(e).__name__}: {str(e)}")
        return False, {"error": f"Processing error occurred: {str(e)}"}
    finally:
        if file:
            file.close()

async def enrich_key_points(key_point: str, full_text: str) -> str:
    """
    Enrich key points with supporting data from the original document text.
    
    This function scans through the full document text to find meaningful supporting 
    information, business nuances, practical examples, and contextual details for each key point.
    
    Args:
        key_point: String containing key points (format: "KEY POINT: text")
        full_text: The full document text to extract supporting data from
        
    Returns:
        Enriched text with both key points and supporting data
    """
    try:
                
        # Enhanced prompt with focus on business nuance and contextual information
        prompt = (
            f"You are a precise knowledge enhancer and business context analyzer tasked with gathering rich supporting data for key points from a document.\n\n"
            f"TASK: For the key point below, thoroughly scan the document to extract COMPREHENSIVE supporting data that includes:\n"
            f"1. Contextual information explaining WHY this point matters in the business context\n"
            f"2. Real-world examples or use cases demonstrating the point's application\n"
            f"3. Related business impact, risks, or opportunities\n"
            f"4. Surrounding knowledge that gives depth to understanding this point\n\n"
            f"GUIDELINES:\n"
            f"1. Extract supporting data VERBATIM from the original text, prioritizing exact scripts, questions, or instructions.\n"
            f"2. Include ALL relevant details (e.g., names, numbers, dates, sequential steps) tied to the key point, ensuring completeness.\n"
            f"3. DISCOVER BUSINESS NUANCE - identify specific business contexts, industry-specific implications, or organizational relevance.\n"
            f"4. FIND PRACTICAL APPLICATIONS - include any examples, case studies, or scenarios that demonstrate how this knowledge is applied.\n"
            f"5. IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all knowledge elements MUST be in that same language. DO NOT translate to English under any circumstances.\n"
            f"6. Preserve all technical terms, proper names, and domain-specific vocabulary exactly as written.\n"
            f"7. Maintain the EXACT ORDER of steps or processes as they appear in the original text.\n"
            f"8. Focus on MEANINGFUL CONTEXT that would help someone understand not just WHAT the key point is, but WHY it matters and HOW it's used.\n"
            f"9. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
            f"10. DO NOT add assumptions, enhancements, or inferred details beyond the original text.\n"
            f"11. IMPORTANT: The supporting data must be rich and comprehensive, providing all context needed to fully understand the key point.\n"
            f"12. NEVER omit critical instructions, conditions, exceptions, or warnings related to the key point.\n"
            f"13. Include ALL numerical thresholds, specific quantities, and measurement criteria exactly as stated.\n"
            f"14. Pay special attention to conditional logic (if/then/else scenarios) and boundary conditions.\n"
            f"15. Capture temporal information - when things should happen, sequence, timing requirements.\n\n"
            f"IMPORTANT WARNING: Missing critical information in supporting data risks business process failure. BE EXHAUSTIVE.\n\n"
            f"Output format: 'KEY POINT: [original key point]\nSUPPORTING DATA: [comprehensive verbatim text from document]'\n\n"
            f"KEY POINTS:\n{key_point}\n\n"
            f"ORIGINAL TEXT:\n{full_text}"
        )
        
        logger.debug(f"Enriching key points, prompt length: {len(prompt)} chars")
        
        # Call LLM
        response = await invoke_llm_with_retry(LLM, prompt)
        enriched_knowledge = response.content.strip()
        
        # Check if response looks like a refusal due to content moderation
        if ("I'm sorry" in enriched_knowledge or 
            "I apologize" in enriched_knowledge or 
            "I cannot" in enriched_knowledge) and "KEY POINT:" not in enriched_knowledge:
            logger.warning("LLM refused to process the content, using fallback")
            
            # Simple fallback approach: use key points as their own supporting data
            fallback_lines = []
            for line in key_point.split('\n'):
                if line.strip() and line.lower().startswith("key point:"):
                    colon_pos = line.find(':')
                    if colon_pos > 0:
                        key_point_text = line[colon_pos + 1:].strip()
                        fallback_lines.append(f"KEY POINT: {key_point_text}")
                        fallback_lines.append(f"SUPPORTING DATA: {key_point_text}")
            
            enriched_knowledge = '\n'.join(fallback_lines)
            logger.info(f"Created fallback enrichment with {len(fallback_lines)//2} key points")
        
        # Log success statistics
        key_point_count = enriched_knowledge.count("KEY POINT:")
        supporting_data_count = enriched_knowledge.count("SUPPORTING DATA:")
        logger.info(f"Enrichment contains {key_point_count} key points and {supporting_data_count} supporting data sections")
        
        return enriched_knowledge
    except Exception as e:
        logger.error(f"Key point enrichment failed: {e}")
        traceback_str = traceback.format_exc()
        logger.error(f"Traceback: {traceback_str}")
        
        # Return a simple error-state response that maintains the expected format
        return f"KEY POINT: {key_point.split('KEY POINT:')[1].strip() if 'KEY POINT:' in key_point else key_point}\nSUPPORTING DATA: Error occurred during enrichment."

async def process_document(text: str = "", file: FileStorage = None, user_id: str = "", mode: str = "default", bank: str = "", knowledge_elements: str = "") -> bool:
    """
    Process a document or text: extract key points from knowledge elements, enrich each with supporting data, 
    and save to Pinecone.
    
    Args:
        text: Raw text input, preferred input method.
        file: Input file (DOCX or PDF), optional for legacy support.
        user_id: User ID for metadata.
        mode: Processing mode (e.g., "default", "pretrain").
        bank: Namespace for Pinecone storage.
        knowledge_elements: Pre-extracted knowledge elements from refine_document function.
    """
    doc_id = str(uuid.uuid4())
    logger.info(f"Processing document with bank={bank}")
    logger.info(f"Knowledge elements received: {knowledge_elements}")

    try:
        # Step 1: Validate text content
        if not text or not text.strip():
            logger.error("No text content provided")
            return False
            
        logger.debug(f"Processing text input of length {len(text)}")
        full_text = text
        
        # Check for minimal content
        if len(text.split()) < 10:
            logger.warning(f"Text content too short: '{text}'")
            return False

        # Step 2: Extract key points from knowledge elements
        knowledge_elements_text = knowledge_elements
        logger.info(f"Knowledge elements provided: {len(knowledge_elements_text)} characters")
        
        # Check if knowledge elements were provided
        if not knowledge_elements_text or "KEY POINT" not in knowledge_elements_text:
            logger.warning("No valid knowledge elements provided, attempting to extract them")
            knowledge_elements_text = await extract_knowledge_with_llm(full_text)
            
            if not knowledge_elements_text or "KEY POINT" not in knowledge_elements_text:
                logger.warning("Failed to extract knowledge elements")
                return False
        
        # Step 3: Extract individual key points
        key_points = []
        
        # Handle numbered format (e.g., "1. KEY POINT: ...")
        if re.search(r'^\d+\.\s+KEY POINT:', knowledge_elements_text, re.MULTILINE):
            logger.info("Detected numbered list format in knowledge elements")
            for line in knowledge_elements_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                # Match both numbered and non-numbered KEY POINT lines
                if re.match(r'^\d+\.\s*KEY POINT:', line, re.IGNORECASE) or line.lower().startswith("key point:"):
                    # Clean up the line
                    clean_line = re.sub(r'^\d+\.\s*', '', line)  # Remove numbering if present
                    colon_pos = clean_line.find(':')
                    if colon_pos > 0:
                        key_point = clean_line[colon_pos + 1:].strip()
                        if key_point:  # Only add non-empty key points
                            key_points.append(key_point)
                            logger.debug(f"Extracted key point: {key_point[:50]}...")
        else:
            # Handle standard format (KEY POINT: ...)
            for line in knowledge_elements_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                if line.lower().startswith("key point:"):
                    colon_pos = line.find(':')
                    if colon_pos > 0:
                        key_point = line[colon_pos + 1:].strip()
                        if key_point:  # Only add non-empty key points
                            key_points.append(key_point)
                            logger.debug(f"Extracted key point: {key_point[:50]}...")
        
        logger.info(f"Extracted {len(key_points)} key points from knowledge elements")
        
        if not key_points:
            logger.warning("No key points extracted from knowledge elements")
            return False
            
        # Step 4: Enrich each key point individually and create chunks
        key_point_chunks = []
        
        for i, key_point in enumerate(key_points):
            # Format a single key point for enrichment with proper prefix
            single_key_point_text = f"KEY POINT: {key_point}"
            logger.info(f"ENRICHING key point:{single_key_point_text}")
            
            # Enrich this specific key point with supporting data
            logger.debug(f"Enriching key point {i+1}/{len(key_points)}: {key_point[:50]}...")
            enriched_text = await enrich_key_points(single_key_point_text, full_text)
            logger.info(f"ENRICHED DATA: {enriched_text}")
            
            # Parse the enrichment result to extract supporting data without prefix
            supporting_data = ""
            for line in enriched_text.split('\n'):
                line = line.strip()
                if line.lower().startswith("supporting data:"):
                    colon_pos = line.find(':')
                    if colon_pos > 0:
                        supporting_data = line[colon_pos + 1:].strip()
                        break
            
            # If no supporting data found, use the key point itself
            if not supporting_data:
                logger.warning(f"No supporting data found for key point {i+1}, using key point as fallback")
                supporting_data = key_point
            
            # Add to chunks - storing raw data without prefixes
            key_point_chunks.append((key_point, supporting_data))
            logger.debug(f"Created chunk {i+1} - Key Point: '{key_point[:50]}...' with Supporting Data: '{supporting_data[:50]}...'")
        
        logger.info(f"Created {len(key_point_chunks)} enriched knowledge chunks")

        # Step 5: Save chunks to database
        processing_tasks = []
        
        for i, (key_point, supporting_data) in enumerate(key_point_chunks):
            chunk_id = f"chunk_{doc_id}_{i}"
            
            # Save the raw supporting data without any prefixes for better vector search
            processing_tasks.append(
                save_training_with_chunk(
                    input=supporting_data,  # No prefix - just the raw content for vectorization
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=True
                )
            )
            
            # For structured data, combine without prefixes for better vector search
            # but keep the relationship between key point and supporting data
            combined_text = f"{key_point}\n{supporting_data}"  # Removed prefixes, kept structured format
            
            processing_tasks.append(
                save_training_with_chunk(
                    input=combined_text,  # Combined text without prefixes for better vector search
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=False
                )
            )

        # Process save tasks
        logger.debug(f"Starting data storage for {len(processing_tasks)} tasks")
        results = await asyncio.gather(*processing_tasks, return_exceptions=True)
        
        success_count = 0
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Processing task failed: {result}")
            elif not result:
                logger.warning("Processing task returned False")
            else:
                success_count += 1
        
        if success_count == 0:
            logger.error("All processing tasks failed")
            return False
            
        logger.info(f"Processed document {doc_id} with {len(key_point_chunks)} chunks ({success_count}/{len(processing_tasks)} tasks succeeded)")
        return True

    except Exception as e:
        logger.error(f"Failed to process document: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False
    finally:
        if file:
            file.close()