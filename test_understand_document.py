import asyncio
import os
import json
import uuid
import networkx as nx
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from io import BytesIO
import tempfile
import sys
import traceback
import numpy as np
import matplotlib.patheffects as path_effects
from matplotlib.font_manager import FontProperties

# Set tokenizers parallelism to false to avoid HuggingFace warnings
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Import the actual function we want to test
from graph_building import understand_document

# Monkey patch NetworkX's write_graphml to handle complex Python types
original_write_graphml = nx.write_graphml

def patched_write_graphml(G, path, encoding='utf-8', prettyprint=True, infer_numeric_types=False,
                         named_key_ids=False, **kwargs):
    """
    Patched version of write_graphml that converts non-serializable types to strings
    before writing the GraphML file.
    """
    # Create a new graph with the same structure
    clean_graph = nx.Graph()
    
    # Copy nodes with cleaned attributes
    for node, attrs in G.nodes(data=True):
        clean_attrs = {}
        for key, value in attrs.items():
            if isinstance(value, np.ndarray):
                # Handle numpy arrays by converting to list and then to string
                clean_attrs[key] = str(value.tolist())
            elif isinstance(value, (list, dict, set, tuple, nx.Graph)):
                clean_attrs[key] = str(value)
            elif hasattr(value, '__dict__') or hasattr(value, '__slots__'):
                # Handle custom objects
                clean_attrs[key] = str(value)
            elif value is None:
                clean_attrs[key] = "None"
            else:
                clean_attrs[key] = value
        clean_graph.add_node(node, **clean_attrs)
    
    # Copy edges with cleaned attributes
    for u, v, attrs in G.edges(data=True):
        clean_attrs = {}
        for key, value in attrs.items():
            if isinstance(value, np.ndarray):
                # Handle numpy arrays by converting to list and then to string
                clean_attrs[key] = str(value.tolist())
            elif isinstance(value, (list, dict, set, tuple, nx.Graph)):
                clean_attrs[key] = str(value)
            elif hasattr(value, '__dict__') or hasattr(value, '__slots__'):
                # Handle custom objects
                clean_attrs[key] = str(value)
            elif value is None:
                clean_attrs[key] = "None"
            else:
                clean_attrs[key] = value
        clean_graph.add_edge(u, v, **clean_attrs)
    
    # Call the original function with the cleaned graph
    return original_write_graphml(clean_graph, path, encoding, prettyprint, 
                                 infer_numeric_types, named_key_ids, **kwargs)

# Apply the monkey patch
nx.write_graphml = patched_write_graphml

def create_enhanced_visualization(G, insights, output_path):
    """
    Create a detailed visualization of the knowledge graph with rich information boxes.
    
    Args:
        G: NetworkX graph object
        insights: Document insights dictionary
        output_path: Path to save the visualization
    """
    # Create a large figure for the detailed visualization
    plt.figure(figsize=(16, 12))
    
    # Set a larger font for all texts
    plt.rcParams.update({'font.size': 12})
    
    # Create a spring layout with more space between nodes
    pos = nx.spring_layout(G, k=0.4, seed=42)
    
    # Get cluster information
    clusters = {cluster['id']: cluster for cluster in insights['document_insights']['clusters']}
    
    # Draw edges first so they appear behind nodes
    nx.draw_networkx_edges(G, pos, alpha=0.3, width=1.5, edge_color='gray')
    
    # Draw nodes with different colors based on node importance (sentence count)
    node_sizes = []
    node_colors = []
    for node in G.nodes:
        # Extract cluster ID by removing "_new" suffix if present
        cluster_id = node.split('_')[0]
        
        if cluster_id in clusters:
            # Size based on sentence count (min 1000, max 3000)
            sentence_count = len(clusters[cluster_id]['sentences'])
            size = 1000 + min(sentence_count * 400, 2000)
            
            # Color based on cluster ID (for visual distinction)
            # Using a categorical colormap
            color_idx = int(cluster_id) % 10  # Modulo to ensure we don't exceed color range
            colors = plt.cm.tab10(range(10))
            color = colors[color_idx]
        else:
            size = 1000
            color = 'lightgray'
            
        node_sizes.append(size)
        node_colors.append(color)
    
    # Draw the nodes
    nodes = nx.draw_networkx_nodes(G, pos, 
                                  node_size=node_sizes, 
                                  node_color=node_colors, 
                                  alpha=0.8,
                                  edgecolors='black',
                                  linewidths=2)
    
    # Add a shadow effect to make nodes stand out
    nodes.set_path_effects([
        path_effects.withStroke(linewidth=5, foreground='#00000022')
    ])
    
    # Create detailed text boxes for each node
    for node in G.nodes:
        # Get node position
        x, y = pos[node]
        
        # Extract cluster ID by removing "_new" suffix if present
        cluster_id = node.split('_')[0]
        
        if cluster_id in clusters:
            cluster = clusters[cluster_id]
            
            # Title (just for the node label)
            title = cluster['title']
            # Show a sample sentence if available (first one)
            sample_text = ""
            if len(cluster['sentences']) > 0:
                sample_text = f"Sample: \"{cluster['sentences'][0][:50]}...\""
            
            # Create a text box with node details
            title_text = plt.text(
                x, y + 0.08, 
                title,
                ha='center', 
                va='center',
                fontsize=10,
                fontweight='bold',
                bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="black", alpha=0.8),
                wrap=True
            )
            
            # Add path effects to make text stand out
            title_text.set_path_effects([
                path_effects.withStroke(linewidth=3, foreground='white')
            ])
        else:
            # Fallback for nodes without cluster info
            plt.text(
                x, y, 
                f"Node {node}",
                ha='center', 
                va='center',
                fontsize=10,
                bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="black", alpha=0.7)
            )
    
    # Add detailed info boxes on the right side for each cluster
    plt.subplots_adjust(right=0.7)  # Make room for text on the right
    
    # Create an axes for the info panel
    info_ax = plt.axes([0.72, 0.1, 0.25, 0.8])
    info_ax.axis('off')
    
    # Set up the text
    info_text = "DOCUMENT CLUSTERS:\n"
    info_text += "=================\n\n"
    
    for i, cluster in enumerate(insights['document_insights']['clusters']):
        info_text += f"CLUSTER {cluster['id']}: {cluster['title']}\n"
        info_text += "-" * 40 + "\n"
        info_text += f"Description: {cluster['description'][:150]}...\n\n"
        
        # Add a sample sentence
        if len(cluster['sentences']) > 0:
            info_text += "Sample sentences:\n"
            for j, sentence in enumerate(cluster['sentences'][:2]):  # Show first 2 sentences
                info_text += f" - \"{sentence[:50]}...\"\n"
        
        info_text += f"Sentence count: {len(cluster['sentences'])}\n\n"
    
    # Add document summary to the info panel
    info_text += "DOCUMENT SUMMARY:\n"
    info_text += "================\n"
    info_text += insights['document_insights']['summary'] + "\n\n"
    
    # Add metadata
    info_text += "METADATA:\n"
    info_text += "========\n"
    for key, value in insights['document_insights']['metadata'].items():
        info_text += f"{key}: {value}\n"
    
    # Add the text to the info panel
    info_ax.text(0, 1, info_text, va='top', fontsize=9, family='monospace', wrap=True)
    
    # Add title and legend to the main graph
    plt.suptitle("Knowledge Graph with Cluster Details", fontsize=20, fontweight='bold')
    
    # Create a custom legend
    legend_text = "Node size represents the number of sentences in the cluster\n"
    legend_text += "Node color distinguishes different clusters"
    plt.figtext(0.02, 0.02, legend_text, fontsize=9, wrap=True)
    
    # Save the figure with high quality
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    print(f"Enhanced visualization saved to: {output_path}")
    
    # Create a second visualization showing cluster relationships
    create_relationship_visualization(insights, os.path.join(os.path.dirname(output_path), "cluster_relationships.png"))

def create_relationship_visualization(insights, output_path):
    """Create a visualization focused on the relationships between clusters"""
    plt.figure(figsize=(14, 10))
    
    clusters = insights['document_insights']['clusters']
    num_clusters = len(clusters)
    
    # Create a manually positioned circular layout
    G = nx.Graph()
    
    # Calculate positions in a circle
    pos = {}
    for i, cluster in enumerate(clusters):
        angle = 2 * np.pi * i / num_clusters
        x = 0.5 * np.cos(angle)
        y = 0.5 * np.sin(angle)
        node_id = cluster['id']
        G.add_node(node_id, **cluster)
        pos[node_id] = (x, y)
    
    # Add edges between all clusters (fully connected graph)
    for i in range(num_clusters):
        for j in range(i+1, num_clusters):
            id1 = clusters[i]['id']
            id2 = clusters[j]['id']
            # Create a sample weight - in a real scenario, this could be from actual similarity
            G.add_edge(id1, id2, weight=0.5)
    
    # Draw edges
    nx.draw_networkx_edges(G, pos, width=1.5, alpha=0.3)
    
    # Draw nodes
    node_sizes = [1000 + len(cluster['sentences']) * 200 for cluster in clusters]
    nodes = nx.draw_networkx_nodes(G, pos, node_size=node_sizes, alpha=0.7, 
                                  node_color=plt.cm.viridis(np.linspace(0, 1, num_clusters)))
    
    # Add cluster titles as node labels
    labels = {cluster['id']: cluster['title'] for cluster in clusters}
    nx.draw_networkx_labels(G, pos, labels=labels, font_size=10, font_weight='bold')
    
    # Add detailed info boxes around each node
    for i, cluster in enumerate(clusters):
        node_id = cluster['id']
        x, y = pos[node_id]
        
        # Calculate box positions - a bit offset from the node
        box_x = x + 0.2 if x > 0 else x - 0.2
        box_y = y + 0.2 if y > 0 else y - 0.2
        
        # Get a sample sentence
        sample = cluster['sentences'][0][:50] + "..." if cluster['sentences'] else "No sample"
        
        # Add a text box with sample info
        plt.text(box_x, box_y, 
                f"Sample: {sample}", 
                bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="black", alpha=0.7),
                fontsize=8, ha='center', va='center', wrap=True)
    
    # Add title
    plt.title("Cluster Relationships and Sample Content", fontsize=16)
    
    # Add legend for node sizes
    plt.figtext(0.01, 0.01, "Node size represents the number of sentences in each cluster", fontsize=10)
    
    plt.axis('off')
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    print(f"Cluster relationship visualization saved to: {output_path}")

async def test_understand_document():
    """
    Test the understand_document function directly with proper handling of complex data types.
    
    This test applies a monkey patch to NetworkX's GraphML serialization to handle 
    complex Python types like lists and embeddings that are not normally supported in GraphML.
    """
    print("Starting test of understand_document with GraphML serialization fix...")
    
    # Create a unique test directory
    test_dir = f"test_understand_{uuid.uuid4().hex[:8]}"
    os.makedirs(test_dir, exist_ok=True)
    print(f"Created test directory: {test_dir}")
    
    # Define test file paths
    graph_path = os.path.join(test_dir, "knowledge_graph.graphml")
    
    # Set font for Vietnamese text support
    plt.rcParams['font.family'] = 'DejaVu Sans'
    
    # Step 1: Check if the Self-driving map.docx exists
    print("\n--- STEP 1: Checking for test document ---")
    if not os.path.exists("Self-driving map.docx"):
        print("Error: Self-driving map.docx not found. Creating a simple test document instead.")
        # Create a simple test document
        doc = Document()
        doc.add_heading('Test Document for Graph Building', 0)
        doc.add_paragraph('This is a simple test document to validate the graph building functionality.')
        doc.add_paragraph('It contains multiple sentences that should be processed into clusters.')
        doc.add_paragraph('The document processing should extract insights from these sentences.')
        doc.add_paragraph('Ideally, these sentences would be grouped by topic or similarity.')
        doc.add_paragraph('The resulting graph should show connections between related concepts.')
        test_doc_path = os.path.join(test_dir, "test_document.docx")
        doc.save(test_doc_path)
    else:
        test_doc_path = "Self-driving map.docx"
        print(f"Using existing document: {test_doc_path}")
    
    # Step 2: Process the document
    print("\n--- STEP 2: Processing document with understand_document ---")
    try:
        # Process the document with a fresh graph path
        insights = await understand_document(
            input_source=test_doc_path,
            file_type="docx",
            existing_graph_path=graph_path
        )
        
        if not insights["success"]:
            print(f"Error processing document: {insights.get('error', 'Unknown error')}")
            print(f"Error type: {insights.get('error_type', 'Unknown')}")
            return
        
        # Save the insights to file
        insights_path = os.path.join(test_dir, "document_insights.json")
        with open(insights_path, "w", encoding="utf-8") as f:
            json.dump(insights, f, indent=2, ensure_ascii=False)
        
        print("Document processing complete!")
        print(f"Number of clusters: {insights['document_insights']['metadata']['cluster_count']}")
        print(f"Number of sentences: {insights['document_insights']['metadata']['sentence_count']}")
        
        # Display clusters
        print("\nClusters:")
        for cluster in insights["document_insights"]["clusters"]:
            print(f"- Cluster {cluster['id']}: {cluster['title']}")
            print(f"  Description: {cluster['description']}")
            print(f"  Sentence count: {len(cluster['sentences'])}")
        
        # Step 3: Validate the GraphML file was created
        print("\n--- STEP 3: Validating GraphML output ---")
        if os.path.exists(graph_path):
            print(f"GraphML file created successfully: {graph_path}")
            # Try to read it back to verify integrity
            try:
                G = nx.read_graphml(graph_path)
                print(f"Successfully read back the graph: {len(G.nodes)} nodes, {len(G.edges)} edges")
                
                # List node titles
                print("\nGraph nodes:")
                for node in G.nodes:
                    title = G.nodes[node].get("title", "No title")
                    print(f"- Node {node}: {title}")
                
                # Create enhanced visualizations
                viz_path = os.path.join(test_dir, "enhanced_visualization.png")
                create_enhanced_visualization(G, insights, viz_path)
                
                # Also create a simple visualization for comparison
                simple_viz_path = os.path.join(test_dir, "simple_graph_visualization.png")
                plt.figure(figsize=(10, 8))
                pos = nx.spring_layout(G, seed=42)
                nx.draw_networkx_nodes(G, pos, node_color="lightblue", node_size=1000)
                nx.draw_networkx_edges(G, pos, alpha=0.5)
                node_labels = {n: G.nodes[n].get("title", n)[:20] for n in G.nodes}
                nx.draw_networkx_labels(G, pos, labels=node_labels, font_size=9)
                plt.title("Basic Knowledge Graph Visualization")
                plt.axis('off')
                plt.tight_layout()
                plt.savefig(simple_viz_path, dpi=300)
                plt.close()
                print(f"Simple graph visualization saved to: {simple_viz_path}")
                
            except Exception as e:
                print(f"Error reading back the GraphML file: {type(e).__name__}: {str(e)}")
                traceback.print_exc()
        else:
            print(f"Error: GraphML file was not created at {graph_path}")
        
        # Return success
        print(f"\nAll test outputs saved to directory: {test_dir}")
        return True
        
    except Exception as e:
        print(f"Error during document processing: {type(e).__name__}: {str(e)}")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print(f"Python version: {sys.version}")
    print("Running understand_document test with GraphML serialization fix...")
    asyncio.run(test_understand_document())
    print("Test completed!") 