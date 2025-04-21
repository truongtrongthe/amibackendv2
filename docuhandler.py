# docuhandler.py
import pdfplumber
from docx import Document
import uuid
from io import BytesIO
from utilities import logger
from database import save_training_with_chunk
import asyncio
from werkzeug.datastructures import FileStorage
from langchain_openai import ChatOpenAI

# Initialize LLM
LLM = ChatOpenAI(model="gpt-4o-mini", streaming=False)

async def summarize_with_llm(text: str, max_length: int = 150) -> str:
    """Summarize text using the LLM in the original language, extracting key knowledge elements."""
    try:
        prompt = (
            f"You are a precise document summarizer specialized in knowledge extraction.\n\n"
            f"TASK: Create a structured summary of the following text in {max_length} words or less.\n\n"
            f"GUIDELINES:\n"
            f"1. Identify and extract the most important FACTUAL information\n"
            f"2. Prioritize extracting actionable knowledge (procedures, requirements, steps)\n"
            f"3. Include specific details like names, numbers, dates, and key terminology\n"
            f"4. Maintain the original language of the text\n"
            f"5. Use concise, direct language without filler words\n"
            f"6. For technical documents, preserve technical accuracy and domain-specific terms\n"
            f"7. For instructional content, capture all key steps in proper sequence\n"
            f"8. For informational content, prioritize concrete facts over general statements\n\n"
            f"TEXT TO SUMMARIZE:\n" + text
        )
        response = await LLM.ainvoke(prompt)
        summary = response.content.strip()
        logger.info(f"Generated summary (approx {len(summary.split())} words): {summary}")
        return summary
    except Exception as e:
        logger.error(f"LLM summarization failed: {e}")
        return "Summary unavailable due to processing error."

async def extract_knowledge_with_llm(text: str, domain: str = "", max_items: int = 10) -> str:
    """
    Extract structured knowledge elements from text using the LLM.
    This is distinct from summarization - it identifies discrete knowledge units.
    
    Args:
        text: The text to extract knowledge from
        domain: Optional domain context (e.g., "finance", "healthcare")
        max_items: Maximum number of knowledge elements to extract
        
    Returns:
        str: Structured knowledge elements in a formatted string
    """
    try:
        domain_context = f" in the {domain} domain" if domain else ""
        
        prompt = (
            f"You are a precise knowledge extractor specialized in identifying actionable information{domain_context}.\n\n"
            f"TASK: Extract the {max_items} most important knowledge elements from the text below.\n\n"
            f"GUIDELINES FOR KNOWLEDGE EXTRACTION:\n"
            f"1. Focus on extracting FACTUAL and ACTIONABLE knowledge (procedures, requirements, rules, criteria, etc.)\n"
            f"2. Prioritize information that would be most useful for solving problems or making decisions\n"
            f"3. Include specific details like quantities, timelines, requirements, and constraints\n"
            f"4. Maintain the original terminology and technical accuracy\n"
            f"5. Format each knowledge element as a clear, concise statement\n"
            f"6. For procedures or processes, preserve the correct sequence\n"
            f"7. For conditions or rules, preserve the exact criteria\n"
            f"8. Ignore general background information unless it contains critical context\n\n"
            
            f"FORMAT YOUR RESPONSE AS:\n"
            f"KNOWLEDGE ELEMENT 1: [concise statement of knowledge]\n"
            f"KNOWLEDGE ELEMENT 2: [concise statement of knowledge]\n"
            f"And so on...\n\n"
            
            f"TEXT TO ANALYZE:\n{text}"
        )
        
        response = await LLM.ainvoke(prompt)
        extracted_knowledge = response.content.strip()
        logger.info(f"Extracted {extracted_knowledge.count('KNOWLEDGE ELEMENT')} knowledge elements")
        return extracted_knowledge
    except Exception as e:
        logger.error(f"LLM knowledge extraction failed: {e}")
        return "Knowledge extraction unavailable due to processing error."

async def summarize_document(file: FileStorage, user_id: str, mode: str = "default") -> tuple[bool, dict]:
    """
    Process a document: extract full text, maintain structural context, summarize, and extract knowledge.
    Returns (success: bool, result_dict: dict). Does not save to database.
    """
    full_text = ""
    document_metadata = {}

    file_extension = file.filename.split('.')[-1].lower()
    try:
        logger.debug(f"Reading file content for {file.filename}")
        file_content = file.read()
        if not file_content:
            logger.warning(f"Empty file content for {file.filename}")
            return False, {"error": "Empty file content."}

        # Step 1: Extract full text with metadata
        if file_extension == 'pdf':
            logger.debug("Opening PDF file")
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                document_metadata = {
                    'page_count': len(pdf.pages),
                    'filename': file.filename,
                    'file_type': 'PDF'
                }
                
                # Extract text with page context
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages_text.append(f"[Page {i+1}] {page_text}")
                
                full_text = "\n\n".join(pages_text)
                if not full_text.strip():
                    logger.warning(f"No text extracted from {file.filename}")
                    return False, {"error": "No text extracted."}
                    
        elif file_extension == 'docx':
            logger.debug("Opening DOCX file")
            doc = Document(BytesIO(file_content))
            
            # Extract document properties
            document_metadata = {
                'filename': file.filename,
                'file_type': 'DOCX',
                'paragraph_count': len(doc.paragraphs)
            }
            
            # Extract text with structural context
            full_text_parts = []
            current_heading = "Document Start"
            
            for para in doc.paragraphs:
                if para.text.strip():
                    if para.style.name.startswith('Heading'):
                        current_heading = para.text.strip()
                        full_text_parts.append(f"\n[{current_heading}]\n{para.text}")
                    else:
                        full_text_parts.append(para.text)
            
            full_text = "\n".join(full_text_parts)
            if not full_text.strip():
                logger.warning(f"No text extracted from {file.filename}")
                return False, {"error": "No text extracted."}
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return False, {"error": f"Unsupported file type: {file_extension}"}

        logger.debug(f"Extracted full_text (first 100 chars): {full_text[:100]}...")

        # Step 2: Prepare contextual information for summarization
        context_prefix = f"Document: {file.filename}\nType: {document_metadata.get('file_type', 'Unknown')}\n"
        
        if document_metadata.get('file_type') == 'PDF':
            context_prefix += f"Pages: {document_metadata.get('page_count', 'Unknown')}\n\n"
        
        # Step 3: Run summarization and knowledge extraction in parallel
        summary_task = summarize_with_llm(context_prefix + full_text)
        knowledge_task = extract_knowledge_with_llm(full_text)
        
        summary, knowledge = await asyncio.gather(summary_task, knowledge_task)
        
        logger.info(f"Full text length: {len(full_text.split())} words")
        logger.info(f"Processing complete for {file.filename}")

        return True, {
            "summary": summary,
            "knowledge_elements": knowledge,
            "metadata": document_metadata,
            "text_length": len(full_text.split())
        }

    except Exception as e:
        logger.error(f"Failed to process document: {type(e).__name__}: {str(e)} - File: {file.filename}")
        return False, {"error": f"Processing error occurred: {str(e)}"}
    finally:
        file.close()

async def extract_knowledge_from_chunk(chunk_text: str, chunk_context: str = "") -> str:
    """
    Extract focused knowledge elements from a document chunk with context awareness.
    Optimized for smaller text segments compared to full document knowledge extraction.
    
    Args:
        chunk_text: The chunk text to analyze
        chunk_context: Optional context about the chunk's position/role in the document
        
    Returns:
        str: Extracted knowledge from the chunk in a structured format
    """
    try:
        # Create a concise, focused prompt for chunk knowledge extraction
        prompt = (
            f"You are a precision knowledge extractor tasked with identifying the most valuable information from a document segment.\n\n"
            f"CONTEXT: This is a segment {chunk_context} from a larger document.\n\n"
            f"TASK: Extract the 3-5 most important, specific knowledge elements from this segment.\n\n"
            f"EXTRACTION GUIDELINES:\n"
            f"1. Prioritize extracting CONCRETE, ACTIONABLE information\n"
            f"2. Focus on specifics: facts, figures, requirements, steps, rules, conditions\n"
            f"3. Maintain precise technical details including numbers, identifiers, and specialized terminology\n"
            f"4. Include contextual relationships for better understanding\n"
            f"5. Remove duplicate or redundant information\n"
            f"6. Preserve the original meaning and critical nuances\n\n"
            
            f"FORMAT YOUR RESPONSE AS:\n"
            f"KEY POINT 1: [precise factual statement with specific details]\n"
            f"KEY POINT 2: [precise factual statement with specific details]\n"
            f"And so on...\n\n"
            
            f"SEGMENT TEXT:\n{chunk_text}"
        )
        
        response = await LLM.ainvoke(prompt)
        extracted_knowledge = response.content.strip()
        
        # Count key points for logging
        key_point_count = 0
        for line in extracted_knowledge.split('\n'):
            if line.strip().startswith("KEY POINT"):
                key_point_count += 1
        
        logger.debug(f"Extracted {key_point_count} key points from chunk")
        return extracted_knowledge
    except Exception as e:
        logger.error(f"Chunk knowledge extraction failed: {e}")
        return "Knowledge extraction failed for this segment."

async def process_document(file: FileStorage, user_id: str, mode: str = "default", bank: str = "") -> bool:
        doc_id = str(uuid.uuid4())
        chunks = []
        extracted_knowledge = []

        logger.info(f"bank name at process-document={bank}")
        # Determine file type
        file_extension = file.filename.split('.')[-1].lower()
        file_content = file.read()

        try:
            if file_extension == 'pdf':
                with pdfplumber.open(BytesIO(file_content)) as pdf:
                    current_section = []
                    current_word_count = 0
                    section_index = 0
                    page_break_indices = []
                    
                    # First, extract text with page breaks marked
                    all_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            all_text.append(page_text)
                            if i > 0:  # Don't add page break before first page
                                page_break_indices.append(len(all_text) - 1)
                    
                    full_text = " ".join(all_text)
                    
                    # Split text into paragraphs
                    paragraphs = []
                    for para in full_text.split('\n'):
                        if para.strip():
                            paragraphs.append(para.strip())
                    
                    # Process paragraphs into semantically coherent chunks
                    for i, para in enumerate(paragraphs):
                        para_words = len(para.split())
                        
                        # If this paragraph alone exceeds max size, split it
                        if para_words > 300:
                            words = para.split()
                            for j in range(0, len(words), 250):  # Smaller chunk size for large paragraphs
                                sub_chunk = " ".join(words[j:j + 250])
                                chunk_id = f"chunk_{section_index}_{j // 250}"
                                chunk_context = f"(large paragraph section {j // 250 + 1})"
                                chunks.append((chunk_id, sub_chunk, chunk_context))
                            section_index += 1
                            continue
                            
                        # Handle normal case - build up sections
                        if current_word_count + para_words <= 300:
                            # Continue current section
                            current_section.append(para)
                            current_word_count += para_words
                        else:
                            # Current section is full, save it and start a new one
                            if current_section:
                                section_text = " ".join(current_section)
                                chunk_id = f"chunk_{section_index}_0"
                                chunk_context = f"(section {section_index + 1})"
                                chunks.append((chunk_id, section_text, chunk_context))
                                section_index += 1
                            
                            # Start new section with current paragraph
                            current_section = [para]
                            current_word_count = para_words
                    
                    # Add any remaining content as final chunk
                    if current_section:
                        section_text = " ".join(current_section)
                        chunk_id = f"chunk_{section_index}_0"
                        chunk_context = f"(final section {section_index + 1})"
                        chunks.append((chunk_id, section_text, chunk_context))

            elif file_extension == 'docx':
                doc = Document(BytesIO(file_content))
                
                current_section = []
                current_word_count = 0
                section_index = 0
                current_heading = "Introduction"
                
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                        
                    # Check if this is a heading - if so, update current heading
                    is_heading = para.style.name.startswith('Heading')
                    if is_heading:
                        current_heading = para.text.strip()
                        
                    para_text = para.text.strip()
                    para_words = len(para_text.split())
                    
                    # If this paragraph alone exceeds max size, split it
                    if para_words > 300:
                        # Process any current section first
                        if current_section:
                            section_text = " ".join(current_section)
                            chunk_id = f"chunk_{section_index}_0"
                            chunk_context = f"(section: {current_heading}, part {section_index + 1})"
                            chunks.append((chunk_id, section_text, chunk_context))
                            section_index += 1
                            current_section = []
                            current_word_count = 0
                            
                        # Then handle the large paragraph
                        words = para_text.split()
                        for j in range(0, len(words), 250):
                            sub_chunk = " ".join(words[j:j + 250])
                            chunk_id = f"chunk_{section_index}_{j // 250}"
                            chunk_context = f"(section: {current_heading}, large paragraph part {j // 250 + 1})"
                            chunks.append((chunk_id, sub_chunk, chunk_context))
                        section_index += 1
                        continue
                        
                    # Start a new section if:
                    # 1. This is a heading
                    # 2. Current section would be too large with this paragraph
                    if is_heading or (current_word_count + para_words > 300):
                        # Save existing section if not empty
                        if current_section:
                            section_text = " ".join(current_section)
                            chunk_id = f"chunk_{section_index}_0"
                            chunk_context = f"(section: {current_heading})"
                            chunks.append((chunk_id, section_text, chunk_context))
                            section_index += 1
                            current_section = []
                            current_word_count = 0
                    
                    # Add paragraph to current section
                    current_section.append(para_text)
                    current_word_count += para_words
                
                # Add any remaining content as final chunk
                if current_section:
                    section_text = " ".join(current_section)
                    chunk_id = f"chunk_{section_index}_0"
                    chunk_context = f"(final section: {current_heading})"
                    chunks.append((chunk_id, section_text, chunk_context))

            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return False

            if not chunks:
                logger.warning(f"No text extracted from {file.filename}")
                return False

            # Process chunks and extract knowledge in parallel
            processing_tasks = []
            knowledge_extraction_tasks = []
            
            for chunk_id, chunk_text, chunk_context in chunks:
                # Add task for knowledge extraction from this chunk
                knowledge_extraction_tasks.append(
                    extract_knowledge_from_chunk(chunk_text, chunk_context)
                )
                
                # Add tasks for saving raw and processed chunks
                processing_tasks.append(
                    save_training_with_chunk(
                        input=chunk_text,
                        user_id=user_id,
                        mode=mode,
                        doc_id=doc_id,
                        chunk_id=chunk_id,
                        bank_name=bank,
                        is_raw=True
                    )
                )
                
                # We'll replace this with the extracted knowledge in a moment
                processing_tasks.append(
                    save_training_with_chunk(
                        input=chunk_text,  # Placeholder - will be replaced with knowledge extraction
                        user_id=user_id,
                        mode=mode,
                        doc_id=doc_id,
                        chunk_id=chunk_id,
                        bank_name=bank,
                        is_raw=False
                    )
                )
            
            # Run knowledge extraction for all chunks
            knowledge_results = await asyncio.gather(*knowledge_extraction_tasks, return_exceptions=True)
            
            # Replace the placeholder tasks with actual knowledge extraction results
            for i, knowledge in enumerate(knowledge_results):
                if isinstance(knowledge, Exception):
                    logger.error(f"Knowledge extraction failed for chunk {i}: {str(knowledge)}")
                    knowledge = f"Knowledge extraction failed: {str(knowledge)}"
                    
                # Replace the corresponding is_raw=False task with extracted knowledge
                task_index = (i * 2) + 1  # Each chunk has 2 tasks, we want the second one
                if task_index < len(processing_tasks):
                    # Get the original task
                    orig_task = processing_tasks[task_index]
                    # Replace it with a new task using the extracted knowledge
                    processing_tasks[task_index] = save_training_with_chunk(
                        input=knowledge,
                        user_id=user_id,
                        mode=mode,
                        doc_id=doc_id,
                        chunk_id=chunks[i][0],  # chunk_id from original chunks list
                        bank_name=bank,
                        is_raw=False
                    )
            
            # Process all save tasks
            results = await asyncio.gather(*processing_tasks, return_exceptions=True)
            for result in results:
                if isinstance(result, Exception) or not result:
                    logger.error(f"Processing failed for one or more chunks: {result}")
                    return False

            logger.info(f"Processed document {doc_id} with {len(chunks)} chunks from {file.filename}")
            return True

        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            return False
        finally:
            file.close()