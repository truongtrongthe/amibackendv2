import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
import hdbscan
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import logging
from docx import Document
import pdfplumber
from io import BytesIO
import os
from typing import List, Tuple, Dict, Union, BinaryIO
import json
from database import save_training_with_chunk
import uuid
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the LLM client with timeout
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)
model = SentenceTransformer('sentence-transformers/LaBSE')
# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt, temperature=0.2, max_tokens=5000, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop, temperature=temperature, max_tokens=max_tokens)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

def load_and_split_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> List[str]:
    """
    Load a document from file path or BytesIO and split it into sentences.
    
    Args:
        input_source: Either a file path (str) or a BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        List of sentences
    """
    logger.info("Loading and splitting document.Filetype: %s", file_type)
    try:
        # Determine whether we have a file path or a file-like object
        if isinstance(input_source, str):
            # It's a file path
            file_path = input_source
            if not file_type:
                # Detect file type from extension if not provided
                file_type = file_path.split('.')[-1].lower()
                
            # Different handling based on file type
            if file_type == 'docx':
                logger.info(f"Loading DOCX from file path: {file_path}")
                doc = Document(file_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info(f"Loading PDF from file path: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            # It's a file-like object (BytesIO)
            file_content = input_source
            
            # Reset position to beginning
            file_content.seek(0)
            
            # Must have file_type for file-like objects
            if not file_type:
                raise ValueError("file_type must be provided when using file-like input")
                
            if file_type == 'docx':
                logger.info("Loading DOCX from BytesIO/file-like object")
                doc = Document(file_content)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info("Loading PDF from BytesIO/file-like object")
                with pdfplumber.open(file_content) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be either a file path or a file-like object")
            
        # Log text extraction info
        logger.info(f"Extracted {len(full_text)} characters of text")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")
        
        # Check if we have enough content
        if not full_text.strip():
            logger.warning("No text content extracted")
            raise ValueError("No text content could be extracted from the document")
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            raise ValueError("Text content too short for meaningful processing")
        
        # Process extracted text into sentences
        sentences = sent_tokenize(full_text)
        clean_sentences = [s.strip() for s in sentences if s.strip()]
        logger.info(f"Split text into {len(clean_sentences)} sentences")
        
        return clean_sentences
        
    except Exception as e:
        logger.error(f"Error loading document: {type(e).__name__}: {str(e)}")
        raise

def generate_sentence_embeddings(sentences: List[str]) -> np.ndarray:
    """
    Generate embeddings for sentences using LaBSE.
    
    Args:
        sentences: List of sentences to embed
        
    Returns:
        numpy array of embeddings
    """
    try:
        
        embeddings = model.encode(sentences, show_progress_bar=True)
        return embeddings
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise

def perform_clustering(embeddings: np.ndarray) -> np.ndarray:
    """
    Cluster sentences using HDBSCAN.
    
    Args:
        embeddings: Sentence embeddings
        
    Returns:
        Array of cluster labels
    """
    try:
        # Add special case handling for very few sentences
        if len(embeddings) == 1:
            logger.warning("Only 1 sentence available - skipping clustering")
            # Put the single point into cluster 0
            return np.zeros(1, dtype=int)
            
        # Special handling for exactly 2 sentences - still try to cluster but with modified parameters
        if len(embeddings) == 2:
            logger.warning("Only 2 sentences available - using simplified clustering")
            # Calculate similarity between the 2 sentences
            similarity = cosine_similarity(embeddings)[0, 1]
            logger.info(f"Similarity between the 2 sentences: {similarity:.4f}")
            
            # If they're similar enough, put them in the same cluster, otherwise different clusters
            if similarity > 0.5:  # Threshold can be adjusted
                logger.info("Sentences are similar - grouping in the same cluster")
                return np.zeros(2, dtype=int)
            else:
                logger.info("Sentences are different - putting in separate clusters")
                return np.array([0, 1], dtype=int)
        
        # For 3+ sentences, use standard HDBSCAN clustering
        # Further adjusted parameters for smaller, more focused clusters
        clusterer = hdbscan.HDBSCAN(
            min_cluster_size=2,           # Keep small clusters
            min_samples=1,                # Allow single points to form clusters
            metric='euclidean',
            cluster_selection_method='eom',
            cluster_selection_epsilon=0.25, # Further reduced to create even tighter clusters
            alpha=0.8,                    # Reduced to encourage more clusters
            prediction_data=True,         # Enable prediction data for better stability
            allow_single_cluster=False    # Prevent all points from being in one cluster
        )
        
        # Fit and predict
        cluster_labels = clusterer.fit_predict(embeddings)
        
        # Log clustering statistics
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                        for label in unique_clusters if label != -1}
        
        logger.info(f"Clustering statistics:")
        logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
        logger.info(f"  Noise points: {noise_count}")
        logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        # If we have too few clusters, try again with different parameters
        if len(unique_clusters) - (1 if -1 in unique_clusters else 0) < 2 and len(embeddings) > 3:
            logger.info("Too few clusters detected, retrying with different parameters...")
            clusterer = hdbscan.HDBSCAN(
                min_cluster_size=2,
                min_samples=1,
                metric='euclidean',
                cluster_selection_method='leaf',
                cluster_selection_epsilon=0.15,
                alpha=0.6,
                prediction_data=True,
                allow_single_cluster=False
            )
            cluster_labels = clusterer.fit_predict(embeddings)
            
            # Log new clustering statistics
            unique_clusters = set(cluster_labels)
            noise_count = sum(1 for label in cluster_labels if label == -1)
            cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                            for label in unique_clusters if label != -1}
            
            logger.info(f"Retry clustering statistics:")
            logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
            logger.info(f"  Noise points: {noise_count}")
            logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        return cluster_labels
    except Exception as e:
        logger.error(f"Error clustering sentences: {type(e).__name__}: {str(e)}")
        raise

def select_all_sentences(
    sentences: List[str],
    cluster_labels: np.ndarray
) -> Dict[int, List[str]]:
    """
    Select all sentences for each cluster, allowing sentences to appear in multiple clusters.
    
    Args:
        sentences: List of sentences
        cluster_labels: Cluster labels for each sentence
        
    Returns:
        Dictionary mapping cluster IDs to lists of sentences
    """
    try:
        cluster_sentences = {}
        unique_clusters = set(cluster_labels)
        
        for cluster_id in unique_clusters:
            if cluster_id == -1:  # Skip noise points
                continue
                
            # Get all sentences for this cluster
            cluster_mask = cluster_labels == cluster_id
            cluster_sentences[cluster_id] = [
                sentences[i] for i in range(len(sentences)) if cluster_mask[i]
            ]
            
        return cluster_sentences
    except Exception as e:
        logger.error(f"Error selecting sentences: {str(e)}")
        raise

async def generate_cluster_title(sentences: List[str]) -> str:
    """
    Generate a descriptive title for a cluster using GPT-4o.
    
    Args:
        sentences: List of sentences
        
    Returns:
        A descriptive title for the cluster
    """
    try:
        sentences_text = "\n".join([f"- {s}" for s in sentences])
        prompt = (
            f"You are a knowledge organization expert.\n\n"
            f"TASK: Based on the following sentences, generate a SHORT, DESCRIPTIVE TITLE (3-5 words) that captures the core theme or topic.\n"
            f"If the sentences are not concise enough, rephrase them into a clear, focused title.\n\n"
            f"GUIDELINES:\n"
            f"1. The title should be in the SAME LANGUAGE as the sentences\n"
            f"2. If the sentences are verbose or unclear, create a concise title that captures the essence\n"
            f"3. Use clear, direct language that immediately conveys the main idea\n"
            f"4. Keep technical terms and domain-specific vocabulary\n"
            f"5. Make the title actionable and specific\n\n"
            f"SENTENCES:\n{sentences_text}\n\n"
            f"TITLE:"
        )
        
        response = await invoke_llm_with_retry(LLM, prompt, max_tokens=25)
        title = response.content.strip()
        
        # Remove quotes if present
        if title.startswith('"') and title.endswith('"'):
            title = title[1:-1]
        
        return title
    except Exception as e:
        logger.error(f"Error generating cluster title: {str(e)}")
        return "Unnamed Cluster"

async def generate_cluster_description(
    cluster_sentences: Dict[int, List[str]],
    cluster_id: int,
    all_clusters: Dict[int, List[str]]
) -> str:
    """
    Generate a description of the cluster and its relationships to other clusters.
    
    Args:
        cluster_sentences: Dictionary mapping cluster IDs to sentences
        cluster_id: ID of the current cluster
        all_clusters: Dictionary of all clusters
        
    Returns:
        A description of the cluster and its relationships
    """
    try:
        # Get sentences from current cluster
        current_sentences = cluster_sentences[cluster_id]
        
        # Get sentences from other clusters
        other_clusters = {cid: sents for cid, sents in all_clusters.items() 
                         if cid != cluster_id and cid != -1}
        
        # Prepare the prompt
        current_text = "\n".join([f"- {s}" for s in current_sentences])
        other_text = "\n".join([f"Cluster {cid}:\n" + "\n".join([f"- {s}" for s in sents])
                              for cid, sents in other_clusters.items()])
        
        prompt = (
            f"You are a knowledge organization expert.\n\n"
            f"TASK: Analyze the following cluster of sentences and describe them in the EXACT SAME LANGUAGE as the sentences.\n"
            f"DO NOT TRANSLATE TO ENGLISH. Use the same words, phrases, and expressions as in the original sentences.\n\n"
            f"EXAMPLE:\n"
            f"If the sentences are in Vietnamese, write the description in Vietnamese.\n"
            f"If the sentences are in Japanese, write the description in Japanese.\n"
            f"If the sentences are in Spanish, write the description in Spanish.\n\n"
            f"REQUIREMENTS:\n"
            f"1. Use the EXACT SAME LANGUAGE as the sentences - no translation allowed\n"
            f"2. Keep all technical terms and domain-specific vocabulary exactly as they appear\n"
            f"3. Maintain the same tone, style, and expressions\n"
            f"4. Preserve all cultural context and nuances\n"
            f"5. Do not add any English words or phrases\n"
            f"6. If the sentences use informal language, use the same informal style\n"
            f"7. If the sentences use formal language, use the same formal style\n\n"
            f"DESCRIBE:\n"
            f"1. The main theme and purpose of this cluster\n"
            f"2. How it relates to other clusters in the document\n"
            f"3. The role it plays in the overall context\n\n"
            f"CURRENT CLUSTER SENTENCES:\n{current_text}\n\n"
            f"OTHER CLUSTERS:\n{other_text}\n\n"
            f"DESCRIPTION (write in the same language as the sentences above):"
        )
        
        response = await invoke_llm_with_retry(LLM, prompt, max_tokens=200)
        return response.content.strip()
    except Exception as e:
        logger.error(f"Error generating cluster description: {str(e)}")
        return ""

async def generate_takeaways(
    cluster_sentences: Dict[int, List[str]],
    sentences: List[str]
) -> Dict[int, Tuple[str, str, str]]:
    """
    Generate titles, descriptions, and application-focused takeaways for each cluster using GPT-4o.
    
    Args:
        cluster_sentences: Dictionary mapping cluster IDs to sentences
        sentences: Complete list of all sentences from the document for context
        
    Returns:
        Dictionary mapping cluster IDs to (title, description, takeaways) tuples
    """
    try:
        results = {}
        
        # Get a sample of the full document for context
        doc_sample = "\n".join(sentences[:30] if len(sentences) > 30 else sentences)
        
        for cluster_id, cluster_sents in cluster_sentences.items():
            # Generate cluster title
            title = await generate_cluster_title(cluster_sents)
            logger.info(f"Generated title for cluster {cluster_id}: {title}")
            
            # Generate cluster description and relationships
            description = await generate_cluster_description(cluster_sentences, cluster_id, cluster_sentences)
            logger.info(f"Generated description for cluster {cluster_id}")
            
            # Prepare the prompt for takeaways
            sentences_text = "\n".join([f"- {s}" for s in cluster_sents])
            prompt = (
                f"You are a knowledge application expert specializing in sales.\n\n"
                f"DOCUMENT CONTEXT (sample from the full document):\n{doc_sample}\n\n"
                f"TASK: Analyze the following sentences and generate detailed, practical takeaways focusing SPECIFICALLY on HOW TO APPLY these insights in sales contexts.\n"
                f"IMPORTANT: Your response MUST be in the SAME LANGUAGE as the sentences.\n\n"
                f"GUIDELINES:\n"
                f"1. Focus on METHODS OF APPLICATION - explain exactly HOW to apply these insights\n"
                f"2. Structure your response with numbered steps or a clear methodology\n"
                f"3. Provide SPECIFIC EXAMPLE SCRIPTS, dialogues, or templates that demonstrate the application\n"
                f"4. Include step-by-step instructions for implementation\n"
                f"5. Consider the document context when interpreting these sentences\n"
                f"6. Format your response as 'Application Method: [title]' followed by the steps\n"
                f"7. Each application method should be immediately actionable\n"
                f"8. IMPORTANT: Maintain the original language of the sentences in your response\n"
                f"9. Include specific factual details and nuanced observations from the text\n"
                f"10. Highlight any subtle but important distinctions or variations in approach\n"
                f"11. Note any potential edge cases or special situations to consider\n"
                f"12. Use the same terminology and expressions as in the original text\n"
                f"13. Keep all technical terms and domain-specific vocabulary in their original form\n\n"
                f"SENTENCES TO ANALYZE:\n{sentences_text}\n\n"
                f"APPLICATION METHODS (with specific examples and step-by-step instructions):"
            )
            
            # Generate takeaway
            response = await invoke_llm_with_retry(LLM, prompt)
            results[cluster_id] = (title, description, response.content.strip())
            
        return results
    except Exception as e:
        logger.error(f"Error generating takeaways: {str(e)}")
        raise

def save_results(
    cluster_sentences: Dict[int, List[str]],
    cluster_results: Dict[int, Tuple[str, str, str]],
    output_path: str
) -> None:
    """
    Save results to a markdown file.
    
    Args:
        cluster_sentences: Dictionary mapping cluster IDs to sentences
        cluster_results: Dictionary mapping cluster IDs to (title, description, takeaways) tuples
        output_path: Path to save the output file
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Key Points Extraction Results\n\n")
            
            for cluster_id in sorted(cluster_sentences.keys()):
                title, description, takeaways = cluster_results[cluster_id]
                f.write(f"## Cluster {cluster_id}: {title}\n\n")
                
                # Write cluster description
                f.write("### Description\n")
                f.write(f"{description}\n\n")
                
                # Write sentences
                f.write("### Sentences\n")
                for sentence in cluster_sentences[cluster_id]:
                    f.write(f"- {sentence}\n")
                f.write("\n")
                
                # Write takeaways
                f.write("### Application Methods\n")
                f.write(f"{takeaways}\n\n")
                
        logger.info(f"Results saved to {output_path}")
    except Exception as e:
        logger.error(f"Error saving results: {str(e)}")
        raise

async def process_document(input_source: Union[str, BytesIO, BinaryIO], output_path: str = None, file_type: str = None) -> Dict:
    """
    Main function to process a document and extract key points.
    
    Args:
        input_source: Either file path (str) or BytesIO/file-like object
        output_path: Optional path to save output markdown file
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        Dictionary containing processing results
    """
    try:
        # Create default output path if not provided
        if output_path is None and isinstance(input_source, str):
            base_name = os.path.splitext(os.path.basename(input_source))[0]
            output_path = f"{base_name}_key_points.md"
            logger.info(f"No output path provided, using: {output_path}")
        elif output_path is None:
            output_path = "document_key_points.md"
            logger.info(f"No output path provided, using default: {output_path}")
        
        # Load and split document
        logger.info("Loading and splitting document...")
        sentences = load_and_split_document(input_source, file_type)
        
        # Special case for documents with only one sentence
        if len(sentences) == 1:
            logger.warning("Document contains only 1 sentence - simplified processing")
            single_sentence = sentences[0]
            
            # Create a simple result with just one cluster
            result = {
                "clusters": {
                    "0": {
                        "title": "Document Content",
                        "description": "The entire document content",
                        "sentences": sentences,
                        "takeaways": "Document is too brief for detailed analysis. Consider adding more content."
                    }
                },
                "metadata": {
                    "sentence_count": 1,
                    "cluster_count": 1,
                    "noise_points": 0
                }
            }
            
            # Save a simplified result if output path is specified
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write("# Document Analysis Results\n\n")
                    f.write("## Document Content\n\n")
                    f.write(f"- {single_sentence}\n\n")
                    f.write("### Note\n")
                    f.write("Document is too brief for detailed analysis. Consider adding more content.\n")
                
                logger.info(f"Saved simplified results to {output_path}")
            
            logger.info("Processing complete with simplified approach!")
            return result
        
        # Continue with normal processing for 2+ sentences
        # Generate embeddings
        logger.info("Generating sentence embeddings...")
        embeddings = generate_sentence_embeddings(sentences)
        
        # Cluster sentences
        logger.info("Clustering sentences...")
        cluster_labels = perform_clustering(embeddings)
        
        # Log clustering results
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_count = len(unique_clusters) - (1 if -1 in unique_clusters else 0)
        logger.info(f"Clustering results: {cluster_count} clusters, {noise_count} noise points")
        
        # Handle the case where all points are noise (no clusters formed)
        if noise_count == len(sentences) and noise_count > 0:
            logger.warning(f"All {noise_count} sentences were classified as noise - creating a single cluster")
            
            # Force all sentences into a single cluster
            cluster_labels = np.zeros(len(sentences), dtype=int)
            unique_clusters = {0}
            noise_count = 0
            cluster_count = 1
            logger.info("Converted all noise points to a single cluster")
        
        # Select all sentences for each cluster
        logger.info("Selecting sentences for each cluster...")
        cluster_sentences = select_all_sentences(sentences, cluster_labels)
        
        # Generate cluster titles and takeaways
        logger.info("Generating cluster titles and application methods...")
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        
        # Save results
        if output_path:
            logger.info(f"Saving results to {output_path}...")
            save_results(cluster_sentences, cluster_results, output_path)
        
        # Create result dictionary
        result = {
            "clusters": {},
            "metadata": {
                "sentence_count": len(sentences),
                "cluster_count": len(unique_clusters) - (1 if -1 in unique_clusters else 0),
                "noise_points": noise_count
            }
        }
        
        # Add cluster info to result
        for cluster_id in sorted(cluster_sentences.keys()):
            title, description, takeaways = cluster_results[cluster_id]
            result["clusters"][str(cluster_id)] = {
                "title": title,
                "description": description,
                "sentences": cluster_sentences[cluster_id],
                "takeaways": takeaways
            }
        
        logger.info("Processing complete!")
        return result
        
    except Exception as e:
        logger.error(f"Error processing document: {type(e).__name__}: {str(e)}")
        logger.error(f"Error details: {str(e)}")
        raise

async def generate_cluster_connections(clusters_list: List[Dict], sentences: List[str]) -> str:
    """
    Generate connections and relationships between different clusters to create a cohesive narrative.
    
    Args:
        clusters_list: List of cluster dictionaries containing titles, descriptions and takeaways
        sentences: Complete list of sentences from the document
        
    Returns:
        String with connections and relationships between clusters
    """
    try:
        # Generate a simple text representation of the clusters for the prompt
        clusters_text = ""
        for i, cluster in enumerate(clusters_list):
            clusters_text += f"Cluster {i}: {cluster['title']}\n"
            clusters_text += f"Description: {cluster['description'][:200]}...\n"
            clusters_text += f"Key sentences: {' '.join(cluster['sentences'])}\n\n"
        
        # Create a prompt to generate connections
        prompt = (
            f"You are a knowledge integration expert who can see connections and relationships between different topics.\n\n"
            f"DOCUMENT OVERVIEW:\nThe document contains {len(clusters_list)} main clusters or topics.\n\n"
            f"CLUSTERS:\n{clusters_text}\n\n"
            f"TASK: Generate one paragraph that connect these clusters together into a cohesive whole. "
            f"Explain how these topics relate to each other and how they contribute to a unified understanding.\n\n"
            f"REQUIREMENTS:\n"
            f"1. The connections should be substantive, specific, and insightful, not just vague generalities\n"
            f"2. IMPORTANT: Use the EXACT SAME LANGUAGE as the source material\n"
            f"3. If the source is in a non-English language like Vietnamese, your response must also be in that language\n"
            f"4. Highlight how these topics work together and why they are placed in the same document\n"
            f"5. Create a narrative flow that connects all clusters\n"
            f"6. Include practical insights about how these topics complement each other\n\n"
            f"CONNECTIONS BETWEEN CLUSTERS:"
        )
        
        # Get response from LLM
        response = await invoke_llm_with_retry(LLM, prompt, temperature=0.3, max_tokens=500)
        return response.content.strip()
    except Exception as e:
        logger.error(f"Error generating cluster connections: {type(e).__name__}: {str(e)}")
        return "Unable to generate connections between clusters due to an error."

async def understand_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> Dict:
    """
    Function to understand the document and extract key points without saving to a file.
    Ready for API endpoint integration.
    
    Args:
        input_source: Either file path (str) or BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided

    Returns:
        Dictionary containing processing results with clusters and insights
    """
    try:
        # Load and split document
        logger.info("Loading and splitting document...")
        sentences = load_and_split_document(input_source, file_type)
        
        # Detect document language
        sample_text = " ".join(sentences[:3]) if len(sentences) > 3 else " ".join(sentences)
        is_english = all(ord(c) < 128 for c in sample_text.replace('\n', ' ').replace(' ', ''))
        language = "English" if is_english else "Non-English (possibly multilingual)"
        logger.info(f"Detected language: {language}")
        
        # Special case for documents with only one sentence
        if len(sentences) == 1:
            logger.warning("Document contains only 1 sentence - simplified processing")
            single_sentence = sentences[0]
            
            # Create a simple result with just one cluster
            result = {
                "success": True,
                "document_insights": {
                    "metadata": {
                        "sentence_count": 1,
                        "cluster_count": 1,
                        "noise_points": 0,
                        "language": language,
                        "processing_level": "minimal"
                    },
                    "summary": single_sentence[:100] + ("..." if len(single_sentence) > 100 else ""),
                    "clusters": [
                        {
                            "id": "0",
                            "title": "Document Content",
                            "description": "The entire document content",
                            "sentences": sentences,
                            "takeaways": "Document is too brief for detailed analysis. Consider adding more content."
                        }
                    ],
                    "connections": "Document contains only a single topic, so no connections are necessary."
                }
            }
            logger.info("Processing complete with simplified approach!")
            return result
        
        # Continue with normal processing for 2+ sentences
        # Generate embeddings
        logger.info("Generating sentence embeddings...")
        embeddings = generate_sentence_embeddings(sentences)
        
        # Cluster sentences
        logger.info("Clustering sentences...")
        cluster_labels = perform_clustering(embeddings)
        
        # Log clustering results
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_count = len(unique_clusters) - (1 if -1 in unique_clusters else 0)
        logger.info(f"Clustering results: {cluster_count} clusters, {noise_count} noise points")
        
        # Handle the case where all or most points are noise (no clusters formed)
        if noise_count == len(sentences) or (noise_count > 0 and cluster_count == 0):
            logger.warning(f"All {noise_count} sentences were classified as noise - creating a single cluster")
            
            # Force all sentences into a single cluster (including noise points)
            cluster_labels = np.zeros(len(sentences), dtype=int)
            unique_clusters = {0}
            noise_count = 0
            cluster_count = 1
            logger.info("Converted all noise points to a single cluster")
            
            # Log the sentences that will be in this cluster
            for i, sentence in enumerate(sentences):
                logger.debug(f"Sentence {i} in cluster 0: {sentence[:50]}...")
        
        # Select all sentences for each cluster
        logger.info("Selecting sentences for each cluster...")
        cluster_sentences = select_all_sentences(sentences, cluster_labels)
        
        # Additional check - if we still have no clusters after selection, create one from all sentences
        if not cluster_sentences and len(sentences) > 0:
            logger.warning("No clusters formed after selection - creating a manual cluster with all sentences")
            cluster_sentences = {0: sentences}
            
        # Generate cluster titles and takeaways
        logger.info(f"Generating cluster titles and application methods for {len(cluster_sentences)} clusters...")
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        
        # Generate document summary from all cluster titles
        all_titles = [cluster_results[cid][0] for cid in sorted(cluster_sentences.keys())]
        summary = f"This document covers the following topics: {', '.join(all_titles)}"
        
        # Create result dictionary in API-friendly format
        clusters_list = []
        for cluster_id in sorted(cluster_sentences.keys()):
            title, description, takeaways = cluster_results[cluster_id]
            clusters_list.append({
                "id": str(cluster_id),
                "title": title,
                "description": description,
                "sentences": cluster_sentences[cluster_id],
                "takeaways": takeaways,
                "sentence_count": len(cluster_sentences[cluster_id])
            })
        
        # Generate connections between clusters if there are multiple
        logger.info(f"Generating connections between {len(clusters_list)} clusters...")
        connections = await generate_cluster_connections(clusters_list, sentences)
        logger.info("Generated cluster connections")
        
        result = {
            "success": True,
            "document_insights": {
                "metadata": {
                    "sentence_count": len(sentences),
                    "cluster_count": cluster_count,
                    "noise_points": noise_count,
                    "language": language,
                    "processing_level": "full"
                },
                "summary": summary,
                "clusters": clusters_list,
                "connections": connections
            }
        }
        
        logger.info("Document understanding complete!")
        return result
        
    except Exception as e:
        logger.error(f"Error understanding document: {type(e).__name__}: {str(e)}")
        logger.error(f"Error details: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }

async def save_document_insights(document_insight: str = "", user_id: str = "", mode: str = "default", bank: str = "") -> bool:
    """
    Save document insights (that were generated by understand_document) to the vector database.
    
    Args:
        document_insight: JSON string containing document insights from understand_document
        user_id: User ID for metadata
        mode: Processing mode (e.g., "default", "pretrain")
        bank: Namespace for vector database storage
        
    Returns:
        bool: True if saving was successful, False otherwise
    """
    doc_id = str(uuid.uuid4())
    logger.info(f"Saving document insights with bank={bank}")
    
    try:
        # Parse the document_insight JSON string
        if not document_insight or not document_insight.strip():
            logger.error("No document insights provided")
            return False
        
        # Parse JSON string to dictionary
        try:
            insights_data = json.loads(document_insight) if isinstance(document_insight, str) else document_insight
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse document insights JSON: {e}")
            return False
        
        # Extract document insights - handle both formats:
        # 1. Format from understand_document: {"success": true, "document_insights": {...}}
        # 2. Direct format: {"clusters": [...], "connections": "...", ...}
        doc_insights = {}
        if "document_insights" in insights_data:
            # Format from understand_document function
            doc_insights = insights_data["document_insights"]
        elif "clusters" in insights_data:
            # Direct format - insights themselves are at the top level
            doc_insights = insights_data
        else:
            logger.error("Invalid document insights format: missing both document_insights and clusters fields")
            return False
        
        # Now extract the components we need
        clusters = doc_insights.get("clusters", [])
        connections = doc_insights.get("connections", "")
        summary = doc_insights.get("summary", "")
        
        if not clusters:
            logger.error("No clusters found in document insights")
            return False
            
        logger.info(f"Processing {len(clusters)} clusters from document insights")
        
        # Create tasks for saving each cluster
        processing_tasks = []
        
        for cluster in clusters:
            cluster_id = cluster.get("id", "unknown")
            cluster_title = cluster.get("title", "Untitled Cluster")
            cluster_description = cluster.get("description", "")
            cluster_sentences = " ".join(cluster.get("sentences", []))
            cluster_takeaways = cluster.get("takeaways", "")
            chunk_id = f"chunk_{doc_id}_{cluster_id}"
            
            # Create a combined text representation of this cluster
            # Include connections with each cluster to ensure they're linked
            combined_text = (
                f"Title: {cluster_title}\n"
                f"Description: {cluster_description}\n"
                f"Content: {cluster_sentences}\n"
                f"Takeaways: {cluster_takeaways}\n"
                f"Document Summary: {summary}\n"
                f"Cross-Cluster Connections: {connections}"
            )
            
            # Save as a searchable document chunk
            processing_tasks.append(
                save_training_with_chunk(
                    input=combined_text,
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=False
                )
            )
            
            # Also save the takeaways with connections for more cohesive retrieval
            combined_takeaways = (
                f"{cluster_takeaways}\n\n"
                f"Related Context: {connections}"
            )
            
            processing_tasks.append(
                save_training_with_chunk(
                    input=combined_takeaways,
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=f"{chunk_id}_takeaways_with_context",
                    bank_name=bank,
                    is_raw=True
                )
            )
        
        # Process all save tasks
        logger.info(f"Executing {len(processing_tasks)} save tasks")
        results = await asyncio.gather(*processing_tasks, return_exceptions=True)
        
        # Track success rate
        success_count = 0
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Processing task failed: {result}")
            elif not result:
                logger.warning("Processing task returned False")
            else:
                success_count += 1
        
        if success_count == 0:
            logger.error("All processing tasks failed")
            return False
            
        logger.info(f"Successfully saved {success_count}/{len(processing_tasks)} chunks from document insights")
        return True
        
    except Exception as e:
        logger.error(f"Failed to save document insights: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False

