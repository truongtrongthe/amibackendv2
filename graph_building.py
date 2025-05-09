import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
import hdbscan
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import logging
from docx import Document
import pdfplumber
from io import BytesIO
import os
from typing import List, Tuple, Dict, Union, BinaryIO
import json
from database import save_training_with_chunk
import uuid
import asyncio
import networkx as nx
import matplotlib.pyplot as plt
from itertools import combinations
from sklearn.feature_extraction.text import TfidfVectorizer

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize models
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)
model = SentenceTransformer('sentence-transformers/LaBSE')

# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt, temperature=0.2, max_tokens=5000, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop, temperature=temperature, max_tokens=max_tokens)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

def load_and_split_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> List[str]:
    """
    Load a document from file path or BytesIO and split it into sentences.
    """
    logger.info("Loading and splitting document. Filetype: %s", file_type)
    try:
        if isinstance(input_source, str):
            file_path = input_source
            if not file_type:
                file_type = file_path.split('.')[-1].lower()
            if file_type == 'docx':
                doc = Document(file_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                with pdfplumber.open(file_path) as pdf:
                    pages_text = [f"[Page {i+1}] {page.extract_text() or ''}" for i, page in enumerate(pdf.pages)]
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            file_content = input_source
            file_content.seek(0)
            if not file_type:
                raise ValueError("file_type must be provided for file-like input")
            if file_type == 'docx':
                doc = Document(file_content)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                with pdfplumber.open(file_content) as pdf:
                    pages_text = [f"[Page {i+1}] {page.extract_text() or ''}" for i, page in enumerate(pdf.pages)]
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be a file path or file-like object")

        if not full_text.strip():
            raise ValueError("No text content extracted from the document")
        if len(full_text.split()) < 10:
            raise ValueError("Text content too short for meaningful processing")

        sentences = sent_tokenize(full_text)
        clean_sentences = [s.strip() for s in sentences if s.strip()]
        logger.info(f"Split text into {len(clean_sentences)} sentences")
        return clean_sentences
    except Exception as e:
        logger.error(f"Error loading document: {type(e).__name__}: {str(e)}")
        raise

def generate_sentence_embeddings(sentences: List[str]) -> np.ndarray:
    """Generate embeddings for sentences using LaBSE."""
    try:
        embeddings = model.encode(sentences, show_progress_bar=True)
        return embeddings
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise

def perform_clustering(embeddings: np.ndarray) -> np.ndarray:
    """Cluster sentences using HDBSCAN."""
    try:
        if len(embeddings) == 1:
            logger.warning("Only 1 sentence available - skipping clustering")
            return np.zeros(1, dtype=int)
        if len(embeddings) == 2:
            similarity = cosine_similarity(embeddings)[0, 1]
            logger.info(f"Similarity between 2 sentences: {similarity:.4f}")
            return np.zeros(2, dtype=int) if similarity > 0.5 else np.array([0, 1], dtype=int)

        clusterer = hdbscan.HDBSCAN(
            min_cluster_size=2, min_samples=1, metric='euclidean',
            cluster_selection_method='eom', cluster_selection_epsilon=0.25,
            alpha=0.8, prediction_data=True, allow_single_cluster=False
        )
        cluster_labels = clusterer.fit_predict(embeddings)

        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                        for label in unique_clusters if label != -1}
        logger.info(f"Clustering: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)} clusters, {noise_count} noise points")

        if len(unique_clusters) - (1 if -1 in unique_clusters else 0) < 2 and len(embeddings) > 3:
            logger.info("Retrying with different parameters...")
            clusterer = hdbscan.HDBSCAN(
                min_cluster_size=2, min_samples=1, metric='euclidean',
                cluster_selection_method='leaf', clusterüssion_epsilon=0.15,
                alpha=0.6, prediction_data=True, allow_single_cluster=False
            )
            cluster_labels = clusterer.fit_predict(embeddings)
            unique_clusters = set(cluster_labels)
            noise_count = sum(1 for label in cluster_labels if label == -1)
            cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                            for label in unique_clusters if label != -1}
            logger.info(f"Retry: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)} clusters, {noise_count} noise points")

        return cluster_labels
    except Exception as e:
        logger.error(f"Error clustering sentences: {type(e).__name__}: {str(e)}")
        raise

def select_all_sentences(sentences: List[str], cluster_labels: np.ndarray) -> Dict[int, List[str]]:
    """Select all sentences for each cluster."""
    try:
        cluster_sentences = {}
        for cluster_id in set(cluster_labels) - {-1}:
            cluster_mask = cluster_labels == cluster_id
            cluster_sentences[cluster_id] = [sentences[i] for i in range(len(sentences)) if cluster_mask[i]]
        return cluster_sentences
    except Exception as e:
        logger.error(f"Error selecting sentences: {str(e)}")
        raise

async def generate_takeaways(cluster_sentences: Dict[int, List[str]], sentences: List[str]) -> Dict[int, Tuple[str, str, str]]:
    """Generate titles, descriptions, and takeaways for each cluster, combining title and description into one LLM call."""
    try:
        doc_sample = "\n".join(sentences[:30] if len(sentences) > 30 else sentences)
        sem = asyncio.Semaphore(5)

        async def generate_title_and_description_for_cluster(cluster_id, cluster_sents, all_clusters):
            sentences_text = "\n".join([f"- {s}" for s in cluster_sents])
            other_clusters = {cid: sents for cid, sents in all_clusters.items() if cid != cluster_id}
            other_text = "\n".join([f"Cluster {cid}:\n" + "\n".join([f"- {s}" for s in sents[:2]]) 
                                   for cid, sents in other_clusters.items()])
            prompt = (
                f"You are a knowledge organization expert specializing in business insights.\n\n"
                f"TASK: For the following cluster, generate:\n"
                f"1. A CONCISE TITLE (5-8 words) capturing the underlying concept.\n"
                f"2. A BRIEF SUMMARY (2-3 sentences) of the cluster.\n"
                f"Return a JSON object with 'title' and 'description' fields.\n\n"
                f"GUIDELINES:\n"
                f"1. Use EXACT SAME LANGUAGE as sentences (no translation).\n"
                f"2. Title must be specific, informative, and use precise vocabulary.\n"
                f"3. Description must focus on core concept, max 2-3 sentences.\n"
                f"4. Incorporate technical terms and maintain tone.\n"
                f"5. Ensure business relevance (e.g., strategic, operational insights).\n"
                f"CURRENT CLUSTER SENTENCES:\n{sentences_text}\n\n"
                f"OTHER CLUSTERS:\n{other_text}\n\n"
                f"RESPONSE (JSON):\n"
                f"{{\"title\": \"...\", \"description\": \"...\"}}"
            )
            response = await invoke_llm_with_retry(LLM, prompt, max_tokens=200)
            try:
                # Strip any markdown formatting that might be in the response
                content = response.content.strip()
                if content.startswith("```json"):
                    content = content.replace("```json", "", 1)
                if content.startswith("```"):
                    content = content.replace("```", "", 1)
                if content.endswith("```"):
                    content = content[:-3]
                    
                content = content.strip()
                
                # Try to parse the JSON
                result = json.loads(content)
                if "title" not in result or "description" not in result:
                    raise ValueError("Missing title or description in response")
                return result["title"].strip('"'), result["description"].strip()
            except Exception as e:
                logger.warning(f"Failed to parse LLM response: {str(e)}. Response content: {response.content[:100]}...")
                # Generate a fallback title and description based on the first sentence
                if cluster_sents:
                    first_sent = cluster_sents[0]
                    fallback_title = first_sent[:30] + ("..." if len(first_sent) > 30 else "")
                    fallback_desc = f"Cluster containing {len(cluster_sents)} sentences about {fallback_title}"
                    return fallback_title, fallback_desc
                return "Unnamed Cluster", "No description available."

        async def generate_takeaway_for_cluster(cluster_sents):
            sentences_text = "\n".join([f"- {s}" for s in cluster_sents])
            prompt = (
                f"You are a knowledge application expert.\n\n"
                f"DOCUMENT CONTEXT:\n{doc_sample[:500]}...\n\n"
                f"TASK: Generate detailed, practical takeaways on HOW TO APPLY these insights.\n"
                f"IMPORTANT: Use EXACT SAME LANGUAGE as sentences.\n"
                f"GUIDELINES:\n"
                f"1. Focus on METHODS OF APPLICATION\n"
                f"2. Infer domain from content\n"
                f"3. Use numbered steps or clear methodology\n"
                f"4. Provide SPECIFIC EXAMPLE SCRIPTS or templates\n"
                f"5. Include step-by-step instructions\n"
                f"6. Format as 'Application Method: [title]' followed by steps\n"
                f"7. Maintain original language\n"
                f"SENTENCES:\n{sentences_text}\n\n"
                f"APPLICATION METHODS:"
            )
            response = await invoke_llm_with_retry(LLM, prompt, temperature=0.05)
            return response.content.strip()

        async def process_cluster(cluster_id, cluster_sents):
            async with sem:
                title_desc, takeaway = await asyncio.gather(
                    generate_title_and_description_for_cluster(cluster_id, cluster_sents, cluster_sentences),
                    generate_takeaway_for_cluster(cluster_sents)
                )
                title, description = title_desc
                logger.info(f"Processed cluster {cluster_id}")
                return cluster_id, (title, description, takeaway)

        tasks = [process_cluster(cid, sents) for cid, sents in cluster_sentences.items()]
        results_list = await asyncio.gather(*tasks)
        return {cid: result for cid, result in results_list}
    except Exception as e:
        logger.error(f"Error generating takeaways: {str(e)}")
        raise

def validate_connection(connection_text: str, cluster1_sents: List[str], cluster2_sents: List[str]) -> bool:
    """Validate connection description for semantic relevance."""
    connection_embedding = model.encode([connection_text])[0]
    cluster1_centroid = np.mean(model.encode(cluster1_sents), axis=0)
    cluster2_centroid = np.mean(model.encode(cluster2_sents), axis=0)
    sim1 = cosine_similarity([connection_embedding], [cluster1_centroid])[0][0]
    sim2 = cosine_similarity([connection_embedding], [cluster2_centroid])[0][0]
    return sim1 > 0.4 and sim2 > 0.4

async def generate_cluster_connections(clusters_list: List[Dict], sentences: List[str]) -> Tuple[nx.Graph, List[Dict]]:
    """Generate a networkx graph and connections for clusters."""
    try:
        G = nx.Graph()
        cluster_sentences = {}
        cluster_embeddings = {}
        for cluster in clusters_list:
            cid = cluster["id"]
            G.add_node(cid, title=cluster["title"], description=cluster["description"],
                      sentences=cluster["sentences"], sentence_count=cluster["sentence_count"])
            cluster_sentences[cid] = cluster["sentences"]
            cluster_embeddings[cid] = np.mean(model.encode(cluster["sentences"]), axis=0)

        connections = []
        for cid1, cid2 in combinations(cluster_sentences.keys(), 2):
            sim = cosine_similarity([cluster_embeddings[cid1]], [cluster_embeddings[cid2]])[0][0]
            if sim > 0.4:
                G.add_edge(cid1, cid2, weight=sim, method="embedding")
                connections.append({
                    "cluster_pair": f"{cid1}-{cid2}",
                    "description": f"Clusters {cid1} and {cid2} share thematic similarities (similarity: {sim:.2f}).",
                    "strength": sim,
                    "method": "embedding"
                })

        vectorizer = TfidfVectorizer(max_features=10)
        cluster_texts = [" ".join(sents) for sents in cluster_sentences.values()]
        tfidf_matrix = vectorizer.fit_transform(cluster_texts)
        feature_names = vectorizer.get_feature_names_out()
        for cid1, cid2 in combinations(cluster_sentences.keys(), 2):
            keywords1 = set(feature_names[tfidf_matrix[list(cluster_sentences.keys()).index(cid1)].toarray()[0].argsort()[-5:]])
            keywords2 = set(feature_names[tfidf_matrix[list(cluster_sentences.keys()).index(cid2)].toarray()[0].argsort()[-5:]])
            common_keywords = keywords1.intersection(keywords2)
            if common_keywords and not G.has_edge(cid1, cid2):
                G.add_edge(cid1, cid2, weight=0.5, method="keyword")
                connections.append({
                    "cluster_pair": f"{cid1}-{cid2}",
                    "description": f"Clusters {cid1} and {cid2} share focus on {', '.join(common_keywords)}.",
                    "strength": 0.5,
                    "method": "keyword"
                })

        high_value_pairs = [(c["cluster_pair"], c["strength"]) for c in connections if c["strength"] > 0.7]
        if high_value_pairs:
            clusters_text = "\n".join(
                f"Cluster {c['id']}: {c['title']}\nDescription: {c['description'][:200]}...\n"
                for c in clusters_list
            )
            prompt = (
                f"You are a business strategy expert.\n"
                f"CLUSTERS:\n{clusters_text}\n"
                f"TASK: Generate business-relevant connection descriptions for the following pairs.\n"
                f"Return JSON array: [{{\"cluster_pair\": \"0-1\", \"description\": \"...\"}}]\n"
                f"Use EXACT SAME LANGUAGE as source.\n"
                f"Focus on strategic/operational relationships.\n"
                f"PAIRS: {json.dumps([p[0] for p in high_value_pairs])}\n"
                f"CONNECTIONS (JSON):"
            )
            response = await invoke_llm_with_retry(LLM, prompt, temperature=0.3, max_tokens=500)
            llm_connections = json.loads(response.content.strip())
            for conn in connections:
                for llm_conn in llm_connections:
                    if conn["cluster_pair"] == llm_conn["cluster_pair"]:
                        conn["description"] = llm_conn["description"]
                        conn["method"] = "llm"
                        cid1, cid2 = conn["cluster_pair"].split("-")
                        if G.has_edge(cid1, cid2):
                            G.edges[cid1, cid2]["method"] = "llm"

        final_connections = [
            conn for conn in connections
            if validate_connection(conn["description"], cluster_sentences[conn["cluster_pair"].split("-")[0]], 
                                  cluster_sentences[conn["cluster_pair"].split("-")[1]])
        ]
        logger.info(f"Generated {len(final_connections)} valid connections")
        return G, final_connections
    except Exception as e:
        logger.error(f"Error generating connections: {type(e).__name__}: {str(e)}")
        return nx.Graph(), [{"cluster_pair": "all", "description": "Unable to generate connections.", "strength": 0}]

def visualize_knowledge_graph(graph: nx.Graph, output_path: str):
    """Visualize the knowledge graph, highlighting new clusters."""
    plt.figure(figsize=(12, 8))
    pos = nx.spring_layout(graph, k=0.5, iterations=50)
    new_nodes = [n for n in graph.nodes if "_new" in n]
    existing_nodes = [n for n in graph.nodes if "_new" not in n]
    nx.draw_networkx_nodes(graph, pos, nodelist=existing_nodes, node_color="lightblue", node_size=1000, label="Existing Clusters")
    nx.draw_networkx_nodes(graph, pos, nodelist=new_nodes, node_color="salmon", node_size=1000, label="New Clusters")
    edge_labels = {(u, v): f"{d['weight']:.2f}" for u, v, d in graph.edges(data=True)}
    nx.draw_networkx_edges(graph, pos, alpha=0.5)
    nx.draw_networkx_edge_labels(graph, pos, edge_labels=edge_labels, font_size=8)
    node_labels = {n: graph.nodes[n]["title"][:20] for n in graph.nodes}
    nx.draw_networkx_labels(graph, pos, labels=node_labels, font_size=10)
    plt.legend()
    plt.title("Knowledge Graph: Cluster Connections")
    plt.tight_layout()
    plt.savefig(output_path, dpi=300)
    plt.close()
    logger.info(f"Saved visualization to {output_path}")

async def understand_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None, 
                           existing_graph_path: str = "knowledge_graph.graphml") -> Dict:
    """
    Process a document, extract insights, and connect clusters to an existing graph.
    Preserves original takeaways format.
    """
    try:
        # Load existing graph
        if existing_graph_path and os.path.exists(existing_graph_path):
            existing_graph = nx.read_graphml(existing_graph_path)
            logger.info(f"Loaded existing graph with {len(existing_graph.nodes)} nodes")
        else:
            existing_graph = nx.Graph()
            logger.info("No existing graph found; starting fresh")

        # Load and process document
        sentences = load_and_split_document(input_source, file_type)
        language = "English" if all(ord(c) < 128 for c in " ".join(sentences[:3]).replace('\n', ' ')) else "Non-English"

        # Handle single-sentence case
        if len(sentences) == 1:
            single_sentence = sentences[0]
            result = {
                "success": True,
                "document_insights": {
                    "metadata": {
                        "sentence_count": 1,
                        "cluster_count": 1,
                        "noise_points": 0,
                        "language": language,
                        "processing_level": "minimal"
                    },
                    "summary": single_sentence[:100] + ("..." if len(single_sentence) > 100 else ""),
                    "clusters": [{
                        "id": "0",
                        "title": "Document Content",
                        "description": "The entire document content",
                        "sentences": sentences,
                        "takeaways": "Document is too brief for detailed analysis."
                    }],
                    "connections": []
                }
            }
            existing_graph.add_node("0_new", title="Document Content", sentences=sentences,
                                  embedding=np.mean(model.encode(sentences), axis=0),
                                  document_id=str(uuid.uuid4()))
            nx.write_graphml(existing_graph, existing_graph_path)
            return result

        # Generate embeddings and clusters
        embeddings = generate_sentence_embeddings(sentences)
        cluster_labels = perform_clustering(embeddings)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_count = len(set(cluster_labels)) - (1 if noise_count > 0 else 0)

        if noise_count == len(sentences) or (noise_count > 0 and cluster_count == 0):
            logger.warning("All sentences classified as noise; creating single cluster")
            cluster_labels = np.zeros(len(sentences), dtype=int)
            noise_count = 0
            cluster_count = 1

        cluster_sentences = select_all_sentences(sentences, cluster_labels)
        if not cluster_sentences:
            logger.warning("No clusters formed; creating manual cluster")
            cluster_sentences = {0: sentences}

        # Generate cluster insights
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        clusters_list = [
            {
                "id": str(cid),
                "title": cluster_results[cid][0],
                "description": cluster_results[cid][1],
                "sentences": cluster_sentences[cid],
                "takeaways": cluster_results[cid][2],
                "sentence_count": len(cluster_sentences[cid])
            }
            for cid in sorted(cluster_sentences.keys())
        ]

        # Generate intra-document connections
        intra_graph, intra_connections = await generate_cluster_connections(clusters_list, sentences)

        # Connect to existing graph
        new_cluster_embeddings = {c["id"]: np.mean(model.encode(c["sentences"]), axis=0) for c in clusters_list}
        for cluster in clusters_list:
            cid = cluster["id"]
            node_id = f"{cid}_new"
            existing_graph.add_node(
                node_id, title=cluster["title"], description=cluster["description"],
                sentences=cluster["sentences"], embedding=new_cluster_embeddings[cid],
                sentence_count=cluster["sentence_count"], document_id=str(uuid.uuid4())
            )

        inter_connections = []
        for new_cid in new_cluster_embeddings:
            new_node_id = f"{new_cid}_new"
            for existing_node in existing_graph.nodes:
                if "_new" not in existing_node and "embedding" in existing_graph.nodes[existing_node]:
                    sim = cosine_similarity(
                        [new_cluster_embeddings[new_cid]],
                        [existing_graph.nodes[existing_node]["embedding"]]
                    )[0][0]
                    if sim > 0.5:
                        existing_graph.add_edge(new_node_id, existing_node, weight=sim, method="embedding")
                        inter_connections.append({
                            "cluster_pair": f"{new_node_id}-{existing_node}",
                            "description": f"New cluster {new_cid} connects to existing cluster {existing_node} (similarity: {sim:.2f}).",
                            "strength": sim,
                            "method": "embedding"
                        })

        # Save and visualize graph
        nx.write_graphml(existing_graph, existing_graph_path)
        visualize_knowledge_graph(existing_graph, f"knowledge_graph_{str(uuid.uuid4())}.png")

        # Combine connections
        all_connections = intra_connections + inter_connections
        summary = f"This document covers: {', '.join(c['title'] for c in clusters_list)}"

        result = {
            "success": True,
            "document_insights": {
                "metadata": {
                    "sentence_count": len(sentences),
                    "cluster_count": cluster_count,
                    "noise_points": noise_count,
                    "language": language,
                    "processing_level": "full"
                },
                "summary": summary,
                "clusters": clusters_list,
                "connections": all_connections
            }
        }
        logger.info("Document processing complete")
        return result
    except Exception as e:
        logger.error(f"Error processing document: {type(e).__name__}: {str(e)}")
        return {"success": False, "error": str(e), "error_type": type(e).__name__}

async def save_document_insights(document_insight: str = "", user_id: str = "", mode: str = "default", bank: str = "") -> bool:
    """Save document insights to vector database."""
    doc_id = str(uuid.uuid4())
    logger.info(f"Saving insights with bank={bank}")
    try:
        insights_data = json.loads(document_insight) if isinstance(document_insight, str) else document_insight
        doc_insights = insights_data.get("document_insights", insights_data)
        clusters = doc_insights.get("clusters", [])
        connections = doc_insights.get("connections", "")
        summary = doc_insights.get("summary", "")

        if not clusters:
            logger.error("No clusters found")
            return False

        processing_tasks = []
        for cluster in clusters:
            cluster_id = cluster.get("id", "unknown")
            cluster_title = cluster.get("title", "Untitled Cluster")
            cluster_description = cluster.get("description", "")
            cluster_sentences = " ".join(cluster.get("sentences", []))
            cluster_takeaways = cluster.get("takeaways", "")
            chunk_id = f"chunk_{doc_id}_{cluster_id}"
            combined_text = (
                f"Title: {cluster_title}\nDescription: {cluster_description}\n"
                f"Content: {cluster_sentences}\nTakeaways: {cluster_takeaways}\n"
                f"Document Summary: {summary}\nCross-Cluster Connections: {connections}"
            )
            processing_tasks.append(
                save_training_with_chunk(
                    input=combined_text, user_id=user_id, mode=mode,
                    doc_id=doc_id, chunk_id=chunk_id, bank_name=bank, is_raw=False
                )
            )

        results = await asyncio.gather(*processing_tasks, return_exceptions=True)
        success_count = sum(1 for r in results if not isinstance(r, Exception) and r)
        logger.info(f"Saved {success_count}/{len(processing_tasks)} chunks")
        return success_count > 0
    except Exception as e:
        logger.error(f"Failed to save insights: {str(e)}")
        return False

async def understand_cluster(sentences: List[str]) -> Dict:
    """Process pre-clustered sentences, preserving takeaways format."""
    try:
        if not sentences:
            return {"success": False, "error": "No sentences provided", "error_type": "InputError"}

        language = "English" if all(ord(c) < 128 for c in " ".join(sentences[:3]).replace('\n', ' ')) else "Non-English"
        if len(sentences) == 1:
            single_sentence = sentences[0]
            return {
                "success": True,
                "document_insights": {
                    "metadata": {
                        "sentence_count": 1,
                        "cluster_count": 1,
                        "noise_points": 0,
                        "language": language,
                        "processing_level": "minimal"
                    },
                    "summary": single_sentence[:100] + ("..." if len(single_sentence) > 100 else ""),
                    "clusters": [{
                        "id": "0",
                        "title": "Content Summary",
                        "description": "The entire content",
                        "sentences": sentences,
                        "takeaways": "Content is too brief for detailed analysis."
                    }],
                    "connections": []
                }
            }

        cluster_sentences = {0: sentences}
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        title, description, takeaways = cluster_results[0]
        clusters_list = [{
            "id": "0",
            "title": title,
            "description": description,
            "sentences": sentences,
            "takeaways": takeaways,
            "sentence_count": len(sentences)
        }]

        return {
            "success": True,
            "document_insights": {
                "metadata": {
                    "sentence_count": len(sentences),
                    "cluster_count": 1,
                    "noise_points": 0,
                    "language": language,
                    "processing_level": "full"
                },
                "summary": f"This content covers: {title}",
                "clusters": clusters_list,
                "connections": []
            }
        }
    except Exception as e:
        logger.error(f"Error understanding sentences: {type(e).__name__}: {str(e)}")
        return {"success": False, "error": str(e), "error_type": type(e).__name__}
