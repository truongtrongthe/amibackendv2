import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
import hdbscan
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import logging
from docx import Document
import pdfplumber
from io import BytesIO
import os
from typing import List, Tuple, Dict, Union, BinaryIO
import json
from database import save_training_with_chunk
import uuid
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the LLM client with timeout
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)
model = SentenceTransformer('sentence-transformers/LaBSE')
# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt, temperature=0.2, max_tokens=5000, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop, temperature=temperature, max_tokens=max_tokens)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

def load_and_split_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> List[str]:
    """
    Load a document from file path or BytesIO and split it into sentences.
    
    Args:
        input_source: Either a file path (str) or a BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        List of sentences
    """
    logger.info("Loading and splitting document.Filetype: %s", file_type)
    try:
        # Determine whether we have a file path or a file-like object
        if isinstance(input_source, str):
            # It's a file path
            file_path = input_source
            if not file_type:
                # Detect file type from extension if not provided
                file_type = file_path.split('.')[-1].lower()
                
            # Different handling based on file type
            if file_type == 'docx':
                logger.info(f"Loading DOCX from file path: {file_path}")
                doc = Document(file_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info(f"Loading PDF from file path: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            # It's a file-like object (BytesIO)
            file_content = input_source
            
            # Reset position to beginning
            file_content.seek(0)
            
            # Must have file_type for file-like objects
            if not file_type:
                raise ValueError("file_type must be provided when using file-like input")
                
            if file_type == 'docx':
                logger.info("Loading DOCX from BytesIO/file-like object")
                doc = Document(file_content)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info("Loading PDF from BytesIO/file-like object")
                with pdfplumber.open(file_content) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be either a file path or a file-like object")
            
        # Log text extraction info
        logger.info(f"Extracted {len(full_text)} characters of text")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")
        
        # Check if we have enough content
        if not full_text.strip():
            logger.warning("No text content extracted")
            raise ValueError("No text content could be extracted from the document")
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            raise ValueError("Text content too short for meaningful processing")
        
        # Process extracted text into sentences
        sentences = sent_tokenize(full_text)
        clean_sentences = [s.strip() for s in sentences if s.strip()]
        logger.info(f"Split text into {len(clean_sentences)} sentences")
        
        return clean_sentences
        
    except Exception as e:
        logger.error(f"Error loading document: {type(e).__name__}: {str(e)}")
        raise

def generate_sentence_embeddings(sentences: List[str]) -> np.ndarray:
    """
    Generate embeddings for sentences using LaBSE.
    
    Args:
        sentences: List of sentences to embed
        
    Returns:
        numpy array of embeddings
    """
    try:
        
        embeddings = model.encode(sentences, show_progress_bar=True)
        return embeddings
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise

def perform_clustering(embeddings: np.ndarray) -> np.ndarray:
    """
    Cluster sentences using HDBSCAN.
    
    Args:
        embeddings: Sentence embeddings
        
    Returns:
        Array of cluster labels
    """
    try:
        # Add special case handling for very few sentences
        if len(embeddings) == 1:
            logger.warning("Only 1 sentence available - skipping clustering")
            # Put the single point into cluster 0
            return np.zeros(1, dtype=int)
            
        # Special handling for exactly 2 sentences - still try to cluster but with modified parameters
        if len(embeddings) == 2:
            logger.warning("Only 2 sentences available - using simplified clustering")
            # Calculate similarity between the 2 sentences
            similarity = cosine_similarity(embeddings)[0, 1]
            logger.info(f"Similarity between the 2 sentences: {similarity:.4f}")
            
            # If they're similar enough, put them in the same cluster, otherwise different clusters
            if similarity > 0.5:  # Threshold can be adjusted
                logger.info("Sentences are similar - grouping in the same cluster")
                return np.zeros(2, dtype=int)
            else:
                logger.info("Sentences are different - putting in separate clusters")
                return np.array([0, 1], dtype=int)
        
        # For 3+ sentences, use standard HDBSCAN clustering
        # Further adjusted parameters for smaller, more focused clusters
        clusterer = hdbscan.HDBSCAN(
            min_cluster_size=2,           # Keep small clusters
            min_samples=1,                # Allow single points to form clusters
            metric='euclidean',
            cluster_selection_method='eom',
            cluster_selection_epsilon=0.25, # Further reduced to create even tighter clusters
            alpha=0.8,                    # Reduced to encourage more clusters
            prediction_data=True,         # Enable prediction data for better stability
            allow_single_cluster=False    # Prevent all points from being in one cluster
        )
        
        # Fit and predict
        cluster_labels = clusterer.fit_predict(embeddings)
        
        # Log clustering statistics
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                        for label in unique_clusters if label != -1}
        
        logger.info(f"Clustering statistics:")
        logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
        logger.info(f"  Noise points: {noise_count}")
        logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        # If we have too few clusters, try again with different parameters
        if len(unique_clusters) - (1 if -1 in unique_clusters else 0) < 2 and len(embeddings) > 3:
            logger.info("Too few clusters detected, retrying with different parameters...")
            clusterer = hdbscan.HDBSCAN(
                min_cluster_size=2,
                min_samples=1,
                metric='euclidean',
                cluster_selection_method='leaf',
                cluster_selection_epsilon=0.15,
                alpha=0.6,
                prediction_data=True,
                allow_single_cluster=False
            )
            cluster_labels = clusterer.fit_predict(embeddings)
            
            # Log new clustering statistics
            unique_clusters = set(cluster_labels)
            noise_count = sum(1 for label in cluster_labels if label == -1)
            cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                            for label in unique_clusters if label != -1}
            
            logger.info(f"Retry clustering statistics:")
            logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
            logger.info(f"  Noise points: {noise_count}")
            logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        return cluster_labels
    except Exception as e:
        logger.error(f"Error clustering sentences: {type(e).__name__}: {str(e)}")
        raise

def select_all_sentences(
    sentences: List[str],
    cluster_labels: np.ndarray
) -> Dict[int, List[str]]:
    """
    Select all sentences for each cluster, allowing sentences to appear in multiple clusters.
    
    Args:
        sentences: List of sentences
        cluster_labels: Cluster labels for each sentence
        
    Returns:
        Dictionary mapping cluster IDs to lists of sentences
    """
    try:
        cluster_sentences = {}
        unique_clusters = set(cluster_labels)
        
        for cluster_id in unique_clusters:
            if cluster_id == -1:  # Skip noise points
                continue
                
            # Get all sentences for this cluster
            cluster_mask = cluster_labels == cluster_id
            cluster_sentences[cluster_id] = [
                sentences[i] for i in range(len(sentences)) if cluster_mask[i]
            ]
            
        return cluster_sentences
    except Exception as e:
        logger.error(f"Error selecting sentences: {str(e)}")
        raise

async def generate_cluster_title(sentences: List[str]) -> str:
    """
    Generate a descriptive title for a cluster using GPT-4o.
    
    Args:
        sentences: List of sentences
        
    Returns:
        A descriptive title for the cluster
    """
    try:
        sentences_text = "\n".join([f"- {s}" for s in sentences])
        prompt = (
            f"You are a knowledge organization expert with excellent abstraction skills.\n\n"
            f"TASK: Based on the following sentences, generate a COMPREHENSIVE yet CONCISE TITLE (5-8 words) that perfectly captures the underlying concept, theme, or principle.\n"
            f"If the sentences are not clear enough, abstract and elevate them to a higher-level concept.\n\n"
            f"GUIDELINES:\n"
            f"1. The title MUST be in the EXACT SAME LANGUAGE as the sentences\n"
            f"2. Create a highly informative and specific title that reveals the core essence\n"
            f"3. Use precise, rich vocabulary that communicates depth of understanding\n"
            f"4. Incorporate key technical terms and domain-specific vocabulary when relevant\n"
            f"5. Balance abstraction with specificity - be conceptual yet clearly descriptive\n"
            f"6. Make the title meaningful even to someone who hasn't read the sentences\n"
            f"7. Identify patterns and principles underlying the specific examples\n"
            f"8. Aim for sophistication and elegance in your wording\n"
            f"9. Ensure the title would serve as an excellent reference point\n\n"
            f"SENTENCES:\n{sentences_text}\n\n"
            f"COMPREHENSIVE TITLE:"
        )
        
        response = await invoke_llm_with_retry(LLM, prompt, max_tokens=50)
        title = response.content.strip()
        
        # Remove quotes if present
        if title.startswith('"') and title.endswith('"'):
            title = title[1:-1]
        
        return title
    except Exception as e:
        logger.error(f"Error generating cluster title: {str(e)}")
        return "Unnamed Cluster"

async def generate_cluster_description(
    cluster_sentences: Dict[int, List[str]],
    cluster_id: int,
    all_clusters: Dict[int, List[str]]
) -> str:
    """
    Generate a description of the cluster and its relationships to other clusters.
    
    Args:
        cluster_sentences: Dictionary mapping cluster IDs to sentences
        cluster_id: ID of the current cluster
        all_clusters: Dictionary of all clusters
        
    Returns:
        A description of the cluster and its relationships
    """
    try:
        # Get sentences from current cluster
        current_sentences = cluster_sentences[cluster_id]
        
        # Get sentences from other clusters
        other_clusters = {cid: sents for cid, sents in all_clusters.items() 
                         if cid != cluster_id and cid != -1}
        
        # Prepare the prompt
        current_text = "\n".join([f"- {s}" for s in current_sentences])
        other_text = "\n".join([f"Cluster {cid}:\n" + "\n".join([f"- {s}" for s in sents[:2]])  # Only use first 2 sentences from other clusters
                              for cid, sents in other_clusters.items()])
        
        prompt = (
            f"You are a knowledge organization expert.\n\n"
            f"TASK: Create a CONCISE, BRIEF summary (2-3 sentences maximum) of the following cluster of sentences in the EXACT SAME LANGUAGE as the sentences.\n"
            f"DO NOT TRANSLATE TO ENGLISH. Use the same words, phrases, and expressions as in the original sentences.\n\n"
            f"EXAMPLE:\n"
            f"If the sentences are in Vietnamese, write the description in Vietnamese.\n"
            f"If the sentences are in Japanese, write the description in Japanese.\n"
            f"If the sentences are in Spanish, write the description in Spanish.\n\n"
            f"REQUIREMENTS:\n"
            f"1. Be EXTREMELY BRIEF - no more than 2-3 sentences total\n"
            f"2. Focus ONLY on the core concept or key insight from these sentences\n"
            f"3. Use the EXACT SAME LANGUAGE as the sentences - no translation allowed\n"
            f"4. Keep all technical terms and domain-specific vocabulary exactly as they appear\n"
            f"5. Maintain the same tone, style, and expressions\n"
            f"6. Preserve all cultural context and nuances\n"
            f"7. Extract the essence without unnecessary details\n"
            f"8. Be direct and to the point\n\n"
            f"CURRENT CLUSTER SENTENCES:\n{current_text}\n\n"
            f"OTHER CLUSTERS:\n{other_text}\n\n"
            f"BRIEF DESCRIPTION (2-3 sentences in the same language as the sentences above):"
        )
        
        response = await invoke_llm_with_retry(LLM, prompt, max_tokens=150)
        return response.content.strip()
    except Exception as e:
        logger.error(f"Error generating cluster description: {str(e)}")
        return ""

async def generate_takeaways(
    cluster_sentences: Dict[int, List[str]],
    sentences: List[str]
) -> Dict[int, Tuple[str, str, str]]:
    """
    Generate titles, descriptions, and application-focused takeaways for each cluster using GPT-4o.
    Uses true parallel processing with proper rate limiting for optimal performance.
    
    Args:
        cluster_sentences: Dictionary mapping cluster IDs to sentences
        sentences: Complete list of all sentences from the document for context
        
    Returns:
        Dictionary mapping cluster IDs to (title, description, takeaways) tuples
    """
    try:
        import asyncio
        from functools import partial
        
        # Get a sample of the full document for context
        doc_sample = "\n".join(sentences[:30] if len(sentences) > 30 else sentences)
        
        # Create helper functions for each type of generation
        async def generate_title_for_cluster(cluster_sents):
            sentences_text = "\n".join([f"- {s}" for s in cluster_sents])
            prompt = (
                f"You are a knowledge organization expert with excellent abstraction skills.\n\n"
                f"TASK: Based on the following sentences, generate a COMPREHENSIVE yet CONCISE TITLE (5-8 words) that perfectly captures the underlying concept, theme, or principle.\n"
                f"If the sentences are not clear enough, abstract and elevate them to a higher-level concept.\n\n"
                f"GUIDELINES:\n"
                f"1. The title MUST be in the EXACT SAME LANGUAGE as the sentences\n"
                f"2. Create a highly informative and specific title that reveals the core essence\n"
                f"3. Use precise, rich vocabulary that communicates depth of understanding\n"
                f"4. Incorporate key technical terms and domain-specific vocabulary when relevant\n"
                f"5. Balance abstraction with specificity - be conceptual yet clearly descriptive\n"
                f"6. Make the title meaningful even to someone who hasn't read the sentences\n"
                f"7. Identify patterns and principles underlying the specific examples\n"
                f"8. Aim for sophistication and elegance in your wording\n"
                f"9. Ensure the title would serve as an excellent reference point\n\n"
                f"SENTENCES:\n{sentences_text}\n\n"
                f"COMPREHENSIVE TITLE:"
            )
            response = await invoke_llm_with_retry(LLM, prompt, max_tokens=50)
            title = response.content.strip()
            # Remove quotes if present
            if title.startswith('"') and title.endswith('"'):
                title = title[1:-1]
            return title
        
        async def generate_description_for_cluster(cluster_id, all_clusters):
            # Get sentences from current cluster
            current_sentences = all_clusters[cluster_id]
            # Get sentences from other clusters
            other_clusters = {cid: sents for cid, sents in all_clusters.items() 
                            if cid != cluster_id and cid != -1}
            
            # Prepare the prompt
            current_text = "\n".join([f"- {s}" for s in current_sentences])
            other_text = "\n".join([f"Cluster {cid}:\n" + "\n".join([f"- {s}" for s in sents[:2]])  # Only use first 2 sentences from other clusters
                                for cid, sents in other_clusters.items()])
            
            prompt = (
                f"You are a knowledge organization expert.\n\n"
                f"TASK: Create a CONCISE, BRIEF summary (2-3 sentences maximum) of the following cluster of sentences in the EXACT SAME LANGUAGE as the sentences.\n"
                f"DO NOT TRANSLATE TO ENGLISH. Use the same words, phrases, and expressions as in the original sentences.\n\n"
                f"REQUIREMENTS:\n"
                f"1. Be EXTREMELY BRIEF - no more than 2-3 sentences total\n"
                f"2. Focus ONLY on the core concept or key insight from these sentences\n"
                f"3. Use the EXACT SAME LANGUAGE as the sentences - no translation allowed\n"
                f"4. Keep all technical terms and domain-specific vocabulary exactly as they appear\n"
                f"5. Maintain the same tone, style, and expressions\n"
                f"6. Extract the essence without unnecessary details\n"
                f"7. Be direct and to the point\n\n"
                f"CURRENT CLUSTER SENTENCES:\n{current_text}\n\n"
                f"BRIEF DESCRIPTION (2-3 sentences in the same language as the sentences above):"
            )
            response = await invoke_llm_with_retry(LLM, prompt, max_tokens=150)
            return response.content.strip()
            
        async def generate_takeaway_for_cluster(cluster_sents):
            sentences_text = "\n".join([f"- {s}" for s in cluster_sents])
            prompt = (
                f"You are a knowledge application expert specializing in practical implementation.\n\n"
                f"DOCUMENT CONTEXT (sample from the full document):\n{doc_sample[:500]}...\n\n"  # Limit context to 500 chars
                f"TASK: Analyze the following sentences and generate detailed, practical takeaways focusing SPECIFICALLY on HOW TO APPLY these insights in relevant practical contexts.\n"
                f"IMPORTANT: Your response MUST be in the SAME LANGUAGE as the sentences.\n\n"
                f"GUIDELINES:\n"
                f"1. Focus on METHODS OF APPLICATION - explain exactly HOW to apply these insights in the domain they belong to\n"
                f"2. Infer the domain/field from the content itself - adapt to whatever topic is being discussed\n"
                f"3. Structure your response with numbered steps or a clear methodology\n"
                f"4. Provide SPECIFIC EXAMPLE SCRIPTS, dialogues, or templates that demonstrate the application\n"
                f"5. Include step-by-step instructions for implementation\n"
                f"6. Format your response as 'Application Method: [title]' followed by the steps\n"
                f"7. IMPORTANT: Maintain the original language of the sentences in your response\n"
                f"8. Include specific factual details and nuanced observations from the text\n"
                f"9. Even if the input sentences are very short, extract meaningful principles and applications\n\n"
                f"SENTENCES TO ANALYZE:\n{sentences_text}\n\n"
                f"APPLICATION METHODS (with specific examples and step-by-step instructions):"
            )
            response = await invoke_llm_with_retry(LLM, prompt, temperature=0.05)
            return response.content.strip()
        
        # This processes all operations for a single cluster concurrently
        async def process_cluster(cluster_id, cluster_sents):
            # Run all three operations concurrently for this cluster
            title, description, takeaway = await asyncio.gather(
                generate_title_for_cluster(cluster_sents),
                generate_description_for_cluster(cluster_id, cluster_sentences),
                generate_takeaway_for_cluster(cluster_sents)
            )
            
            logger.info(f"Completed processing for cluster {cluster_id}")
            return cluster_id, (title, description, takeaway)
            
        # Create tasks for all clusters - we'll use a semaphore to limit concurrency
        sem = asyncio.Semaphore(5)  # Process up to 5 clusters concurrently to avoid rate limits
        
        async def process_with_rate_limit(cluster_id, cluster_sents):
            async with sem:
                return await process_cluster(cluster_id, cluster_sents)
        
        # Create task for each cluster
        logger.info(f"Beginning parallel processing for {len(cluster_sentences)} clusters...")
        tasks = [
            process_with_rate_limit(cluster_id, cluster_sents)
            for cluster_id, cluster_sents in cluster_sentences.items()
        ]
        
        # Run all tasks concurrently and get results
        results_list = await asyncio.gather(*tasks)
        
        # Convert list of results to dictionary
        return {cluster_id: result for cluster_id, result in results_list}
            
    except Exception as e:
        logger.error(f"Error generating takeaways: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise

async def generate_cluster_connections(clusters_list: List[Dict], sentences: List[str]) -> str:
    """
    Generate connections and relationships between different clusters to create a cohesive narrative.
    
    Args:
        clusters_list: List of cluster dictionaries containing titles, descriptions and takeaways
        sentences: Complete list of sentences from the document
        
    Returns:
        String with connections and relationships between clusters
    """
    try:
        # Generate a simple text representation of the clusters for the prompt
        clusters_text = ""
        for i, cluster in enumerate(clusters_list):
            clusters_text += f"Cluster {i}: {cluster['title']}\n"
            clusters_text += f"Description: {cluster['description'][:200]}...\n"
            clusters_text += f"Key sentences: {' '.join(cluster['sentences'])}\n\n"
        
        # Create a prompt to generate connections
        prompt = (
            f"You are a knowledge integration expert who can see connections and relationships between different topics.\n\n"
            f"DOCUMENT OVERVIEW:\nThe document contains {len(clusters_list)} main clusters or topics.\n\n"
            f"CLUSTERS:\n{clusters_text}\n\n"
            f"TASK: Generate one paragraph that connect these clusters together into a cohesive whole. "
            f"Explain how these topics relate to each other and how they contribute to a unified understanding.\n\n"
            f"REQUIREMENTS:\n"
            f"1. The connections should be substantive, specific, and insightful, not just vague generalities\n"
            f"2. IMPORTANT: Use the EXACT SAME LANGUAGE as the source material\n"
            f"3. If the source is in a non-English language like Vietnamese, your response must also be in that language\n"
            f"4. Highlight how these topics work together and why they are placed in the same document\n"
            f"5. Create a narrative flow that connects all clusters\n"
            f"6. Include practical insights about how these topics complement each other\n\n"
            f"CONNECTIONS BETWEEN CLUSTERS:"
        )
        
        # Get response from LLM
        response = await invoke_llm_with_retry(LLM, prompt, temperature=0.3, max_tokens=500)
        return response.content.strip()
    except Exception as e:
        logger.error(f"Error generating cluster connections: {type(e).__name__}: {str(e)}")
        return "Unable to generate connections between clusters due to an error."

async def understand_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> Dict:
    """
    Function to understand the document and extract key points without saving to a file.
    Ready for API endpoint integration.
    
    Args:
        input_source: Either file path (str) or BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided

    Returns:
        Dictionary containing processing results with clusters and insights
    """
    try:
        # Load and split document
        logger.info("Loading and splitting document...")
        sentences = load_and_split_document(input_source, file_type)
        
        # Detect document language
        sample_text = " ".join(sentences[:3]) if len(sentences) > 3 else " ".join(sentences)
        is_english = all(ord(c) < 128 for c in sample_text.replace('\n', ' ').replace(' ', ''))
        language = "English" if is_english else "Non-English (possibly multilingual)"
        logger.info(f"Detected language: {language}")
        
        # Special case for documents with only one sentence
        if len(sentences) == 1:
            logger.warning("Document contains only 1 sentence - simplified processing")
            single_sentence = sentences[0]
            
            # Create a simple result with just one cluster
            result = {
                "success": True,
                "document_insights": {
                    "metadata": {
                        "sentence_count": 1,
                        "cluster_count": 1,
                        "noise_points": 0,
                        "language": language,
                        "processing_level": "minimal"
                    },
                    "summary": single_sentence[:100] + ("..." if len(single_sentence) > 100 else ""),
                    "clusters": [
                        {
                            "id": "0",
                            "title": "Document Content",
                            "description": "The entire document content",
                            "sentences": sentences,
                            "takeaways": "Document is too brief for detailed analysis. Consider adding more content."
                        }
                    ],
                    "connections": "Document contains only a single topic, so no connections are necessary."
                }
            }
            logger.info("Processing complete with simplified approach!")
            return result
        
        # Continue with normal processing for 2+ sentences
        # Generate embeddings
        logger.info("Generating sentence embeddings...")
        embeddings = generate_sentence_embeddings(sentences)
        
        # Cluster sentences
        logger.info("Clustering sentences...")
        cluster_labels = perform_clustering(embeddings)
        
        # Log clustering results
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_count = len(unique_clusters) - (1 if -1 in unique_clusters else 0)
        logger.info(f"Clustering results: {cluster_count} clusters, {noise_count} noise points")
        
        # Handle the case where all or most points are noise (no clusters formed)
        if noise_count == len(sentences) or (noise_count > 0 and cluster_count == 0):
            logger.warning(f"All {noise_count} sentences were classified as noise - creating a single cluster")
            
            # Force all sentences into a single cluster (including noise points)
            cluster_labels = np.zeros(len(sentences), dtype=int)
            unique_clusters = {0}
            noise_count = 0
            cluster_count = 1
            logger.info("Converted all noise points to a single cluster")
            
            # Log the sentences that will be in this cluster
            for i, sentence in enumerate(sentences):
                logger.debug(f"Sentence {i} in cluster 0: {sentence[:50]}...")
        
        # Select all sentences for each cluster
        logger.info("Selecting sentences for each cluster...")
        cluster_sentences = select_all_sentences(sentences, cluster_labels)
        
        # Additional check - if we still have no clusters after selection, create one from all sentences
        if not cluster_sentences and len(sentences) > 0:
            logger.warning("No clusters formed after selection - creating a manual cluster with all sentences")
            cluster_sentences = {0: sentences}
            
        # Generate cluster titles and takeaways
        logger.info(f"Generating cluster titles and application methods for {len(cluster_sentences)} clusters...")
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        
        # Generate document summary from all cluster titles
        all_titles = [cluster_results[cid][0] for cid in sorted(cluster_sentences.keys())]
        summary = f"This document covers the following topics: {', '.join(all_titles)}"
        
        # Create result dictionary in API-friendly format
        clusters_list = []
        for cluster_id in sorted(cluster_sentences.keys()):
            title, description, takeaways = cluster_results[cluster_id]
            clusters_list.append({
                "id": str(cluster_id),
                "title": title,
                "description": description,
                "sentences": cluster_sentences[cluster_id],
                "takeaways": takeaways,
                "sentence_count": len(cluster_sentences[cluster_id])
            })
        
        # Generate connections between clusters if there are multiple
        logger.info(f"Generating connections between {len(clusters_list)} clusters...")
        #connections = await generate_cluster_connections(clusters_list, sentences)
        connections = ""
        logger.info("Generated cluster connections")
        
        result = {
            "success": True,
            "document_insights": {
                "metadata": {
                    "sentence_count": len(sentences),
                    "cluster_count": cluster_count,
                    "noise_points": noise_count,
                    "language": language,
                    "processing_level": "full"
                },
                "summary": summary,
                "clusters": clusters_list,
                "connections": connections
            }
        }
        
        logger.info("Document understanding complete!")
        return result
        
    except Exception as e:
        logger.error(f"Error understanding document: {type(e).__name__}: {str(e)}")
        logger.error(f"Error details: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }

async def save_document_insights(document_insight: str = "", user_id: str = "", mode: str = "default", bank: str = "") -> bool:
    """
    Save document insights (that were generated by understand_document) to the vector database.
    
    Args:
        document_insight: JSON string containing document insights from understand_document
        user_id: User ID for metadata
        mode: Processing mode (e.g., "default", "pretrain")
        bank: Namespace for vector database storage
        
    Returns:
        bool: True if saving was successful, False otherwise
    """
    doc_id = str(uuid.uuid4())
    logger.info(f"Saving document insights with bank={bank}")
    
    try:
        # Parse the document_insight JSON string
        if not document_insight or not document_insight.strip():
            logger.error("No document insights provided")
            return False
        
        # Parse JSON string to dictionary
        try:
            insights_data = json.loads(document_insight) if isinstance(document_insight, str) else document_insight
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse document insights JSON: {e}")
            return False
        
        # Extract document insights - handle both formats:
        # 1. Format from understand_document: {"success": true, "document_insights": {...}}
        # 2. Direct format: {"clusters": [...], "connections": "...", ...}
        doc_insights = {}
        if "document_insights" in insights_data:
            # Format from understand_document function
            doc_insights = insights_data["document_insights"]
        elif "clusters" in insights_data:
            # Direct format - insights themselves are at the top level
            doc_insights = insights_data
        else:
            logger.error("Invalid document insights format: missing both document_insights and clusters fields")
            return False
        
        # Now extract the components we need
        clusters = doc_insights.get("clusters", [])
        connections = doc_insights.get("connections", "")
        summary = doc_insights.get("summary", "")
        
        if not clusters:
            logger.error("No clusters found in document insights")
            return False
            
        logger.info(f"Processing {len(clusters)} clusters from document insights")
        
        # Create tasks for saving each cluster
        processing_tasks = []
        
        for cluster in clusters:
            cluster_id = cluster.get("id", "unknown")
            cluster_title = cluster.get("title", "Untitled Cluster")
            cluster_description = cluster.get("description", "")
            cluster_sentences = " ".join(cluster.get("sentences", []))
            cluster_takeaways = cluster.get("takeaways", "")
            chunk_id = f"chunk_{doc_id}_{cluster_id}"
            
            # Create a combined text representation of this cluster
            # Include connections with each cluster to ensure they're linked
            combined_text = (
                f"Title: {cluster_title}\n"
                f"Description: {cluster_description}\n"
                f"Content: {cluster_sentences}\n"
                f"Takeaways: {cluster_takeaways}\n"
                f"Document Summary: {summary}\n"
                f"Cross-Cluster Connections: {connections}"
            )
            
            # Save as a searchable document chunk
            processing_tasks.append(
                save_training_with_chunk(
                    input=combined_text,
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=False
                )
            )
            
        # Process all save tasks
        logger.info(f"Executing {len(processing_tasks)} save tasks")
        results = await asyncio.gather(*processing_tasks, return_exceptions=True)
        
        # Track success rate
        success_count = 0
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Processing task failed: {result}")
            elif not result:
                logger.warning("Processing task returned False")
            else:
                success_count += 1
        
        if success_count == 0:
            logger.error("All processing tasks failed")
            return False
            
        logger.info(f"Successfully saved {success_count}/{len(processing_tasks)} chunks from document insights")
        return True
        
    except Exception as e:
        logger.error(f"Failed to save document insights: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False

async def understand_cluster(sentences: List[str]) -> Dict:
    """
    Function to process a pre-clustered set of sentences and extract insights.
    Assumes the input sentences belong to a single cluster.
    Preserves the same response structure as understand_document for compatibility.
    
    Args:
        sentences: List of sentences to process (pre-clustered)
        
    Returns:
        Dictionary containing processing results in the same format as understand_document
    """
    try:
        # Validate input
        if not sentences:
            logger.warning("No sentences provided")
            return {
                "success": False,
                "error": "No sentences provided",
                "error_type": "InputError"
            }
            
        logger.info(f"Processing {len(sentences)} sentences")
        
        # Detect language
        sample_text = " ".join(sentences[:3]) if len(sentences) > 3 else " ".join(sentences)
        is_english = all(ord(c) < 128 for c in sample_text.replace('\n', ' ').replace(' ', ''))
        language = "English" if is_english else "Non-English (possibly multilingual)"
        logger.info(f"Detected language: {language}")
        
        # Special case for only one sentence
        if len(sentences) == 1:
            logger.warning("Only 1 sentence provided - simplified processing")
            single_sentence = sentences[0]
            
            # Create a simple result with just one cluster - matching understand_document format
            result = {
                "success": True,
                "document_insights": {
                    "metadata": {
                        "sentence_count": 1,
                        "cluster_count": 1,
                        "noise_points": 0,
                        "language": language,
                        "processing_level": "minimal"
                    },
                    "summary": single_sentence[:100] + ("..." if len(single_sentence) > 100 else ""),
                    "clusters": [
                        {
                            "id": "0",
                            "title": "Content Summary",
                            "description": "The entire content",
                            "sentences": sentences,
                            "takeaways": "Content is too brief for detailed analysis. Consider adding more sentences."
                        }
                    ],
                    "connections": ""
                }
            }
            logger.info("Processing complete with simplified approach")
            return result
        
        # Prepare the cluster data - assume all sentences belong to a single cluster
        # We'll use the cluster ID 0 as a key
        cluster_sentences = {0: sentences}
        
        # Generate cluster titles and takeaways
        logger.info(f"Generating title, description and application methods...")
        cluster_results = await generate_takeaways(cluster_sentences, sentences)
        
        # Get the results for the single cluster
        title, description, takeaways = cluster_results[0]
        
        # Create result dictionary in the same format as understand_document for compatibility
        clusters_list = [{
            "id": "0",
            "title": title,
            "description": description,
            "sentences": sentences,
            "takeaways": takeaways,
            "sentence_count": len(sentences)
        }]
        
        result = {
            "success": True,
            "document_insights": {
                "metadata": {
                    "sentence_count": len(sentences),
                    "cluster_count": 1,
                    "noise_points": 0,
                    "language": language,
                    "processing_level": "full"
                },
                "summary": f"This content covers: {title}",
                "clusters": clusters_list,
                "connections": ""
            }
        }
        
        logger.info("Sentence understanding complete!")
        return result
        
    except Exception as e:
        logger.error(f"Error understanding sentences: {type(e).__name__}: {str(e)}")
        logger.error(f"Error details: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }

