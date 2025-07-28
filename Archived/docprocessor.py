import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
import hdbscan
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import logging
from docx import Document
import pdfplumber
from io import BytesIO
from typing import List, Tuple, Dict, Union, BinaryIO


# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the LLM client with timeout
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)
model = SentenceTransformer('sentence-transformers/LaBSE')
# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt, temperature=0.2, max_tokens=5000, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop, temperature=temperature, max_tokens=max_tokens)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

def load_and_split_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> List[str]:
    """
    Load a document from file path or BytesIO and split it into sentences.
    
    Args:
        input_source: Either a file path (str) or a BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        List of sentences
    """
    logger.info("Loading and splitting document.Filetype: %s", file_type)
    try:
        # Determine whether we have a file path or a file-like object
        if isinstance(input_source, str):
            # It's a file path
            file_path = input_source
            if not file_type:
                # Detect file type from extension if not provided
                file_type = file_path.split('.')[-1].lower()
                
            # Different handling based on file type
            if file_type == 'docx':
                logger.info(f"Loading DOCX from file path: {file_path}")
                doc = Document(file_path)
                
                # Extract paragraphs
                paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
                
                # Extract tables
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text for cell in row.cells]
                        table_rows.append(" | ".join(row_cells))
                    
                    if table_rows:  # Only add non-empty tables
                        tables_text.append(f"[Table {i+1}]\n" + "\n".join(table_rows))
                
                # Combine paragraphs and tables
                full_text = "\n\n".join(paragraphs_text + tables_text)
                logger.info(f"Extracted {len(paragraphs_text)} paragraphs and {len(tables_text)} tables")
            elif file_type == 'pdf':
                logger.info(f"Loading PDF from file path: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            # It's a file-like object (BytesIO)
            file_content = input_source
            
            # Reset position to beginning
            file_content.seek(0)
            
            # Must have file_type for file-like objects
            if not file_type:
                raise ValueError("file_type must be provided when using file-like input")
                
            if file_type == 'docx':
                logger.info("Loading DOCX from BytesIO/file-like object")
                doc = Document(file_content)
                
                # Extract paragraphs
                paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
                
                # Extract tables
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text for cell in row.cells]
                        table_rows.append(" | ".join(row_cells))
                    
                    if table_rows:  # Only add non-empty tables
                        tables_text.append(f"[Table {i+1}]\n" + "\n".join(table_rows))
                
                # Combine paragraphs and tables
                full_text = "\n\n".join(paragraphs_text + tables_text)
                logger.info(f"Extracted {len(paragraphs_text)} paragraphs and {len(tables_text)} tables")
            elif file_type == 'pdf':
                logger.info("Loading PDF from BytesIO/file-like object")
                with pdfplumber.open(file_content) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be either a file path or a file-like object")
            
        # Log text extraction info
        logger.info(f"Extracted {len(full_text)} characters of text")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")
        
        # Check if we have enough content
        if not full_text.strip():
            logger.warning("No text content extracted")
            raise ValueError("No text content could be extracted from the document")
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            raise ValueError("Text content too short for meaningful processing")
        
        # Process extracted text into sentences
        sentences = sent_tokenize(full_text)
        clean_sentences = [s.strip() for s in sentences if s.strip()]
        logger.info(f"Split text into {len(clean_sentences)} sentences")
        
        return clean_sentences
        
    except Exception as e:
        logger.error(f"Error loading document: {type(e).__name__}: {str(e)}")
        raise

def generate_sentence_embeddings(sentences: List[str]) -> np.ndarray:
    """
    Generate embeddings for sentences using LaBSE.
    
    Args:
        sentences: List of sentences to embed
        
    Returns:
        numpy array of embeddings
    """
    try:
        
        embeddings = model.encode(sentences, show_progress_bar=True)
        return embeddings
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise

def perform_clustering(embeddings: np.ndarray) -> np.ndarray:
    """
    Cluster sentences using HDBSCAN.
    
    Args:
        embeddings: Sentence embeddings
        
    Returns:
        Array of cluster labels
    """
    try:
        # Add special case handling for very few sentences
        if len(embeddings) == 1:
            logger.warning("Only 1 sentence available - skipping clustering")
            # Put the single point into cluster 0
            return np.zeros(1, dtype=int)
            
        # Special handling for exactly 2 sentences - still try to cluster but with modified parameters
        if len(embeddings) == 2:
            logger.warning("Only 2 sentences available - using simplified clustering")
            # Calculate similarity between the 2 sentences
            similarity = cosine_similarity(embeddings)[0, 1]
            logger.info(f"Similarity between the 2 sentences: {similarity:.4f}")
            
            # If they're similar enough, put them in the same cluster, otherwise different clusters
            if similarity > 0.5:  # Threshold can be adjusted
                logger.info("Sentences are similar - grouping in the same cluster")
                return np.zeros(2, dtype=int)
            else:
                logger.info("Sentences are different - putting in separate clusters")
                return np.array([0, 1], dtype=int)
        
        # For 3+ sentences, use standard HDBSCAN clustering
        # Further adjusted parameters for smaller, more focused clusters
        clusterer = hdbscan.HDBSCAN(
            min_cluster_size=2,           # Keep small clusters
            min_samples=1,                # Allow single points to form clusters
            metric='euclidean',
            cluster_selection_method='eom',
            cluster_selection_epsilon=0.25, # Further reduced to create even tighter clusters
            alpha=0.8,                    # Reduced to encourage more clusters
            prediction_data=True,         # Enable prediction data for better stability
            allow_single_cluster=False    # Prevent all points from being in one cluster
        )
        
        # Fit and predict
        cluster_labels = clusterer.fit_predict(embeddings)
        
        # Log clustering statistics
        unique_clusters = set(cluster_labels)
        noise_count = sum(1 for label in cluster_labels if label == -1)
        cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                        for label in unique_clusters if label != -1}
        
        logger.info(f"Clustering statistics:")
        logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
        logger.info(f"  Noise points: {noise_count}")
        logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        # If we have too few clusters, try again with different parameters
        if len(unique_clusters) - (1 if -1 in unique_clusters else 0) < 2 and len(embeddings) > 3:
            logger.info("Too few clusters detected, retrying with different parameters...")
            clusterer = hdbscan.HDBSCAN(
                min_cluster_size=2,
                min_samples=1,
                metric='euclidean',
                cluster_selection_method='leaf',
                cluster_selection_epsilon=0.15,
                alpha=0.6,
                prediction_data=True,
                allow_single_cluster=False
            )
            cluster_labels = clusterer.fit_predict(embeddings)
            
            # Log new clustering statistics
            unique_clusters = set(cluster_labels)
            noise_count = sum(1 for label in cluster_labels if label == -1)
            cluster_sizes = {label: sum(1 for l in cluster_labels if l == label) 
                            for label in unique_clusters if label != -1}
            
            logger.info(f"Retry clustering statistics:")
            logger.info(f"  Total clusters: {len(unique_clusters) - (1 if -1 in unique_clusters else 0)}")
            logger.info(f"  Noise points: {noise_count}")
            logger.info(f"  Cluster sizes: {cluster_sizes}")
        
        return cluster_labels
    except Exception as e:
        logger.error(f"Error clustering sentences: {type(e).__name__}: {str(e)}")
        raise

def select_all_sentences(
    sentences: List[str],
    cluster_labels: np.ndarray
) -> Dict[int, List[str]]:
    """
    Select all sentences for each cluster, allowing sentences to appear in multiple clusters.
    
    Args:
        sentences: List of sentences
        cluster_labels: Cluster labels for each sentence
        
    Returns:
        Dictionary mapping cluster IDs to lists of sentences
    """
    try:
        cluster_sentences = {}
        unique_clusters = set(cluster_labels)
        
        for cluster_id in unique_clusters:
            if cluster_id == -1:  # Skip noise points
                continue
                
            # Get all sentences for this cluster
            cluster_mask = cluster_labels == cluster_id
            cluster_sentences[cluster_id] = [
                sentences[i] for i in range(len(sentences)) if cluster_mask[i]
            ]
            
        return cluster_sentences
    except Exception as e:
        logger.error(f"Error selecting sentences: {str(e)}")
        raise

def extract_paragraphs(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> List[str]:
    """
    Load a document from file path or BytesIO and extract paragraphs.
    Unlike load_and_split_document, this function maintains paragraph structure rather than splitting into sentences.
    
    Args:
        input_source: Either a file path (str) or a BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        List of paragraphs
    """
    logger.info("Extracting paragraphs from document. Filetype: %s", file_type)
    try:
        # Determine whether we have a file path or a file-like object
        if isinstance(input_source, str):
            # It's a file path
            file_path = input_source
            if not file_type:
                # Detect file type from extension if not provided
                file_type = file_path.split('.')[-1].lower()
                
            # Different handling based on file type
            if file_type == 'docx':
                logger.info(f"Loading DOCX from file path: {file_path}")
                doc = Document(file_path)
                
                # Instead of treating each paragraph object as a separate paragraph,
                # we'll group related paragraphs based on context and structure
                
                # First, collect all paragraphs
                raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # Group related paragraphs - we'll use heuristics to determine what belongs together
                all_paragraphs = []
                current_paragraph = ""
                
                # Identifiers for headers, titles, and list items
                header_patterns = [
                    # Numbered section headers (e.g., "1. Introduction", "1.1 Definition")
                    lambda x: bool(any(x.startswith(f"{i}.") for i in range(1, 10)) or 
                                 any(x.startswith(f"{i}.{j}") for i in range(1, 10) for j in range(1, 10))),
                    # Short phrases that likely represent headings (< 60 chars and doesn't end with period)
                    lambda x: len(x) < 60 and not x.endswith('.') and not x.endswith(':')
                ]
                
                list_item_pattern = lambda x: bool(x.startswith('-') or x.startswith('•') or 
                                               any(x.startswith(f"{i}.") for i in range(1, 100)))
                
                for i, p in enumerate(raw_paragraphs):
                    is_header = any(pattern(p) for pattern in header_patterns)
                    is_list_item = list_item_pattern(p)
                    
                    # Headers and section titles get their own paragraph
                    if is_header:
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                        
                        # Add the header as its own paragraph
                        all_paragraphs.append(p)
                        
                    # List items are also separate paragraphs
                    elif is_list_item:
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                        
                        # Add the list item
                        all_paragraphs.append(p)
                        
                    # If this is a short line followed by more text, it might be a subheading
                    elif len(p) < 100 and i < len(raw_paragraphs) - 1 and p.endswith(':'):
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                            
                        # Start a new paragraph with this as a heading
                        current_paragraph = p + " "
                        
                    # Otherwise, it's content - append to the current paragraph
                    else:
                        # If the current paragraph is empty, start a new one
                        if not current_paragraph:
                            current_paragraph = p
                        else:
                            # Check if this seems to continue the previous paragraph
                            # If the previous paragraph ends with a period, question mark, or exclamation,
                            # and this one starts with a capital letter, it's likely a new thought
                            if (current_paragraph.endswith('.') or 
                                current_paragraph.endswith('?') or 
                                current_paragraph.endswith('!')) and p[0].isupper():
                                all_paragraphs.append(current_paragraph.strip())
                                current_paragraph = p
                            else:
                                # Otherwise, append to the current paragraph with a space
                                current_paragraph += " " + p
                
                # Add the last paragraph if there is one
                if current_paragraph:
                    all_paragraphs.append(current_paragraph.strip())
                
                # Extract tables (each table is treated as a paragraph)
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text for cell in row.cells]
                        table_rows.append(" | ".join(row_cells))
                    
                    if table_rows:  # Only add non-empty tables
                        tables_text.append(f"[Table {i+1}]\n" + "\n".join(table_rows))
                
                # Add tables to the paragraphs
                all_paragraphs.extend(tables_text)
                logger.info(f"Extracted and grouped into {len(all_paragraphs)} paragraphs and {len(tables_text)} tables")
                
            elif file_type == 'pdf':
                logger.info(f"Loading PDF from file path: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    # For PDFs, we need to be smarter about paragraph extraction
                    all_text = ""
                    
                    # First, extract all text from all pages
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            # Keep track of page number for reference
                            all_text += f"[Page {i+1}]\n{page_text}\n\n"
                    
                    # Now split the text into paragraphs by identifying patterns
                    # Double newlines often indicate paragraph breaks
                    initial_paragraphs = all_text.split("\n\n")
                    
                    # Further process these paragraphs to handle common PDF formatting issues
                    all_paragraphs = []
                    current_paragraph = ""
                    
                    for p in initial_paragraphs:
                        p = p.strip()
                        if not p:
                            continue
                            
                        # Check if this is a new page marker
                        if p.startswith("[Page "):
                            # Add previous paragraph if it exists
                            if current_paragraph:
                                all_paragraphs.append(current_paragraph.strip())
                                current_paragraph = ""
                            
                            # Extract page number
                            page_number = p.split("]")[0] + "]"
                            remaining_text = p[len(page_number):].strip()
                            
                            # Add the page marker to the current paragraph
                            current_paragraph = page_number + " "
                            
                            # If there's text after the page marker, add it
                            if remaining_text:
                                current_paragraph += remaining_text
                        
                        # If the paragraph is very short, it might be a heading
                        elif len(p) < 60 and not p.endswith('.'):
                            # Add previous paragraph if it exists
                            if current_paragraph:
                                all_paragraphs.append(current_paragraph.strip())
                            
                            # Add this as a standalone paragraph (likely a heading)
                            all_paragraphs.append(p)
                            current_paragraph = ""
                            
                        # Otherwise, it's content
                        else:
                            # If we have a current paragraph, check if this seems to continue it
                            if current_paragraph:
                                # Append with a space
                                current_paragraph += " " + p
                            else:
                                current_paragraph = p
                    
                    # Add the last paragraph if there is one
                    if current_paragraph:
                        all_paragraphs.append(current_paragraph.strip())
                
                logger.info(f"Extracted {len(all_paragraphs)} paragraphs from PDF")
                
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            # It's a file-like object (BytesIO)
            file_content = input_source
            
            # Reset position to beginning
            file_content.seek(0)
            
            # Must have file_type for file-like objects
            if not file_type:
                raise ValueError("file_type must be provided when using file-like input")
                
            if file_type == 'docx':
                logger.info("Loading DOCX from BytesIO/file-like object")
                doc = Document(file_content)
                
                # Use the same paragraph grouping logic as for file paths
                raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # Group related paragraphs - we'll use heuristics to determine what belongs together
                all_paragraphs = []
                current_paragraph = ""
                
                # Identifiers for headers, titles, and list items
                header_patterns = [
                    # Numbered section headers (e.g., "1. Introduction", "1.1 Definition")
                    lambda x: bool(any(x.startswith(f"{i}.") for i in range(1, 10)) or 
                                 any(x.startswith(f"{i}.{j}") for i in range(1, 10) for j in range(1, 10))),
                    # Short phrases that likely represent headings (< 60 chars and doesn't end with period)
                    lambda x: len(x) < 60 and not x.endswith('.') and not x.endswith(':')
                ]
                
                list_item_pattern = lambda x: bool(x.startswith('-') or x.startswith('•') or 
                                               any(x.startswith(f"{i}.") for i in range(1, 100)))
                
                for i, p in enumerate(raw_paragraphs):
                    is_header = any(pattern(p) for pattern in header_patterns)
                    is_list_item = list_item_pattern(p)
                    
                    # Headers and section titles get their own paragraph
                    if is_header:
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                        
                        # Add the header as its own paragraph
                        all_paragraphs.append(p)
                        
                    # List items are also separate paragraphs
                    elif is_list_item:
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                        
                        # Add the list item
                        all_paragraphs.append(p)
                        
                    # If this is a short line followed by more text, it might be a subheading
                    elif len(p) < 100 and i < len(raw_paragraphs) - 1 and p.endswith(':'):
                        # Add the previous paragraph if it exists
                        if current_paragraph:
                            all_paragraphs.append(current_paragraph.strip())
                            current_paragraph = ""
                            
                        # Start a new paragraph with this as a heading
                        current_paragraph = p + " "
                        
                    # Otherwise, it's content - append to the current paragraph
                    else:
                        # If the current paragraph is empty, start a new one
                        if not current_paragraph:
                            current_paragraph = p
                        else:
                            # Check if this seems to continue the previous paragraph
                            # If the previous paragraph ends with a period, question mark, or exclamation,
                            # and this one starts with a capital letter, it's likely a new thought
                            if (current_paragraph.endswith('.') or 
                                current_paragraph.endswith('?') or 
                                current_paragraph.endswith('!')) and p[0].isupper():
                                all_paragraphs.append(current_paragraph.strip())
                                current_paragraph = p
                            else:
                                # Otherwise, append to the current paragraph with a space
                                current_paragraph += " " + p
                
                # Add the last paragraph if there is one
                if current_paragraph:
                    all_paragraphs.append(current_paragraph.strip())
                
                # Extract tables (each table is treated as a paragraph)
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text for cell in row.cells]
                        table_rows.append(" | ".join(row_cells))
                    
                    if table_rows:  # Only add non-empty tables
                        tables_text.append(f"[Table {i+1}]\n" + "\n".join(table_rows))
                
                # Add tables to the paragraphs
                all_paragraphs.extend(tables_text)
                logger.info(f"Extracted and grouped into {len(all_paragraphs)} paragraphs and {len(tables_text)} tables")
                
            elif file_type == 'pdf':
                logger.info("Loading PDF from BytesIO/file-like object")
                with pdfplumber.open(file_content) as pdf:
                    # Use the same PDF processing logic as for file paths
                    all_text = ""
                    
                    # First, extract all text from all pages
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            # Keep track of page number for reference
                            all_text += f"[Page {i+1}]\n{page_text}\n\n"
                    
                    # Now split the text into paragraphs by identifying patterns
                    # Double newlines often indicate paragraph breaks
                    initial_paragraphs = all_text.split("\n\n")
                    
                    # Further process these paragraphs to handle common PDF formatting issues
                    all_paragraphs = []
                    current_paragraph = ""
                    
                    for p in initial_paragraphs:
                        p = p.strip()
                        if not p:
                            continue
                            
                        # Check if this is a new page marker
                        if p.startswith("[Page "):
                            # Add previous paragraph if it exists
                            if current_paragraph:
                                all_paragraphs.append(current_paragraph.strip())
                                current_paragraph = ""
                            
                            # Extract page number
                            page_number = p.split("]")[0] + "]"
                            remaining_text = p[len(page_number):].strip()
                            
                            # Add the page marker to the current paragraph
                            current_paragraph = page_number + " "
                            
                            # If there's text after the page marker, add it
                            if remaining_text:
                                current_paragraph += remaining_text
                        
                        # If the paragraph is very short, it might be a heading
                        elif len(p) < 60 and not p.endswith('.'):
                            # Add previous paragraph if it exists
                            if current_paragraph:
                                all_paragraphs.append(current_paragraph.strip())
                            
                            # Add this as a standalone paragraph (likely a heading)
                            all_paragraphs.append(p)
                            current_paragraph = ""
                            
                        # Otherwise, it's content
                        else:
                            # If we have a current paragraph, check if this seems to continue it
                            if current_paragraph:
                                # Append with a space
                                current_paragraph += " " + p
                            else:
                                current_paragraph = p
                    
                    # Add the last paragraph if there is one
                    if current_paragraph:
                        all_paragraphs.append(current_paragraph.strip())
                
                logger.info(f"Extracted {len(all_paragraphs)} paragraphs from PDF")
                
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be either a file path or a file-like object")
            
        # Log extraction info
        logger.info(f"Extracted {len(all_paragraphs)} total paragraphs")
        
        # Check if we have enough content
        if not all_paragraphs:
            logger.warning("No paragraphs extracted")
            raise ValueError("No paragraphs could be extracted from the document")
            
        # Filter out any empty paragraphs and clean whitespace
        clean_paragraphs = [p.strip() for p in all_paragraphs if p.strip()]
        logger.info(f"Final count: {len(clean_paragraphs)} non-empty paragraphs")
        
        return clean_paragraphs
        
    except Exception as e:
        logger.error(f"Error extracting paragraphs: {type(e).__name__}: {str(e)}")
        raise
