# Install required libraries
# pip install langchain sentence-transformers pinecone-client PyPDF2 python-docx pdfplumber scikit-learn nltk spacy
# python -m spacy download en_core_web_sm

import asyncio
import logging
import re
import numpy as np
from typing import List, Union, BinaryIO, Optional, Dict
from io import BytesIO
from docx import Document
import pdfplumber
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
import spacy
from utilities import logger
from pccontroller import save_knowledge
import json
import uuid

# Configure logging
logging.basicConfig(level=logging.INFO)

# Initialize sentence transformer (use lighter model for performance)
try:
    semantic_model = SentenceTransformer('all-MiniLM-L6-v2')  # Faster than LaBSE
    logger.info("Successfully loaded sentence transformer: all-MiniLM-L6-v2")
except Exception as e:
    logger.warning(f"Could not load sentence transformer: {e}. Falling back to simple chunking.")
    semantic_model = None

# Initialize spaCy for better sentence tokenization
try:
    nlp = spacy.load('en_core_web_sm', disable=['ner', 'lemmatizer'])
    logger.info("Successfully loaded spaCy model")
except Exception as e:
    logger.warning(f"Could not load spaCy model: {e}. Falling back to NLTK.")
    nlp = None

def preprocess_text(text: str) -> str:
    """Remove noise and normalize text."""
    # Remove multiple newlines and excessive whitespace
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r'\s+', ' ', text.strip())
    # Remove common boilerplate (e.g., page numbers, headers)
    text = re.sub(r'Page \d+', '', text)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    # Remove repeated punctuation
    text = re.sub(r'[\.\,\!]{2,}', '', text)
    return text

def load_document(input_source: Union[str, BytesIO, BinaryIO], file_type: str = None) -> str:
    """
    Load a document from file path or BytesIO and extract text content.
    
    Args:
        input_source: Either a file path (str) or a BytesIO/file-like object
        file_type: Optional file type override ('pdf', 'docx'), detected from path if not provided
        
    Returns:
        Extracted text content as string
    """
    logger.info(f"Loading document. File type: {file_type}")
    
    try:
        # Determine whether we have a file path or a file-like object
        if isinstance(input_source, str):
            file_path = input_source
            if not file_type:
                file_type = file_path.split('.')[-1].lower()
                
            if file_type == 'docx':
                logger.info(f"Loading DOCX from file path: {file_path}")
                doc = Document(file_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info(f"Loading PDF from file path: {file_path}")
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        elif hasattr(input_source, 'read') and hasattr(input_source, 'seek'):
            file_content = input_source
            file_content.seek(0)
            
            if not file_type:
                raise ValueError("file_type must be provided when using file-like input")
                
            if file_type == 'docx':
                logger.info("Loading DOCX from BytesIO/file-like object")
                doc = Document(file_content)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == 'pdf':
                logger.info("Loading PDF from BytesIO/file-like object")
                with pdfplumber.open(file_content) as pdf:
                    pages_text = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                            pages_text.append(f"[Page {i+1}] {page_text}")
                        else:
                            logger.debug(f"PDF page {i+1}: empty or non-text content")
                    full_text = "\n\n".join(pages_text)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            raise TypeError("input_source must be either a file path or a file-like object")
            
        full_text = preprocess_text(full_text)
        logger.info(f"Extracted {len(full_text)} characters of text after preprocessing")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")
        
        if not full_text.strip():
            logger.warning("No text content extracted")
            raise ValueError("No text content could be extracted from the document")
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            raise ValueError("Text content too short for meaningful processing")
        
        return full_text
        
    except Exception as e:
        logger.error(f"Error loading document: {type(e).__name__}: {str(e)}")
        raise

class SemanticChunker:
    """
    Principled semantic chunking using embeddings and topic coherence.
    """
    
    def __init__(self, model=None, similarity_threshold=0.7, coherence_threshold=0.5):
        self.model = model or semantic_model
        self.similarity_threshold = similarity_threshold
        self.coherence_threshold = coherence_threshold
    
    def _get_adaptive_thresholds(self, text_length: int) -> tuple:
        """Adjust thresholds based on document length."""
        if text_length < 1000:
            return self.similarity_threshold * 1.1, self.coherence_threshold * 1.1
        elif text_length > 10000:
            return self.similarity_threshold * 0.9, self.coherence_threshold * 0.9
        return self.similarity_threshold, self.coherence_threshold
    
    def _tokenize_sentences(self, text: str) -> List[str]:
        """Tokenize text into sentences using spaCy or NLTK."""
        if nlp:
            doc = nlp(text)
            return [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        else:
            try:
                sentences = nltk.sent_tokenize(text)
            except LookupError:
                nltk.download('punkt', quiet=True)
                sentences = nltk.sent_tokenize(text)
            return sentences
    
    def detect_semantic_boundaries(self, text: str, window_size: int = 3) -> List[int]:
        """
        Detect semantic boundaries using sentence embeddings and similarity analysis.
        """
        if not self.model:
            return self._detect_structural_boundaries(text)
        
        try:
            sentences = self._tokenize_sentences(text)
            if len(sentences) <= window_size:
                return []
            
            # Batch embedding generation for performance
            logger.debug(f"Generating embeddings for {len(sentences)} sentences")
            embeddings = self.model.encode(sentences, batch_size=32, show_progress_bar=False)
            
            boundaries = []
            sim_threshold, _ = self._get_adaptive_thresholds(len(text))
            
            for i in range(window_size, len(sentences) - window_size):
                current_window = embeddings[i-window_size:i]
                next_window = embeddings[i:i+window_size]
                
                current_centroid = np.mean(current_window, axis=0)
                next_centroid = np.mean(next_window, axis=0)
                
                similarity = cosine_similarity(
                    current_centroid.reshape(1, -1),
                    next_centroid.reshape(1, -1)
                )[0][0]
                
                if similarity < sim_threshold:
                    boundaries.append(i)
                    logger.debug(f"Semantic boundary detected at sentence {i} (similarity: {similarity:.3f})")
            
            logger.info(f"Detected {len(boundaries)} semantic boundaries")
            return boundaries
            
        except Exception as e:
            logger.error(f"Error in semantic boundary detection: {e}")
            return self._detect_structural_boundaries(text)
    
    def _detect_structural_boundaries(self, text: str) -> List[int]:
        """
        Fallback method using structural patterns.
        """
        boundaries = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            conditions = [
                re.match(r'^\d+\.?\s+', line_stripped),
                line_stripped.isupper() and len(line_stripped.split()) <= 10,
                re.match(r'^[\-\=\*]{3,}', line_stripped),
                line_stripped.endswith(':') and len(line_stripped.split()) <= 8,
                len(line) - len(line.lstrip()) > 4,
            ]
            
            if any(conditions):
                boundaries.append(i)
        
        return boundaries
    
    def calculate_coherence(self, sentences: List[str]) -> float:
        """
        Calculate topic coherence score for a group of sentences.
        """
        if not self.model or len(sentences) < 2:
            return 0.5
        
        try:
            embeddings = self.model.encode(sentences, batch_size=32, show_progress_bar=False)
            similarities = cosine_similarity(embeddings)
            
            n = len(similarities)
            mask = ~np.eye(n, dtype=bool)
            return float(np.mean(similarities[mask]))
            
        except Exception as e:
            logger.error(f"Error calculating coherence: {e}")
            return 0.5
    
    def adaptive_merge(self, chunks: List[Dict], min_size: int, max_size: int) -> List[Dict]:
        """
        Adaptively merge or split chunks based on size and coherence.
        """
        if not chunks:
            return chunks
        
        merged_chunks = []
        current_chunk = None
        
        for chunk in chunks:
            if current_chunk is None:
                current_chunk = chunk.copy()
                continue
            
            combined_content = current_chunk['content'] + ' ' + chunk['content']
            combined_size = len(combined_content)
            
            should_merge = False
            
            if combined_size <= max_size:
                current_sentences = self._tokenize_sentences(current_chunk['content'])
                next_sentences = self._tokenize_sentences(chunk['content'])
                combined_sentences = current_sentences + next_sentences
                
                combined_coherence = self.calculate_coherence(combined_sentences)
                avg_individual_coherence = (
                    current_chunk.get('coherence_score', 0.5) + 
                    chunk.get('coherence_score', 0.5)
                ) / 2
                
                if (combined_coherence >= avg_individual_coherence * 0.9 or 
                    current_chunk['char_count'] < min_size or 
                    chunk['char_count'] < min_size):
                    should_merge = True
            
            if should_merge:
                current_chunk['content'] = combined_content
                current_chunk['char_count'] = combined_size
                current_chunk['coherence_score'] = self.calculate_coherence(
                    self._tokenize_sentences(combined_content)
                )
            else:
                merged_chunks.append(current_chunk)
                current_chunk = chunk.copy()
        
        if current_chunk:
            merged_chunks.append(current_chunk)
        
        return merged_chunks
    
    def split_text(self, text: str, target_size: int = 1000, 
                   min_size: int = 200, max_size: int = 2000, 
                   overlap_sentences: int = 2) -> List[Dict[str, str]]:
        """
        Main semantic splitting method with overlap.
        """
        try:
            logger.info("Starting semantic text splitting")
            
            boundaries = self.detect_semantic_boundaries(text)
            sentences = self._tokenize_sentences(text)
            
            chunks = []
            current_sentences = []
            boundary_set = set(boundaries)
            
            for i, sentence in enumerate(sentences):
                current_sentences.append(sentence)
                
                if (i in boundary_set or 
                    len(' '.join(current_sentences)) >= target_size or 
                    i == len(sentences) - 1):
                    
                    if current_sentences:
                        chunk_text = ' '.join(current_sentences)
                        chunks.append({
                            'content': chunk_text,
                            'title': f"Semantic Chunk {len(chunks) + 1}",
                            'type': 'semantic_chunk',
                            'char_count': len(chunk_text),
                            'sentence_count': len(current_sentences),
                            'coherence_score': self.calculate_coherence(current_sentences)
                        })
                        current_sentences = current_sentences[-overlap_sentences:] if overlap_sentences > 0 else []
            
            final_chunks = self.adaptive_merge(chunks, min_size, max_size)
            
            processed_chunks = []
            for chunk in final_chunks:
                if chunk['char_count'] > max_size:
                    sub_chunks = self._split_large_chunk(chunk['content'], target_size)
                    processed_chunks.extend(sub_chunks)
                else:
                    processed_chunks.append(chunk)
            
            logger.info(f"Created {len(processed_chunks)} semantic chunks")
            return processed_chunks
            
        except Exception as e:
            logger.error(f"Error in semantic splitting: {e}")
            return self._fallback_split(text, target_size, min_size)
    
    def _split_large_chunk(self, text: str, target_size: int) -> List[Dict[str, str]]:
        """Split large chunks using sentence boundaries."""
        sentences = self._tokenize_sentences(text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > target_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'content': chunk_text,
                    'title': f"Split Chunk {len(chunks) + 1}",
                    'type': 'split_chunk',
                    'char_count': len(chunk_text),
                    'sentence_count': len(current_chunk),
                    'coherence_score': self.calculate_coherence(current_chunk)
                })
                current_chunk = []
                current_size = 0
            
            current_chunk.append(sentence)
            current_size += sentence_size
        
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'content': chunk_text,
                'title': f"Split Chunk {len(chunks) + 1}",
                'type': 'split_chunk',
                'char_count': len(chunk_text),
                'sentence_count': len(current_chunk),
                'coherence_score': self.calculate_coherence(current_chunk)
            })
        
        return chunks
    
    def _fallback_split(self, text: str, target_size: int, min_size: int) -> List[Dict[str, str]]:
        """Fallback splitting method."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=target_size,
            chunk_overlap=target_size // 5,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        return [{
            "content": chunk.strip(),
            "title": f"Fallback Chunk {i + 1}",
            "type": "fallback_chunk",
            "char_count": len(chunk),
            "sentence_count": len(self._tokenize_sentences(chunk)),
            "coherence_score": 0.5,
        } for i, chunk in enumerate(chunks)]

def save_to_local(chunks: List[Dict], user_id: str, bank_name: str, title: str) -> List[Dict]:
    """Fallback to save chunks locally as JSON."""
    try:
        os.makedirs(f"local_storage/{user_id}/{bank_name}", exist_ok=True)
        file_path = f"local_storage/{user_id}/{bank_name}/{title.replace(' ', '_')}_{uuid.uuid4()}.json"
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(chunks, f, indent=2, ensure_ascii=False)
        logger.info(f"Saved {len(chunks)} chunks locally to {file_path}")
        return [{"success": True, "file_path": file_path, "chunk_id": str(uuid.uuid4())} for _ in chunks]
    except Exception as e:
        logger.error(f"Failed to save chunks locally: {e}")
        return [{"success": False, "error": str(e)} for _ in chunks]

async def process_document(
    input_source: Union[str, BytesIO, BinaryIO], 
    user_id: str,
    bank_name: str = "documents",
    file_type: str = None,
    title: str = "",
    thread_id: Optional[str] = None,
    topic: Optional[str] = None,
    categories: Optional[List[str]] = None,
    ttl_days: Optional[int] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    use_semantic_chunking: bool = True,
    similarity_threshold: float = 0.7,
    coherence_threshold: float = 0.5,
    overlap_sentences: int = 2
) -> Dict:
    """
    Process document with principled semantic chunking approach.
    
    Args:
        input_source: Either file path (str) or BytesIO/file-like object
        user_id: Identifier for the user
        bank_name: Namespace for the knowledge (default: "documents")
        file_type: Optional file type override ('pdf', 'docx')
        title: Title for the document (optional)
        thread_id: Conversation thread identifier (optional)
        topic: Core topic of the knowledge (optional)
        categories: List of categories (optional)
        ttl_days: Time-to-live in days (optional)
        chunk_size: Target chunk size (default: 1000)
        chunk_overlap: Overlap for fallback method (default: 200)
        use_semantic_chunking: Whether to use semantic-aware chunking (default: True)
        similarity_threshold: Threshold for semantic boundary detection (default: 0.7)
        coherence_threshold: Minimum coherence score (default: 0.5)
        overlap_sentences: Number of sentences to overlap in semantic chunking (default: 2)
        
    Returns:
        Dictionary containing processing results and quality metrics
    """
    try:
        logger.info(f"Starting principled document processing for user_id={user_id}")
        
        text_content = load_document(input_source, file_type)
        
        if not title:
            if isinstance(input_source, str):
                title = os.path.splitext(os.path.basename(input_source))[0]
            else:
                title = f"Document processed at {topic or 'unknown topic'}"
        
        if use_semantic_chunking:
            chunker = SemanticChunker(
                similarity_threshold=similarity_threshold,
                coherence_threshold=coherence_threshold
            )
            chunks = chunker.split_text(
                text_content,
                target_size=chunk_size,
                min_size=chunk_size // 5,
                max_size=chunk_size * 2,
                overlap_sentences=overlap_sentences
            )
        else:
            chunks = split_text_recursive(text_content, chunk_size, chunk_overlap)
        
        if not chunks:
            return {
                "success": False,
                "error": "No chunks could be created from the document",
                "chunks_processed": 0,
                "chunks_saved": 0,
                "quality_metrics": {}
            }
        
        logger.info(f"Saving {len(chunks)} chunks to Pinecone...")
        
        saved_results = []
        failed_chunks = 0
        batch_size = 10
        
        for i in range(0, len(chunks), batch_size):
            batch_chunks = chunks[i:i + batch_size]
            batch_tasks = []
            
            for j, chunk_data in enumerate(batch_chunks):
                chunk_title = f"{title} - {chunk_data['title']}"
                enhanced_categories = (categories or []) + [
                    chunk_data['type'],
                    f"coherence_{chunk_data.get('coherence_score', 0.5):.1f}"
                ]
                
                task = save_knowledge(
                    input=chunk_data['content'],
                    user_id=user_id,
                    org_id="unknown",
                    title=chunk_title,
                    bank_name=bank_name,
                    thread_id=thread_id,
                    topic=topic,
                    categories=enhanced_categories,
                    ttl_days=ttl_days
                )
                batch_tasks.append(task)
            
            batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)
            
            for k, result in enumerate(batch_results):
                if isinstance(result, Exception) or not result.get("success", False):
                    failed_chunks += 1
                else:
                    saved_results.append(result)
        
        # Fallback to local storage if Pinecone fails
        if failed_chunks > 0:
            logger.warning(f"{failed_chunks} chunks failed to save to Pinecone. Using local storage.")
            local_results = save_to_local(chunks, user_id, bank_name, title)
            saved_results.extend([r for r in local_results if r.get("success")])
            failed_chunks = sum(1 for r in local_results if not r.get("success"))
        
        # Calculate quality metrics
        chunks_saved = len(saved_results)
        success_rate = chunks_saved / len(chunks) if chunks else 0
        avg_coherence = np.mean([c.get('coherence_score', 0.5) for c in chunks])
        avg_chunk_size = np.mean([c['char_count'] for c in chunks])
        chunk_types = {}
        for chunk in chunks:
            chunk_type = chunk['type']
            chunk_types[chunk_type] = chunk_types.get(chunk_type, 0) + 1
        
        quality_metrics = {
            "average_coherence": float(avg_coherence),
            "average_chunk_size": float(avg_chunk_size),
            "chunk_size_std": float(np.std([c['char_count'] for c in chunks])),
            "success_rate": float(success_rate),
            "chunk_type_distribution": chunk_types
        }
        
        logger.info(f"Processing complete: {chunks_saved}/{len(chunks)} saved ({success_rate:.1%})")
        logger.info(f"Quality metrics: {quality_metrics}")
        
        return {
            "success": chunks_saved > 0,
            "chunks_processed": len(chunks),
            "chunks_saved": chunks_saved,
            "chunks_failed": failed_chunks,
            "success_rate": success_rate,
            "title": title,
            "bank_name": bank_name,
            "semantic_chunking_used": use_semantic_chunking,
            "quality_metrics": quality_metrics,
            "saved_results": saved_results[:3] if saved_results else []
        }
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return {
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__,
            "chunks_processed": 0,
            "chunks_saved": 0,
            "quality_metrics": {}
        }

def split_text_recursive(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, str]]:
    """
    Simple recursive text splitting (fallback method).
    """
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        
        return [{
            "content": chunk.strip(),
            "title": f"Text Chunk {i + 1}",
            "type": "text_chunk",
            "char_count": len(chunk),
            "sentence_count": len(nlp(chunk).sents) if nlp else len(nltk.sent_tokenize(chunk)),
            "coherence_score": 0.5,
            "has_context": False
        } for i, chunk in enumerate(chunks)]
        
    except Exception as e:
        logger.error(f"Error in recursive splitting: {e}")
        raise

async def example_usage():
    """Example of principled semantic chunking."""
    result = await process_document(
        input_source="training.pdf",
        user_id="user123",
        bank_name="documents",
        topic="training_materials",
        use_semantic_chunking=True,
        similarity_threshold=0.75,
        coherence_threshold=0.6,
        chunk_size=1200,
        overlap_sentences=2
    )
    print(f"Semantic processing result: {result}")

if __name__ == "__main__":
    asyncio.run(example_usage())