import asyncio
import os
import uuid
from io import BytesIO

import matplotlib.pyplot as plt
import networkx as nx
from docx import Document

from graph_building import understand_document, visualize_knowledge_graph

async def test_multi_document_graph():
    """
    Test knowledge graph creation with multiple documents
    """
    print("Setting up documents...")
    
    # Check if Self-driving map.docx exists
    if not os.path.exists("Self-driving map.docx"):
        print("Warning: 'Self-driving map.docx' not found in current directory")
        print("Proceeding with only synthetic documents")
        use_real_doc = False
    else:
        use_real_doc = True
        print("Found 'Self-driving map.docx', will include it in the test")
    
    # Create document 1 - AI focused
    print("Creating synthetic documents...")
    doc1 = Document()
    doc1.add_heading('AI Technology Overview', 0)
    doc1.add_paragraph('Artificial intelligence systems are revolutionizing industries worldwide.')
    doc1.add_paragraph('Machine learning algorithms can process vast amounts of data for insights.')
    doc1.add_paragraph('Deep learning networks mimic human neural structures for complex tasks.')
    doc1.add_paragraph('Reinforcement learning enables AI to improve through experience and rewards.')
    doc1.add_paragraph('Natural language processing allows machines to understand and generate text.')
    
    # Create document 2 - Business focused
    doc2 = Document()
    doc2.add_heading('Business Applications of AI', 0)
    doc2.add_paragraph('Companies are leveraging AI for competitive advantages in the marketplace.')
    doc2.add_paragraph('Machine learning helps businesses predict customer behavior and preferences.')
    doc2.add_paragraph('Automation of routine tasks increases operational efficiency and reduces costs.')
    doc2.add_paragraph('Data analytics provides valuable insights for strategic decision making.')
    doc2.add_paragraph('AI adoption requires careful change management and employee training.')
    
    # Save documents to BytesIO objects
    docs = []
    doc_names = []
    
    # Add Self-driving map.docx as the first document if it exists
    if use_real_doc:
        docs.append("Self-driving map.docx")
        doc_names.append("Self-driving map.docx")
    
    # Add synthetic documents
    for i, doc in enumerate([doc1, doc2]):
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        docs.append(doc_bytes)
        doc_names.append(f"Synthetic document {i+1}")
    
    # Create unique graph file for this test
    test_graph_file = f"multi_doc_graph_{uuid.uuid4()}.graphml"
    
    # Process documents sequentially to build the knowledge graph
    for i, doc_input in enumerate(docs):
        print(f"Processing document {i+1}: {doc_names[i]}...")
        insights = await understand_document(
            input_source=doc_input,
            file_type='docx',
            existing_graph_path=test_graph_file
        )
        
        if not insights["success"]:
            print(f"Error processing document {i+1}: {insights.get('error', 'Unknown error')}")
            continue
            
        print(f"Document {i+1} processing succeeded!")
        print(f"Number of clusters: {insights['document_insights']['metadata']['cluster_count']}")
        
        # After each document, print the current graph state
        if os.path.exists(test_graph_file):
            graph = nx.read_graphml(test_graph_file)
            print(f"Current graph has {len(graph.nodes)} nodes and {len(graph.edges)} edges")
    
    # Visualization after all documents are processed
    if os.path.exists(test_graph_file):
        print("\nFinal graph analysis:")
        graph = nx.read_graphml(test_graph_file)
        
        # Get connected components
        components = list(nx.connected_components(graph))
        print(f"Graph has {len(components)} connected components")
        
        # Centrality measures
        degree_centrality = nx.degree_centrality(graph)
        top_nodes = sorted(degree_centrality.items(), key=lambda x: x[1], reverse=True)[:5]
        print("Top connected nodes:")
        for node, centrality in top_nodes:
            print(f"  - {node}: {graph.nodes[node].get('title', 'No title')} (centrality: {centrality:.3f})")
        
        # Create a visualization
        plt.figure(figsize=(16, 12))
        pos = nx.spring_layout(graph, k=0.5, iterations=50, seed=42)
        
        # Color nodes by their original document (based on node_id pattern)
        colors = []
        for node in graph.nodes():
            if "_new" not in node:  # Existing node
                colors.append("lightgray")
            else:
                node_num = int(node.split("_")[0])
                if node_num == 0:
                    colors.append("darkblue")  # Self-driving map.docx
                elif node_num == 1:
                    colors.append("lightgreen")  # AI Tech document
                else:
                    colors.append("salmon")  # Business document
        
        # Draw the graph with custom settings
        nx.draw_networkx_nodes(graph, pos, node_color=colors, node_size=1000, alpha=0.8, edgecolors='black')
        
        # Edge weights determine thickness
        edge_weights = [graph.edges[e].get('weight', 0.5) * 3 for e in graph.edges]
        nx.draw_networkx_edges(graph, pos, width=edge_weights, alpha=0.6, edge_color='gray')
        
        # Add labels
        node_labels = {n: graph.nodes[n].get('title', n)[:15] for n in graph.nodes}
        nx.draw_networkx_labels(graph, pos, labels=node_labels, font_size=9, font_weight='bold')
        
        # Create legend with document names
        legend_items = []
        if use_real_doc:
            legend_items.append(("darkblue", "Self-driving map.docx"))
        legend_items.extend([
            ("lightgreen", "AI Technology Doc"),
            ("salmon", "Business Applications Doc"),
            ("lightgray", "Existing Nodes")
        ])
        
        for color, label in legend_items:
            plt.plot([0], [0], 'o', color=color, label=label)
            
        plt.legend(loc='upper left', bbox_to_anchor=(1, 1))
        
        # Save the final visualization
        plt.title("Multi-Document Knowledge Graph")
        plt.axis('off')
        plt.tight_layout()
        
        output_path = "multi_doc_self_driving_map_graph.png"
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"Multi-document graph visualization saved to {output_path}")
        
        # Also save using the module's visualization function
        output_path2 = "module_viz_self_driving_map.png"
        visualize_knowledge_graph(graph, output_path2)
        print(f"Module's graph visualization saved to {output_path2}")
    else:
        print(f"Graph file {test_graph_file} not found")

if __name__ == "__main__":
    print("Starting multi-document graph test...")
    asyncio.run(test_multi_document_graph())
    print("Test completed!") 