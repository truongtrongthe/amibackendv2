import pdfplumber
from docx import Document
import uuid
from io import BytesIO
from utilities import logger
from database import save_training_with_chunk
import asyncio
from werkzeug.datastructures import FileStorage
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import openai
import tiktoken

# Initialize the LLM client with timeout
LLM = ChatOpenAI(model="gpt-4o", streaming=False, request_timeout=60)  # 60 second timeout
FAST_LLM = ChatOpenAI(model="gpt-4o-mini", streaming=False, request_timeout=45)  # 45 second timeout

# Retry decorator for LLM calls
@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=1, max=10),
    retry=retry_if_exception_type(openai.APITimeoutError)
)
async def invoke_llm_with_retry(llm, prompt, stop=None):
    """Call LLM with retry logic for timeouts and transient errors"""
    try:
        return await llm.ainvoke(prompt, stop=stop)
    except Exception as e:
        logger.warning(f"LLM call error: {type(e).__name__}: {str(e)}")
        raise

async def summarize_with_llm(text: str, max_length: int = 150) -> str:
    """Summarize text using the LLM, preserving key factual and actionable elements in the original language."""
    try:
        # Log the length of the text being summarized
        text_length = len(text)
        first_100_chars = text[:100].replace('\n', ' ') if text else "EMPTY TEXT"
        last_100_chars = text[-100:].replace('\n', ' ') if text_length > 100 else ""
        logger.debug(f"Summarizing text of length {text_length}. First 100 chars: '{first_100_chars}...'")
        if last_100_chars:
            logger.debug(f"Last 100 chars: '...{last_100_chars}'")
        
        if not text or text_length < 50:
            logger.warning(f"Text is too short or empty for summarization: '{text}'")
            return "Text provided is too short or empty for meaningful summarization."
        
        # Check if text appears to be a placeholder or metadata only
        if text.count('\n') < 3 and "Document:" in text and "Type:" in text:
            logger.warning(f"Text appears to be metadata only, no content: '{text}'")
            return "The provided text appears to contain only metadata without substantial content to summarize."
        
        # Detect language to explicitly tell the model to maintain it
        # This is a simple detection - the LLM will handle the actual language preservation    
        prompt = (
            f"You are a precise document summarizer specialized in knowledge extraction and preserving cultural and business nuance.\n\n"
            f"TASK: Create a concise summary of the following text in 250 words or less.\n\n"
            f"GUIDELINES:\n"
            f"1. Extract the most important FACTUAL information (e.g., names, numbers, dates).\n"
            f"2. Prioritize actionable knowledge (e.g., procedures, requirements, steps).\n"
            f"3. Preserve EXACT PHRASING for scripted instructions, questions, and sales prompts.\n"
            f"4. Maintain the EXACT ORDER of steps and processes as in the original text.\n"
            f"5. IMPORTANT: You MUST maintain the EXACT original language of the document. If the document is in Vietnamese or any non-English language, the summary MUST be in that same language. DO NOT translate to English under any circumstances.\n"
            f"6. Preserve all technical terms, proper names, and domain-specific vocabulary exactly as written.\n"
            f"7. Use clear, professional language without filler words.\n"
            f"8. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
            f"9. INCLUDE real-life lessons, practical advice, and applied knowledge from the text.\n"
            f"10. Maintain the emotional tone and perspective of the original document.\n\n"
            f"TEXT TO SUMMARIZE:\n{{text}}"
        )
        
        # Log prompt length for debugging
        logger.debug(f"Full prompt length: {len(prompt)} chars")
        
        response = await invoke_llm_with_retry(LLM, prompt)
        summary = response.content.strip()
        logger.info(f"Generated summary (approx {len(summary.split())} words): {summary}")
        return summary
    except Exception as e:
        logger.error(f"LLM summarization failed: {e}")
        return "Summary unavailable due to processing error."

async def extract_knowledge_with_llm(text: str, domain: str = "", max_items: int = 10) -> str:
    """Extract structured knowledge elements from text using the LLM while preserving the original language."""
    try:
        # Log the length of the text being analyzed
        text_length = len(text)
        first_100_chars = text[:100].replace('\n', ' ') if text else "EMPTY TEXT"
        logger.debug(f"Extracting knowledge from text of length {text_length}. First 100 chars: '{first_100_chars}...'")
        
        if not text or text_length < 50:
            logger.warning(f"Text is too short or empty for knowledge extraction: '{text}'")
            return "Text provided is too short or empty for meaningful knowledge extraction."
             
        prompt = (
                f"You are a precise knowledge extractor specialized in identifying actionable information{{domain_context}} while preserving cultural context and nuance.\n\n"
                f"TASK: Extract up to 20 key knowledge elements from the text below.\n\n"
                f"GUIDELINES:\n"
                f"1. Focus on FACTUAL and ACTIONABLE knowledge (e.g., procedures, requirements).\n"
                f"2. Include all specific questions, scripts, and sequential steps EXACTLY as written.\n"
                f"3. Include specific details (e.g., quantities, timelines, criteria).\n"
                f"4. IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all knowledge elements MUST be in that same language. DO NOT translate to English under any circumstances.\n"
                f"5. Maintain all technical terms, proper names, and specific vocabulary exactly as written.\n"
                f"6. Format each element as: 'KEY POINT: [concise statement]'.\n"
                f"7. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
                f"8. PRIORITIZE real-life lessons, practical advice, and applied knowledge.\n"
                f"9. Capture the underlying reasoning and wisdom, not just surface instructions.\n"
                f"10. Include both explicit statements and implied knowledge from the context.\n"
                f"11. Maintain the EXACT ORDER of steps and processes where applicable.\n\n"
                f"TEXT TO ANALYZE:\n{{text}}"
            )

        logger.debug(f"Knowledge extraction prompt length: {len(prompt)} chars")
        
        response = await invoke_llm_with_retry(LLM, prompt)
        extracted_knowledge = response.content.strip()
        logger.info(f"Extracted {extracted_knowledge.count('KEY POINT')} knowledge elements")
        return extracted_knowledge
    except Exception as e:
        logger.error(f"LLM knowledge extraction failed: {e}")
        return "Knowledge extraction unavailable due to processing error."

async def extract_knowledge_from_chunk(chunk_text: str, chunk_context: str = "") -> str:
    """Extract focused knowledge elements from a document chunk while preserving the original language."""
    try:
        prompt = (
            f"You are a precision knowledge extractor for a document segment, specializing in cultural nuance and contextual wisdom.\n\n"
            f"CONTEXT: This is a segment {chunk_context} from a larger document.\n\n"
            f"TASK: Extract 3-5 key knowledge elements from this segment, ensuring all critical details are captured.\n\n"
            f"GUIDELINES:\n"
            f"1. Prioritize CONCRETE, ACTIONABLE information (e.g., procedures, requirements, steps).\n"
            f"2. Include all specific questions, scripts, and sequential steps EXACTLY as written.\n"
            f"3. Capture specific details (e.g., names, numbers, dates, technical terms).\n"
            f"4. IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all knowledge elements MUST be in that same language. DO NOT translate to English under any circumstances.\n"
            f"5. Preserve all technical terms, proper names, and domain-specific vocabulary exactly as written.\n"
            f"6. Format each element as: 'KEY POINT: [concise statement]'.\n"
            f"7. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
            f"8. PRIORITIZE real-life lessons, practical advice, and applied knowledge.\n"
            f"9. Capture the underlying reasoning and implicit wisdom, but DO NOT add assumptions or enhancements beyond the text.\n"
            f"10. Maintain the EXACT ORDER of steps and processes as they appear in the segment.\n"
            f"11. If the segment contains partial instructions, include them verbatim and note their role in the context.\n\n"
            f"SEGMENT TEXT:\n{chunk_text}"
        )
        # Log a sample of the chunk text for debugging
        logger.debug(f"Processing chunk of length {len(chunk_text)}. First 50 chars: '{chunk_text[:50].replace(chr(10), ' ')}...'")
        
        response = await invoke_llm_with_retry(FAST_LLM, prompt)
        extracted_knowledge = response.content.strip()
        key_point_count = extracted_knowledge.count("KEY POINT")
        logger.debug(f"Extracted {key_point_count} key points from chunk")
        return extracted_knowledge
    except Exception as e:
        logger.error(f"Chunk knowledge extraction failed: {e}")
        return "Knowledge extraction failed for this segment."


async def refine_document(file: FileStorage = None, text: str = None, user_id: str = "", mode: str = "default", reformat_text: bool = False) -> tuple[bool, dict]:
    """
    Refine a document or text: extract content, summarize, extract knowledge, and optionally reformat as structured text.
    Returns (success: bool, result_dict: dict). Does not save to database.
    
    Args:
        file: Input file (DOCX or PDF), optional if text is provided.
        text: Raw text input, optional if file is provided.
        user_id: User ID for metadata.
        mode: Processing mode (e.g., "default", "pretrain").
        reformat_text: If True, return a structured text representation.
    """
    full_text = ""
    document_metadata = {}
    sections = []

    try:
        # Step 1: Extract text
        if text and text.strip():  # Ensure text is not just whitespace
            logger.debug(f"Processing text input of length {len(text)}")
            full_text = text
            document_metadata = {
                'filename': 'text_input',
                'file_type': 'TXT',
                'paragraph_count': len([p for p in text.split('\n') if p.strip()])
            }
            # Parse text sections (assuming possible Markdown or plain text)
            current_heading = "Document Start"
            current_section = []
            for line in text.split('\n'):
                if line.strip():
                    if line.startswith('# ') or line.startswith('## '):
                        if current_section:
                            sections.append((current_heading, "\n".join(current_section)))
                            current_section = []
                        current_heading = line.lstrip('# ').strip()
                    else:
                        current_section.append(line)
            if current_section:
                sections.append((current_heading, "\n".join(current_section)))
                
            # Verify that we have actual content not just metadata
            if len(sections) == 0 or (len(sections) == 1 and len(sections[0][1]) < 50):
                logger.warning(f"Insufficient text content: {text}")
                return False, {"error": "Insufficient text content."}
                
        elif file:
            file_extension = file.filename.split('.')[-1].lower()
            logger.debug(f"Processing file: {file.filename} ({file_extension})")
            
            # Read file content
            file_content = file.read()
            if not file_content:
                logger.warning(f"Empty file content for {file.filename}")
                return False, {"error": "Empty file content."}

            # Create a fresh BytesIO object for processing
            file_content_bytes = BytesIO(file_content)
            
            if file_extension == 'pdf':
                logger.debug("Opening PDF file")
                try:
                    with pdfplumber.open(file_content_bytes) as pdf:
                        document_metadata = {
                            'page_count': len(pdf.pages),
                            'filename': file.filename,
                            'file_type': 'PDF'
                        }
                        
                        # Extract text with page context
                        pages_text = []
                        for i, page in enumerate(pdf.pages):
                            page_text = page.extract_text() or ""
                            if page_text.strip():
                                logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                                pages_text.append(f"[Page {i+1}] {page_text}")
                            else:
                                logger.debug(f"PDF page {i+1}: empty or non-text content")
                        
                        # Join pages with clear separation
                        full_text = "\n\n".join(pages_text)
                        
                        if not full_text.strip():
                            logger.warning(f"No text extracted from {file.filename}")
                            return False, {"error": "No text extracted from PDF."}
                            
                        # Create sections from pages for better processing
                        sections = [(f"Page {i+1}", page) for i, page in enumerate(pages_text) if page.strip()]
                        logger.debug(f"Created {len(sections)} sections from PDF pages")
                except Exception as pdf_error:
                    logger.error(f"PDF processing error: {type(pdf_error).__name__}: {str(pdf_error)}")
                    return False, {"error": f"PDF processing error: {str(pdf_error)}"}
                    
            elif file_extension == 'docx':
                logger.debug("Opening DOCX file")
                try:
                    # Reset BytesIO position
                    file_content_bytes.seek(0)
                    doc = Document(file_content_bytes)
                    
                    # Extract document properties
                    document_metadata = {
                        'filename': file.filename,
                        'file_type': 'DOCX',
                        'paragraph_count': len(doc.paragraphs)
                    }
                    
                    # Extract text with structural context - similar to docuhandler approach
                    full_text_parts = []
                    current_heading = "Document Start"
                    current_section = []
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            if para.style.name.startswith('Heading'):
                                # Save current section if it exists
                                if current_section:
                                    sections.append((current_heading, "\n".join(current_section)))
                                    current_section = []
                                
                                current_heading = para.text.strip()
                                full_text_parts.append(f"\n[{current_heading}]\n{para.text}")
                            else:
                                current_section.append(para.text)
                                full_text_parts.append(para.text)
                    
                    # Don't forget the last section
                    if current_section:
                        sections.append((current_heading, "\n".join(current_section)))
                    
                    # Combine all the text parts
                    full_text = "\n".join(full_text_parts)
                    
                    if not full_text.strip():
                        logger.warning(f"No text extracted from {file.filename}")
                        return False, {"error": "No text extracted from DOCX."}
                except Exception as docx_error:
                    logger.error(f"DOCX processing error: {type(docx_error).__name__}: {str(docx_error)}")
                    return False, {"error": f"DOCX processing error: {str(docx_error)}"}
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return False, {"error": f"Unsupported file type: {file_extension}"}
        else:
            logger.error("No file or valid text provided")
            return False, {"error": "No file or valid text provided."}

        # Additional validation to ensure we have content to process
        if not full_text.strip():
            logger.warning("No text content extracted")
            return False, {"error": "No text content could be extracted."}
            
        if len(full_text.split()) < 10:
            logger.warning(f"Text content too short: '{full_text}'")
            return False, {"error": "Text content too short for meaningful processing."}

        # Detect language for logging and validation
        sample_text = full_text[:200]
        is_english = all(ord(c) < 128 for c in sample_text.replace('\n', ' ').replace(' ', ''))
        detected_language = "English" if is_english else "Non-English (possibly multilingual)"
        logger.info(f"Detected language: {detected_language}")

        logger.info(f"Extracted text content of {len(full_text.split())} words from {document_metadata.get('filename', 'unknown')}")
        logger.debug(f"First 100 chars: {full_text[:100].replace(chr(10), ' ')}...")

        # Step 2: Prepare contextual information - DO NOT include user_id or mode in the context
        context_prefix = f"Document: {document_metadata.get('filename', 'unknown')}\nType: {document_metadata.get('file_type', 'Unknown')}\n"
        if document_metadata.get('file_type') == 'PDF' and document_metadata.get('page_count'):
            context_prefix += f"Pages: {document_metadata.get('page_count')}\n\n"
        else:
            context_prefix += "\n"  # Add a blank line after metadata
        
        # Combine content for summarization but don't include mode or user_id
        combined_text = context_prefix + full_text
        logger.debug(f"Prepared text for summarization: prefix length {len(context_prefix)} chars, content length {len(full_text)} chars")
        
        # Step 3: Run summarization and knowledge extraction
        # For knowledge extraction, use just the full text without the metadata
        summary_task = summarize_with_llm(combined_text)
        knowledge_task = extract_knowledge_with_llm(full_text)
        summary, knowledge = await asyncio.gather(summary_task, knowledge_task)

        result = {
            "summary": summary,
            "knowledge_elements": knowledge,
            "metadata": document_metadata,
            "text_length": len(full_text.split())
        }

        # Step 4: Reformat as structured text if requested
        if reformat_text and sections:
            async def reformat_document(sections: list[tuple[str, str]]) -> str:
                """Reformat document sections into structured text with clear sections and subsections in original language."""
               
                prompt = (
                    f"You are a document reformatter tasked with creating a structured text output from raw document sections while preserving cultural nuance and contextual wisdom.\n\n"
                    f"INPUT SECTIONS:\n"
                    f"{{'\n'.join([f'Section: {{h}}\nContent: {{c}}' for h, c in sections])}}\n\n"
                    f"TASK:\n"
                    f"1. Summarize each section concisely, keeping factual and actionable details.\n"
                    f"2. Split mixed topics into subsections with descriptive headings.\n"
                    f"3. Preserve VERBATIM PHRASING for all scripted instructions, questions, and sales prompts.\n"
                    f"4. Maintain the EXACT ORDER of steps and processes as in the original text.\n"
                    f"5. IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all content MUST be in that same language. DO NOT translate to English under any circumstances.\n"
                    f"6. Preserve all technical terms, proper names, and specific vocabulary exactly as written.\n"
                    f"7. Use professional language, fixing informal phrases and errors only if they do not alter scripted content.\n"
                    f"8. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
                    f"9. PRIORITIZE real-life lessons, practical advice, and applied knowledge.\n"
                    f"10. Maintain the emotional tone and perspective of the original document.\n"
                    f"11. Output as plain text with sections ('# Heading') and subsections ('## Subheading').\n\n"
                    f"OUTPUT FORMAT:\n"
                    f"# Section Heading\n## Subsection Heading\nContent\n## Subsection Heading\nContent\n\n"
                )
                response = await invoke_llm_with_retry(LLM, prompt)
                return response.content.strip()

            reformatted_text = await reformat_document(sections)
            result["reformatted_text"] = reformatted_text

        logger.info(f"Processing complete for {document_metadata.get('filename', 'unknown')}")
        return True, result

    except Exception as e:
        logger.error(f"Failed to refine document: {type(e).__name__}: {str(e)}")
        return False, {"error": f"Processing error occurred: {str(e)}"}
    finally:
        if file:
            file.close()

async def enrich_key_points(key_points: str, full_text: str) -> str:
    """Enrich key points with supporting data from the original document text."""
    try:
        is_english = all(ord(c) < 128 for c in full_text[:100].replace('\n', ' ').replace(' ', ''))
        language_instruction = "IMPORTANT: You MUST use the EXACT SAME LANGUAGE as the original document. If the document is in Vietnamese or any non-English language, all content MUST be in that same language. DO NOT translate to English under any circumstances." if not is_english else "Keep the document's original language."
        
        prompt = (
            f"You are a precise knowledge enhancer tasked with gathering supporting data for key points from a document, preserving cultural nuance and exact phrasing.\n\n"
            f"TASK: For each key point below, extract clear, concise supporting data from the original document text that provides context, descriptions, or relevant details. The supporting data must include all critical facts, scripts, or steps directly related to the key point.\n\n"
            f"GUIDELINES:\n"
            f"1. Extract supporting data VERBATIM from the original text, prioritizing exact scripts, questions, or instructions.\n"
            f"2. Include ALL relevant details (e.g., names, numbers, dates, sequential steps) tied to the key point, ensuring completeness.\n"
            f"3. Keep supporting data concise, targeting 50-100 words, but include all critical information without a strict sentence limit.\n"
            f"4. {language_instruction}\n"
            f"5. Preserve all technical terms, proper names, and domain-specific vocabulary exactly as written.\n"
            f"6. Maintain the EXACT ORDER of steps or processes as they appear in the original text.\n"
            f"7. If the key point references partial instructions, include surrounding text to clarify context.\n"
            f"8. PRESERVE the NUANCE and CULTURAL CONTEXT of the original text, including tone and perspective.\n"
            f"9. DO NOT add assumptions, enhancements, or inferred details beyond the original text.\n"
            f"10. Output format: 'KEY POINT: [original key point]\nSUPPORTING DATA: [verbatim text from document]'.\n\n"
            f"KEY POINTS:\n{key_points}\n\n"
            f"ORIGINAL TEXT:\n{full_text}"
        )
        
        logger.debug(f"Enriching key points with prompt length: {len(prompt)} chars")
        response = await invoke_llm_with_retry(LLM, prompt)
        enriched_knowledge = response.content.strip()
        logger.info(f"Enriched {enriched_knowledge.count('KEY POINT')} key points with supporting data")
        return enriched_knowledge
    except Exception as e:
        logger.error(f"Key point enrichment failed: {e}")
        return "Key point enrichment failed due to processing error."

async def process_document(text: str = "", file: FileStorage = None, user_id: str = "", mode: str = "default", bank: str = "") -> bool:
    """
    Process a document or text: split into semantic chunks, extract knowledge, and save to Pinecone.
    
    Args:
        text: Raw text input, preferred input method.
        file: Input file (DOCX or PDF), optional for legacy support.
        user_id: User ID for metadata.
        mode: Processing mode (e.g., "default", "pretrain").
        bank: Namespace for Pinecone storage.
    """
    doc_id = str(uuid.uuid4())
    chunks = []
    logger.info(f"Processing document with bank={bank}")

    try:
        # Step 1: Extract text and create paragraphs
        if text and text.strip():  # Ensure text is not just whitespace
            logger.debug(f"Processing text input of length {len(text)}")
            full_text = text
            
            # Check for minimal content
            if len(text.split()) < 10:
                logger.warning(f"Text content too short: '{text}'")
                return False
                
            # Parse text with Markdown headings
            paragraphs = []
            current_heading = "Introduction"
            current_subheading = ""
            for line in full_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('# '):
                    current_heading = line[2:].strip()
                    current_subheading = ""
                elif line.startswith('## '):
                    current_subheading = line[3:].strip()
                else:
                    context = f"section: {current_heading}"
                    if current_subheading:
                        context += f", subsection: {current_subheading}"
                    paragraphs.append((line, context))
                    
            logger.debug(f"Extracted {len(paragraphs)} paragraphs from text input")
            
        elif file:
            file_extension = file.filename.split('.')[-1].lower()
            logger.debug(f"Processing file: {file.filename} ({file_extension})")
            
            # Read file content
            file_content = file.read()
            if not file_content:
                logger.warning(f"Empty file content for {file.filename}")
                return False
            
            logger.debug(f"Successfully read {len(file_content)} bytes from {file.filename}")
            
            # Create a fresh BytesIO object for processing
            file_content_bytes = BytesIO(file_content)

            if file_extension == 'pdf':
                logger.debug("Opening PDF file")
                try:
                    with pdfplumber.open(file_content_bytes) as pdf:
                        page_count = len(pdf.pages)
                        logger.debug(f"PDF processing: found {page_count} pages")
                        all_text = []
                        
                        for i, page in enumerate(pdf.pages):
                            page_text = page.extract_text() or ""
                            if page_text.strip():
                                logger.debug(f"PDF page {i+1}: extracted {len(page_text)} characters")
                                all_text.append(page_text)
                            else:
                                logger.debug(f"PDF page {i+1}: empty or non-text content")
                                
                        if not all_text:
                            logger.warning(f"No text extracted from {file.filename}")
                            return False
                                
                        full_text = " ".join(all_text)
                        logger.debug(f"Combined PDF text: {len(full_text)} characters")
                        
                        # Split the text into paragraphs
                        paragraphs = []
                        for para in full_text.split('\n'):
                            if para.strip():
                                paragraphs.append((para.strip(), f"PDF document, page info not preserved in combined text"))
                        
                        logger.debug(f"Extracted {len(paragraphs)} paragraphs from PDF")
                except Exception as pdf_error:
                    logger.error(f"PDF processing error: {type(pdf_error).__name__}: {str(pdf_error)}")
                    return False
                            
            elif file_extension == 'docx':
                logger.debug("Opening DOCX file")
                try:
                    # Reset BytesIO position
                    file_content_bytes.seek(0)
                    doc = Document(file_content_bytes)
                    paragraphs = []
                    current_heading = "Introduction"
                    
                    for para in doc.paragraphs:
                        if not para.text.strip():
                            continue
                        if para.style.name.startswith('Heading'):
                            current_heading = para.text.strip()
                            # Include headings as their own paragraphs with context
                            paragraphs.append((para.text.strip(), "heading"))
                        else:
                            paragraphs.append((para.text.strip(), f"section: {current_heading}"))
                    
                    logger.debug(f"Extracted {len(paragraphs)} paragraphs from DOCX")
                    
                    if not paragraphs:
                        logger.warning(f"No paragraphs extracted from {file.filename}")
                        return False
                        
                except Exception as docx_error:
                    logger.error(f"DOCX processing error: {type(docx_error).__name__}: {str(docx_error)}")
                    return False
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return False
        else:
            logger.error("No file or valid text provided")
            return False

        # Check if we have paragraphs to process
        if not paragraphs:
            logger.warning("No paragraphs extracted from input")
            return False
            
        logger.info(f"Successfully extracted {len(paragraphs)} paragraphs for processing")

        # Step 2: Group paragraphs into semantic chunks
        current_section = []
        current_word_count = 0
        section_index = 0
        max_chunk_words = 300
        min_chunk_words = 50

        for para_text, para_context in paragraphs:
            para_words = len(para_text.split())
            
            # Skip empty or very small paragraphs
            if para_words < 2:
                continue
                
            # Detect implicit headings (e.g., lines with ":" or short capitalized phrases)
            is_implicit_heading = ':' in para_text or (para_text[0].isupper() and len(para_text.split()) < 10)
            if is_implicit_heading and current_section:
                # Save current section as a chunk
                if current_word_count >= min_chunk_words:
                    section_text = " ".join(current_section)
                    chunk_id = f"chunk_{section_index}_0"
                    chunks.append((chunk_id, section_text, para_context))
                    section_index += 1
                current_section = [para_text]
                current_word_count = para_words
                continue

            # Handle large paragraphs
            if para_words > max_chunk_words:
                words = para_text.split()
                for j in range(0, len(words), max_chunk_words - 50):
                    sub_chunk = " ".join(words[j:j + (max_chunk_words - 50)])
                    chunk_id = f"chunk_{section_index}_{j // (max_chunk_words - 50)}"
                    chunks.append((chunk_id, sub_chunk, f"{para_context}, large paragraph part {j // (max_chunk_words - 50) + 1}"))
                section_index += 1
                continue

            # Add to current section
            if current_word_count + para_words <= max_chunk_words or current_word_count < min_chunk_words:
                current_section.append(para_text)
                current_word_count += para_words
            else:
                # Save current section
                section_text = " ".join(current_section)
                chunk_id = f"chunk_{section_index}_0"
                chunks.append((chunk_id, section_text, para_context))
                section_index += 1
                current_section = [para_text]
                current_word_count = para_words

        # Save any remaining section
        if current_section and current_word_count >= min_chunk_words:
            section_text = " ".join(current_section)
            chunk_id = f"chunk_{section_index}_0"
            chunks.append((chunk_id, section_text, "final section"))
            
        if not chunks:
            logger.warning("No chunks created from paragraphs")
            return False
            
        logger.info(f"Created {len(chunks)} chunks for processing")

        # Step 3: Process chunks and extract knowledge
        processing_tasks = []
        knowledge_extraction_tasks = []
        
        for chunk_id, chunk_text, chunk_context in chunks:
            # Log the first chunk to help with debugging
            if chunk_id.endswith("_0"):
                logger.debug(f"Sample chunk {chunk_id}: {chunk_text[:100]}...")
                
            knowledge_extraction_tasks.append(
                extract_knowledge_from_chunk(chunk_text, chunk_context)
            )
            processing_tasks.append(
                save_training_with_chunk(
                    input=chunk_text,
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=True
                )
            )
            processing_tasks.append(
                save_training_with_chunk(
                    input=chunk_text,  # Placeholder for knowledge
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunk_id,
                    bank_name=bank,
                    is_raw=False
                )
            )

        # Run knowledge extraction
        logger.debug(f"Starting knowledge extraction for {len(knowledge_extraction_tasks)} chunks")
        knowledge_results = await asyncio.gather(*knowledge_extraction_tasks, return_exceptions=True)
        
        # Replace placeholder tasks with knowledge
        for i, knowledge in enumerate(knowledge_results):
            if isinstance(knowledge, Exception):
                logger.error(f"Knowledge extraction failed for chunk {i}: {str(knowledge)}")
                knowledge = f"Knowledge extraction failed: {str(knowledge)}"
            task_index = (i * 2) + 1
            if task_index < len(processing_tasks):
                processing_tasks[task_index] = save_training_with_chunk(
                    input=knowledge,
                    user_id=user_id,
                    mode=mode,
                    doc_id=doc_id,
                    chunk_id=chunks[i][0],
                    bank_name=bank,
                    is_raw=False
                )

        # Process save tasks
        logger.debug(f"Starting data storage for {len(processing_tasks)} tasks")
        results = await asyncio.gather(*processing_tasks, return_exceptions=True)
        
        success_count = 0
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Processing task failed: {result}")
            elif not result:
                logger.warning("Processing task returned False")
            else:
                success_count += 1
        
        if success_count == 0:
            logger.error("All processing tasks failed")
            return False
            
        logger.info(f"Processed document {doc_id} with {len(chunks)} chunks ({success_count}/{len(processing_tasks)} tasks succeeded)")
        return True

    except Exception as e:
        logger.error(f"Failed to process document: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False
    finally:
        if file:
            file.close()