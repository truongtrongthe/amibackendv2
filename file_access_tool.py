"""
File Access Tool - Handle both local and Google Drive file operations
Supports reading DOCX, PDF files from local filesystem and Google Drive
"""

import os
import logging
from typing import List, Dict, Any, Optional
import json

logger = logging.getLogger(__name__)

# Configure file access specific logging
file_logger = logging.getLogger("file_access")
file_logger.setLevel(logging.INFO)
if not file_logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter('📁 [FILE] %(asctime)s - %(message)s', datefmt='%H:%M:%S')
    handler.setFormatter(formatter)
    file_logger.addHandler(handler)


class FileAccessTool:
    """Tool for accessing files from local filesystem and Google Drive"""
    
    def __init__(self):
        self.name = "file_access_tool"
        self.google_drive_service = None
        self._initialize_google_drive()
    
    def _initialize_google_drive(self):
        """Initialize Google Drive API service"""
        try:
            from googleapiclient.discovery import build
            from google.oauth2.service_account import Credentials
            
            # Try to load service account credentials
            creds_path = os.getenv('GOOGLE_DRIVE_CREDENTIALS_PATH', 'google_drive_credentials.json')
            
            # If default path doesn't exist, try to find any amiai-*.json file
            if not os.path.exists(creds_path):
                import glob
                amiai_files = glob.glob('amiai-*.json')
                if amiai_files:
                    creds_path = amiai_files[0]  # Use the first found file
                    file_logger.info(f"Found Google Drive credentials: {creds_path}")
            
            if os.path.exists(creds_path):
                credentials = Credentials.from_service_account_file(
                    creds_path,
                    scopes=['https://www.googleapis.com/auth/drive.readonly']
                )
                self.google_drive_service = build('drive', 'v3', credentials=credentials)
                logger.info("Google Drive service initialized successfully")
            else:
                logger.warning(f"Google Drive credentials not found at {creds_path}")
                
        except Exception as e:
            logger.error(f"Failed to initialize Google Drive service: {e}")
            self.google_drive_service = None
    
    # LOCAL FILE OPERATIONS
    
    def read_local_docx(self, file_path: str) -> str:
        """
        Read DOCX file from local filesystem
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Text content of the DOCX file
        """
        try:
            from docx import Document
            
            # Security: Ensure file is within project directory
            abs_path = os.path.abspath(file_path)
            if not abs_path.startswith(os.getcwd()):
                return "Error: File access denied - path outside project directory"
            
            if not os.path.exists(file_path):
                return f"Error: File not found at {file_path}"
            
            doc = Document(file_path)
            content = []
            
            # Read paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Read tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_content.append(" | ".join(row_text))
                content.extend(table_content)
            
            return "\n".join(content)
            
        except ImportError:
            return "Error: python-docx library not installed. Run: pip install python-docx"
        except Exception as e:
            return f"Error reading local DOCX file: {str(e)}"
    
    def read_local_pdf(self, file_path: str) -> str:
        """
        Read PDF file from local filesystem
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Text content of the PDF file
        """
        try:
            import PyPDF2
            
            # Security check
            abs_path = os.path.abspath(file_path)
            if not abs_path.startswith(os.getcwd()):
                return "Error: File access denied - path outside project directory"
            
            if not os.path.exists(file_path):
                return f"Error: File not found at {file_path}"
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = []
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content.append(f"--- Page {page_num + 1} ---")
                        content.append(page_text)
                
                return "\n".join(content)
                
        except ImportError:
            return "Error: PyPDF2 library not installed. Run: pip install PyPDF2"
        except Exception as e:
            return f"Error reading local PDF file: {str(e)}"
    
    def list_local_files(self, directory: str) -> str:
        """
        List files in local directory
        
        Args:
            directory: Directory path to list
            
        Returns:
            List of files in the directory
        """
        try:
            # Security check
            abs_dir = os.path.abspath(directory)
            if not abs_dir.startswith(os.getcwd()):
                return "Error: Directory access denied - path outside project directory"
            
            if not os.path.exists(directory):
                return f"Error: Directory not found at {directory}"
            
            files = []
            for item in os.listdir(directory):
                item_path = os.path.join(directory, item)
                if os.path.isfile(item_path):
                    size = os.path.getsize(item_path)
                    files.append(f"📄 {item} ({self._format_file_size(size)})")
                elif os.path.isdir(item_path):
                    files.append(f"📁 {item}/")
            
            if not files:
                return f"No files found in {directory}"
            
            return f"Files in {directory}:\n" + "\n".join(files)
            
        except Exception as e:
            return f"Error listing local files: {str(e)}"
    
    # GOOGLE DRIVE OPERATIONS
    
    def list_gdrive_files(self, folder_id: str = None, folder_name: str = None) -> str:
        """
        List files in Google Drive folder
        
        Args:
            folder_id: Google Drive folder ID (optional)
            folder_name: Folder name to search for (optional)
            
        Returns:
            List of files in the Google Drive folder
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # If folder_name provided, find the folder ID first
            if folder_name and not folder_id:
                folder_query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
                folder_results = self.google_drive_service.files().list(
                    q=folder_query,
                    fields="files(id, name)"
                ).execute()
                
                folders = folder_results.get('files', [])
                if not folders:
                    return f"Error: Folder '{folder_name}' not found in Google Drive"
                
                folder_id = folders[0]['id']
                logger.info(f"Found folder '{folder_name}' with ID: {folder_id}")
            
            # List files in the folder
            if folder_id:
                query = f"'{folder_id}' in parents and trashed=false"
            else:
                query = "trashed=false"
            
            results = self.google_drive_service.files().list(
                q=query,
                fields="files(id, name, mimeType, size, modifiedTime)",
                orderBy="name"
            ).execute()
            
            files = results.get('files', [])
            
            if not files:
                return "No files found in Google Drive folder"
            
            file_list = []
            for file in files:
                name = file['name']
                mime_type = file['mimeType']
                size = file.get('size', 'Unknown')
                modified = file.get('modifiedTime', 'Unknown')
                
                # Format file info
                if 'folder' in mime_type:
                    file_list.append(f"📁 {name}/")
                elif 'document' in mime_type:
                    file_list.append(f"📄 {name} (Google Doc)")
                elif 'pdf' in mime_type:
                    file_list.append(f"📄 {name} (PDF, {self._format_file_size(int(size) if size != 'Unknown' else 0)})")
                else:
                    file_list.append(f"📄 {name} ({mime_type.split('/')[-1]})")
            
            return f"Google Drive files:\n" + "\n".join(file_list)
            
        except Exception as e:
            return f"Error listing Google Drive files: {str(e)}"
    
    def read_gdrive_docx(self, file_id: str = None, file_name: str = None, folder_id: str = None) -> str:
        """
        Read DOCX file from Google Drive
        
        Args:
            file_id: Google Drive file ID (optional)
            file_name: File name to search for (optional)
            folder_id: Folder ID to search in (optional)
            
        Returns:
            Text content of the DOCX file
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # If file_name provided, find the file ID first
            if file_name and not file_id:
                if folder_id:
                    query = f"name='{file_name}' and '{folder_id}' in parents and trashed=false"
                else:
                    query = f"name='{file_name}' and trashed=false"
                
                results = self.google_drive_service.files().list(
                    q=query,
                    fields="files(id, name, mimeType)"
                ).execute()
                
                files = results.get('files', [])
                if not files:
                    return f"Error: File '{file_name}' not found in Google Drive"
                
                file_id = files[0]['id']
                logger.info(f"Found file '{file_name}' with ID: {file_id}")
            
            if not file_id:
                return "Error: Either file_id or file_name must be provided"
            
            # Get file metadata to check MIME type
            file_metadata = self.google_drive_service.files().get(
                fileId=file_id,
                fields="id,name,mimeType"
            ).execute()
            
            mime_type = file_metadata.get('mimeType', '')
            file_logger.info(f"READ_GDRIVE_DOCX - File MIME type: {mime_type}")
            
            # Handle different file types
            if mime_type == 'application/vnd.google-apps.document':
                # Google Docs - export as DOCX
                file_logger.info(f"READ_GDRIVE_DOCX - Exporting Google Docs as DOCX")
                file_content = self.google_drive_service.files().export(
                    fileId=file_id,
                    mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ).execute()
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Native DOCX - download directly
                file_logger.info(f"READ_GDRIVE_DOCX - Downloading native DOCX")
                file_content = self.google_drive_service.files().get_media(fileId=file_id).execute()
            else:
                return f"Error: Unsupported file type: {mime_type}. Only Google Docs and DOCX files are supported."
            
            # Save temporarily and read with python-docx
            import tempfile
            from docx import Document
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                content = []
                
                # Read paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                # Read tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        content.append(" | ".join(row_text))
                
                content_text = "\n".join(content)
                logger.info(f"READ_GDRIVE_DOCX - Successfully read {len(content_text)} characters from file")
                return content_text
                
            finally:
                # Clean up temp file
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temp file {temp_file_path}: {cleanup_error}")
            
        except ImportError:
            return "Error: python-docx library not installed. Run: pip install python-docx"
        except Exception as e:
            logger.error(f"READ_GDRIVE_DOCX - Exception occurred: {str(e)}")
            logger.error(f"READ_GDRIVE_DOCX - Exception type: {type(e).__name__}")
            
            # Handle specific Google Drive errors
            if "404" in str(e) and "File not found" in str(e):
                return f"Error: File not found or access denied. Please ensure the file is shared with the service account. File ID: {file_id}"
            elif "403" in str(e) and "Forbidden" in str(e):
                return f"Error: Access denied. Please check file sharing permissions for file ID: {file_id}"
            elif "401" in str(e) and "Unauthorized" in str(e):
                return f"Error: Authentication failed. Please check Google Drive credentials."
            
            import traceback
            logger.error(f"READ_GDRIVE_DOCX - Traceback: {traceback.format_exc()}")
            return f"Error reading Google Drive DOCX file: {str(e)}"
    
    def read_gdrive_pdf(self, file_id: str = None, file_name: str = None, folder_id: str = None) -> str:
        """
        Read PDF file from Google Drive
        
        Args:
            file_id: Google Drive file ID (optional)
            file_name: File name to search for (optional)
            folder_id: Folder ID to search in (optional)
            
        Returns:
            Text content of the PDF file
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # If file_name provided, find the file ID first
            if file_name and not file_id:
                if folder_id:
                    query = f"name='{file_name}' and '{folder_id}' in parents and trashed=false"
                else:
                    query = f"name='{file_name}' and trashed=false"
                
                results = self.google_drive_service.files().list(
                    q=query,
                    fields="files(id, name, mimeType)"
                ).execute()
                
                files = results.get('files', [])
                if not files:
                    return f"Error: File '{file_name}' not found in Google Drive"
                
                file_id = files[0]['id']
                logger.info(f"Found file '{file_name}' with ID: {file_id}")
            
            if not file_id:
                return "Error: Either file_id or file_name must be provided"
            
            # Download the file
            file_content = self.google_drive_service.files().get_media(fileId=file_id).execute()
            
            # Save temporarily and read with PyPDF2
            import tempfile
            import PyPDF2
            
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                with open(temp_file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    content = []
                    
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            content.append(f"--- Page {page_num + 1} ---")
                            content.append(page_text)
                    
                    return "\n".join(content)
                    
            finally:
                # Clean up temp file
                os.unlink(temp_file_path)
            
        except ImportError:
            return "Error: PyPDF2 library not installed. Run: pip install PyPDF2"
        except Exception as e:
            return f"Error reading Google Drive PDF file: {str(e)}"
    
    def read_gdrive_link_docx(self, drive_link: str) -> str:
        """
        Read DOCX file from Google Drive using a sharing link
        
        Args:
            drive_link: Google Drive sharing link (e.g., https://drive.google.com/file/d/ABC123/view)
            
        Returns:
            Text content of the DOCX file
        """
        file_logger.info(f"READ_GDRIVE_LINK_DOCX - Starting with link: {drive_link}")
        
        if not self.google_drive_service:
            file_logger.error("READ_GDRIVE_LINK_DOCX - Google Drive service not initialized")
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # Extract file ID from the link
            file_logger.info(f"READ_GDRIVE_LINK_DOCX - Extracting file ID from link")
            file_id = self._extract_file_id_from_link(drive_link)
            if not file_id:
                file_logger.error(f"READ_GDRIVE_LINK_DOCX - Could not extract file ID from: {drive_link}")
                return f"Error: Could not extract file ID from link: {drive_link}"
            
            file_logger.info(f"READ_GDRIVE_LINK_DOCX - Extracted file ID: {file_id}")
            
            # Use existing method with extracted file ID
            file_logger.info(f"READ_GDRIVE_LINK_DOCX - Reading file with ID: {file_id}")
            result = self.read_gdrive_docx(file_id=file_id)
            
            if result.startswith("Error:"):
                file_logger.error(f"READ_GDRIVE_LINK_DOCX - File read failed: {result[:100]}")
            else:
                content_length = len(result)
                file_logger.info(f"READ_GDRIVE_LINK_DOCX - Successfully read {content_length} characters from file")
            
            return result
            
        except Exception as e:
            file_logger.error(f"READ_GDRIVE_LINK_DOCX - Exception: {str(e)}")
            file_logger.error(f"READ_GDRIVE_LINK_DOCX - Exception type: {type(e).__name__}")
            import traceback
            file_logger.error(f"READ_GDRIVE_LINK_DOCX - Traceback: {traceback.format_exc()}")
            return f"Error reading Google Drive DOCX from link: {str(e)}"
    
    def read_gdrive_link_pdf(self, drive_link: str) -> str:
        """
        Read PDF file from Google Drive using a sharing link
        
        Args:
            drive_link: Google Drive sharing link (e.g., https://drive.google.com/file/d/ABC123/view)
            
        Returns:
            Text content of the PDF file
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # Extract file ID from the link
            file_id = self._extract_file_id_from_link(drive_link)
            if not file_id:
                return f"Error: Could not extract file ID from link: {drive_link}"
            
            # Use existing method with extracted file ID
            return self.read_gdrive_pdf(file_id=file_id)
            
        except Exception as e:
            return f"Error reading Google Drive PDF from link: {str(e)}"
    
    def get_file_info_from_link(self, drive_link: str) -> str:
        """
        Get file information from Google Drive sharing link
        
        Args:
            drive_link: Google Drive sharing link
            
        Returns:
            File information including name, type, size
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # Extract file ID from the link
            file_id = self._extract_file_id_from_link(drive_link)
            if not file_id:
                return f"Error: Could not extract file ID from link: {drive_link}"
            
            # Get file metadata
            file_metadata = self.google_drive_service.files().get(
                fileId=file_id,
                fields="id,name,mimeType,size,modifiedTime,owners"
            ).execute()
            
            info_parts = []
            info_parts.append(f"📄 File Information:")
            info_parts.append(f"   Name: {file_metadata.get('name', 'Unknown')}")
            info_parts.append(f"   ID: {file_metadata.get('id', 'Unknown')}")
            info_parts.append(f"   Type: {file_metadata.get('mimeType', 'Unknown')}")
            
            if 'size' in file_metadata:
                size = int(file_metadata['size'])
                info_parts.append(f"   Size: {self._format_file_size(size)}")
            
            if 'modifiedTime' in file_metadata:
                info_parts.append(f"   Modified: {file_metadata['modifiedTime']}")
            
            if 'owners' in file_metadata and file_metadata['owners']:
                owner = file_metadata['owners'][0]
                info_parts.append(f"   Owner: {owner.get('displayName', 'Unknown')}")
            
            return "\n".join(info_parts)
            
        except Exception as e:
            return f"Error getting file info from link: {str(e)}"
    
    def _extract_file_id_from_link(self, drive_link: str) -> str:
        """
        Extract file ID from various Google Drive link formats
        
        Args:
            drive_link: Google Drive sharing link
            
        Returns:
            File ID or None if not found
        """
        import re
        
        file_logger.info(f"EXTRACT_FILE_ID - Processing link: {drive_link}")
        
        # Remove any whitespace
        drive_link = drive_link.strip()
        
        # Pattern 1: https://drive.google.com/file/d/FILE_ID/view?usp=sharing
        pattern1 = r'https://drive\.google\.com/file/d/([a-zA-Z0-9-_]+)'
        match1 = re.search(pattern1, drive_link)
        if match1:
            file_id = match1.group(1)
            file_logger.info(f"EXTRACT_FILE_ID - Matched Pattern 1 (drive.google.com/file): {file_id}")
            return file_id
        
        # Pattern 2: https://docs.google.com/document/d/FILE_ID/edit?usp=sharing
        pattern2 = r'https://docs\.google\.com/document/d/([a-zA-Z0-9-_]+)'
        match2 = re.search(pattern2, drive_link)
        if match2:
            file_id = match2.group(1)
            file_logger.info(f"EXTRACT_FILE_ID - Matched Pattern 2 (docs.google.com/document): {file_id}")
            return file_id
        
        # Pattern 3: https://drive.google.com/open?id=FILE_ID
        pattern3 = r'https://drive\.google\.com/open\?id=([a-zA-Z0-9-_]+)'
        match3 = re.search(pattern3, drive_link)
        if match3:
            file_id = match3.group(1)
            file_logger.info(f"EXTRACT_FILE_ID - Matched Pattern 3 (drive.google.com/open): {file_id}")
            return file_id
        
        # Pattern 4: Direct file ID (if someone just pastes the ID)
        pattern4 = r'^[a-zA-Z0-9-_]{25,}$'
        if re.match(pattern4, drive_link):
            file_logger.info(f"EXTRACT_FILE_ID - Matched Pattern 4 (direct file ID): {drive_link}")
            return drive_link
        
        # If no pattern matches
        file_logger.warning(f"EXTRACT_FILE_ID - No pattern matched for link: {drive_link}")
        logger.warning(f"Could not extract file ID from link: {drive_link}")
        return None

    def find_file_by_name(self, file_name: str, folder_name: str = None) -> str:
        """
        Find file information by name in Google Drive
        
        Args:
            file_name: Name of the file to find
            folder_name: Optional folder name to search within
            
        Returns:
            File information including ID and location
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # Build search query
            query = f"name='{file_name}' and trashed=false"
            
            if folder_name:
                # First find the folder
                folder_query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
                folder_results = self.google_drive_service.files().list(
                    q=folder_query,
                    fields="files(id, name)"
                ).execute()
                
                folders = folder_results.get('files', [])
                if folders:
                    folder_id = folders[0]['id']
                    query += f" and '{folder_id}' in parents"
            
            results = self.google_drive_service.files().list(
                q=query,
                fields="files(id, name, mimeType, size, parents)"
            ).execute()
            
            files = results.get('files', [])
            
            if not files:
                location = f" in folder '{folder_name}'" if folder_name else ""
                return f"File '{file_name}' not found{location}"
            
            file_info = []
            for file in files:
                info = f"📄 {file['name']}\n"
                info += f"   ID: {file['id']}\n"
                info += f"   Type: {file['mimeType']}\n"
                if 'size' in file:
                    info += f"   Size: {self._format_file_size(int(file['size']))}\n"
                file_info.append(info)
            
            return f"Found {len(files)} file(s):\n" + "\n".join(file_info)
            
        except Exception as e:
            return f"Error finding file: {str(e)}"
    
    def read_gdrive_folder(self, folder_id: str = None, folder_name: str = None, file_types: List[str] = None, max_chars: int = 50000) -> str:
        """
        Read all supported files from a Google Drive folder
        
        Args:
            folder_id: Google Drive folder ID (optional)
            folder_name: Folder name to search for (optional)
            file_types: List of file types to read (default: ['docx', 'pdf'])
            max_chars: Maximum characters to return (default: 50000 to avoid token limits)
            
        Returns:
            Combined text content of all supported files in the folder
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        # Default file types to read
        if file_types is None:
            file_types = ['docx', 'pdf']
        
        try:
            # If folder_name provided, find the folder ID first
            if folder_name and not folder_id:
                folder_query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
                folder_results = self.google_drive_service.files().list(
                    q=folder_query,
                    fields="files(id, name)"
                ).execute()
                
                folders = folder_results.get('files', [])
                if not folders:
                    return f"Error: Folder '{folder_name}' not found in Google Drive"
                
                folder_id = folders[0]['id']
                logger.info(f"Found folder '{folder_name}' with ID: {folder_id}")
            
            if not folder_id:
                return "Error: Either folder_id or folder_name must be provided"
            
            # List all files in the folder
            query = f"'{folder_id}' in parents and trashed=false"
            results = self.google_drive_service.files().list(
                q=query,
                fields="files(id, name, mimeType)",
                orderBy="name"
            ).execute()
            
            files = results.get('files', [])
            
            if not files:
                return "No files found in Google Drive folder"
            
            # Filter files by supported types
            supported_files = []
            for file in files:
                mime_type = file.get('mimeType', '')
                file_name = file.get('name', '').lower()
                
                # Check if file type is supported
                is_supported = False
                if 'docx' in file_types:
                    if (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
                        mime_type == 'application/vnd.google-apps.document' or
                        file_name.endswith('.docx')):
                        is_supported = True
                
                if 'pdf' in file_types:
                    if (mime_type == 'application/pdf' or
                        file_name.endswith('.pdf')):
                        is_supported = True
                
                if is_supported:
                    supported_files.append(file)
            
            if not supported_files:
                return f"No supported files found in folder. Supported types: {', '.join(file_types)}"
            
            # Read content from all supported files
            folder_content = []
            folder_content.append(f"=== FOLDER CONTENTS ===\n")
            folder_content.append(f"Reading {len(supported_files)} supported files:\n")
            
            total_chars = 0
            files_read = 0
            files_truncated = 0
            
            for i, file in enumerate(supported_files, 1):
                file_id = file['id']
                file_name = file['name']
                mime_type = file.get('mimeType', '')
                
                folder_content.append(f"\n--- FILE {i}: {file_name} ---")
                
                try:
                    # Determine file type and read accordingly
                    if (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
                        mime_type == 'application/vnd.google-apps.document' or
                        file_name.lower().endswith('.docx')):
                        
                        content = self.read_gdrive_docx(file_id=file_id)
                        if not content.startswith("Error:"):
                            files_read += 1
                            content_length = len(content)
                            total_chars += content_length
                            
                            # Check if we need to truncate this file
                            if total_chars > max_chars:
                                # Truncate this file to fit within limits
                                remaining_chars = max_chars - (total_chars - content_length)
                                if remaining_chars > 100:  # Only include if we have meaningful content
                                    content = content[:remaining_chars] + "\n\n[CONTENT TRUNCATED - File too large for complete analysis]"
                                    files_truncated += 1
                                else:
                                    folder_content.append(f"[FILE TOO LARGE - Skipped to avoid token limits]")
                                    continue
                            
                            folder_content.append(content)
                        else:
                            folder_content.append(f"Error reading file: {content}")
                    
                    elif (mime_type == 'application/pdf' or
                          file_name.lower().endswith('.pdf')):
                        
                        content = self.read_gdrive_pdf(file_id=file_id)
                        if not content.startswith("Error:"):
                            files_read += 1
                            content_length = len(content)
                            total_chars += content_length
                            
                            # Check if we need to truncate this file
                            if total_chars > max_chars:
                                # Truncate this file to fit within limits
                                remaining_chars = max_chars - (total_chars - content_length)
                                if remaining_chars > 100:  # Only include if we have meaningful content
                                    content = content[:remaining_chars] + "\n\n[CONTENT TRUNCATED - File too large for complete analysis]"
                                    files_truncated += 1
                                else:
                                    folder_content.append(f"[FILE TOO LARGE - Skipped to avoid token limits]")
                                    continue
                            
                            folder_content.append(content)
                        else:
                            folder_content.append(f"Error reading file: {content}")
                    
                    else:
                        folder_content.append(f"Unsupported file type: {mime_type}")
                
                except Exception as e:
                    folder_content.append(f"Error reading {file_name}: {str(e)}")
            
            # Add summary information
            summary = f"\n\n=== FOLDER READING SUMMARY ===\n"
            summary += f"Files successfully read: {files_read}\n"
            summary += f"Files truncated due to size: {files_truncated}\n"
            summary += f"Total characters processed: {total_chars:,}\n"
            summary += f"Character limit: {max_chars:,}\n"
            
            if files_truncated > 0:
                summary += f"⚠️  Some files were truncated to avoid token limits. Consider reading individual files for complete analysis.\n"
            
            folder_content.append(summary)
            
            return "\n".join(folder_content)
            
        except Exception as e:
            return f"Error reading Google Drive folder: {str(e)}"
    
    def check_file_access(self, file_id: str = None, drive_link: str = None) -> str:
        """
        Check if a file is accessible and get basic information
        
        Args:
            file_id: Google Drive file ID (optional)
            drive_link: Google Drive sharing link (optional)
            
        Returns:
            File access status and information
        """
        if not self.google_drive_service:
            return "Error: Google Drive service not initialized. Check credentials."
        
        try:
            # Extract file ID from link if provided
            if drive_link and not file_id:
                file_id = self._extract_file_id_from_link(drive_link)
                if not file_id:
                    return f"Error: Could not extract file ID from link: {drive_link}"
            
            if not file_id:
                return "Error: Either file_id or drive_link must be provided"
            
            # Try to get file metadata
            try:
                file_metadata = self.google_drive_service.files().get(
                    fileId=file_id,
                    fields="id,name,mimeType,size,owners,permissions"
                ).execute()
                
                # File is accessible
                info_parts = []
                info_parts.append(f"✅ File Access: SUCCESS")
                info_parts.append(f"📄 File Information:")
                info_parts.append(f"   Name: {file_metadata.get('name', 'Unknown')}")
                info_parts.append(f"   ID: {file_metadata.get('id', 'Unknown')}")
                info_parts.append(f"   Type: {file_metadata.get('mimeType', 'Unknown')}")
                
                if 'size' in file_metadata:
                    size = int(file_metadata['size'])
                    info_parts.append(f"   Size: {self._format_file_size(size)}")
                
                if 'owners' in file_metadata and file_metadata['owners']:
                    owner = file_metadata['owners'][0]
                    info_parts.append(f"   Owner: {owner.get('displayName', 'Unknown')}")
                
                info_parts.append(f"   Status: File is accessible and ready to read")
                
                return "\n".join(info_parts)
                
            except Exception as e:
                if "404" in str(e) and "File not found" in str(e):
                    return f"❌ File Access: DENIED\n\nError: File not found or access denied.\nFile ID: {file_id}\n\nTo fix this:\n1. Ensure the file exists\n2. Share the file with the service account\n3. Check file permissions"
                elif "403" in str(e) and "Forbidden" in str(e):
                    return f"❌ File Access: DENIED\n\nError: Access forbidden.\nFile ID: {file_id}\n\nTo fix this:\n1. Share the file with the service account\n2. Grant 'Viewer' permissions\n3. Check if file is in shared drive with proper access"
                elif "401" in str(e) and "Unauthorized" in str(e):
                    return f"❌ File Access: DENIED\n\nError: Authentication failed.\n\nTo fix this:\n1. Check Google Drive credentials\n2. Verify service account setup\n3. Ensure API is enabled"
                else:
                    return f"❌ File Access: ERROR\n\nUnexpected error: {str(e)}\nFile ID: {file_id}"
                    
        except Exception as e:
            return f"Error checking file access: {str(e)}"
    
    # UTILITY METHODS
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        import math
        i = int(math.floor(math.log(size_bytes, 1024)))
        p = math.pow(1024, i)
        s = round(size_bytes / p, 2)
        return f"{s} {size_names[i]}"
    
    def get_tool_description(self):
        """Return tool descriptions for LLM function calling"""
        return [
            {
                "name": "read_local_docx",
                "description": "Read DOCX file from local filesystem",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to the local DOCX file"
                        }
                    },
                    "required": ["file_path"]
                }
            },
            {
                "name": "read_local_pdf", 
                "description": "Read PDF file from local filesystem",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to the local PDF file"
                        }
                    },
                    "required": ["file_path"]
                }
            },
            {
                "name": "list_local_files",
                "description": "List files in local directory",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "directory": {
                            "type": "string",
                            "description": "Local directory path to list"
                        }
                    },
                    "required": ["directory"]
                }
            },
            {
                "name": "list_gdrive_files",
                "description": "List files in Google Drive folder",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "folder_id": {
                            "type": "string",
                            "description": "Google Drive folder ID (optional)"
                        },
                        "folder_name": {
                            "type": "string", 
                            "description": "Folder name to search for (optional)"
                        }
                    }
                }
            },
            {
                "name": "read_gdrive_docx",
                "description": "Read DOCX file from Google Drive",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "Google Drive file ID (optional)"
                        },
                        "file_name": {
                            "type": "string",
                            "description": "File name to search for (optional)"
                        },
                        "folder_id": {
                            "type": "string",
                            "description": "Folder ID to search in (optional)"
                        }
                    }
                }
            },
            {
                "name": "read_gdrive_pdf",
                "description": "Read PDF file from Google Drive", 
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "Google Drive file ID (optional)"
                        },
                        "file_name": {
                            "type": "string",
                            "description": "File name to search for (optional)"
                        },
                        "folder_id": {
                            "type": "string",
                            "description": "Folder ID to search in (optional)"
                        }
                    }
                }
            },
            {
                "name": "read_gdrive_link_docx",
                "description": "CRITICAL: Read DOCX file from Google Drive using a sharing link. Use this when user provides a Google Drive URL to analyze document content.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "drive_link": {
                            "type": "string",
                            "description": "Google Drive sharing link (e.g., https://drive.google.com/file/d/ABC123/view or https://docs.google.com/document/d/ABC123/edit)"
                        }
                    },
                    "required": ["drive_link"]
                }
            },
            {
                "name": "read_gdrive_link_pdf",
                "description": "CRITICAL: Read PDF file from Google Drive using a sharing link. Use this when user provides a Google Drive URL to analyze PDF document content.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "drive_link": {
                            "type": "string",
                            "description": "Google Drive sharing link (e.g., https://drive.google.com/file/d/ABC123/view or https://docs.google.com/document/d/ABC123/edit)"
                        }
                    },
                    "required": ["drive_link"]
                }
            },
            {
                "name": "get_file_info_from_link",
                "description": "Get file information from Google Drive sharing link",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "drive_link": {
                            "type": "string",
                            "description": "Google Drive sharing link"
                        }
                    },
                    "required": ["drive_link"]
                }
            },
            {
                "name": "find_file_by_name",
                "description": "Find file information by name in Google Drive",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_name": {
                            "type": "string",
                            "description": "Name of the file to find"
                        },
                        "folder_name": {
                            "type": "string",
                            "description": "Optional folder name to search within"
                        }
                    },
                    "required": ["file_name"]
                }
            },
            {
                "name": "read_gdrive_folder",
                "description": "Read all supported files (DOCX, PDF) from a Google Drive folder. Use this when you need to analyze multiple files from a folder.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "folder_id": {
                            "type": "string",
                            "description": "Google Drive folder ID (optional)"
                        },
                        "folder_name": {
                            "type": "string",
                            "description": "Folder name to search for (optional)"
                        },
                        "file_types": {
                            "type": "array",
                            "items": {
                                "type": "string",
                                "description": "List of file types to read (e.g., ['docx', 'pdf'])"
                            },
                            "description": "List of file types to read (default: ['docx', 'pdf'])"
                        },
                        "max_chars": {
                            "type": "integer",
                            "description": "Maximum characters to return (default: 50000 to avoid token limits)"
                        }
                    }
                }
            },
            {
                "name": "check_file_access",
                "description": "Check if a Google Drive file is accessible and get file information. Use this to diagnose permission issues before reading files.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "Google Drive file ID (optional)"
                        },
                        "drive_link": {
                            "type": "string",
                            "description": "Google Drive sharing link (optional)"
                        }
                    }
                }
            }
        ] 